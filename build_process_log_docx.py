"""Sinh file Word ghi lại toàn bộ quy trình hoàn thiện bài RG-SCSO (Scientific
Reports) — nhật ký thực hiện từ đầu phiên làm việc đến hiện tại, phục vụ mục
đích tham khảo nội bộ của tác giả, KHÔNG phải một phần bản thảo nộp tạp chí.

Chạy: python build_process_log_docx.py
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT_DOCX = "RG-SCSO_Quy_Trinh_Thuc_Hien.docx"


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def build() -> None:
    doc = Document()

    title = doc.add_heading("Quy trình hoàn thiện bài báo RG-SCSO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Nhật ký thực hiện — bản Scientific Reports "
                     "(Relevance-Guided Sand Cat Swarm Optimization)")
    r.italic = True
    r.font.size = Pt(13)

    note = doc.add_paragraph()
    r = note.add_run(
        "Tài liệu nội bộ, ghi lại toàn bộ quy trình từ lúc bắt đầu đợt sửa "
        "bài đến hiện tại — không phải một phần bản thảo nộp tạp chí. Mục "
        "đích: giúp tác giả (và người đọc lại sau này) hiểu rõ đã làm gì, "
        "tại sao làm, và kết quả cụ thể ra sao ở mỗi bước, theo đúng trình "
        "tự thời gian thực tế."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ------------------------------------------------------------ Giai đoạn 0
    h1(doc, "Giai đoạn 0 — Điểm xuất phát")
    para(
        doc,
        "Có hai tài liệu phản biện độc lập được dùng làm căn cứ cho toàn bộ "
        "đợt sửa bài:",
    )
    bullet(
        doc,
        " — một bản phản biện đầy đủ theo chuẩn Area Chair/Senior Reviewer "
        "Q1, kết luận Weak Reject (độ tin cậy 0,86), gồm 13 điểm yếu đánh số "
        "W1–W13 và 6 nhóm “Required Revisions” ưu tiên (Priority 1–6).",
        bold_lead="RG-SCSO_Q1_Review_Final.md",
    )
    bullet(
        doc,
        " — một bản tự phản biện (adversarial self-review) riêng, ban đầu "
        "được viết cho bản Applied Intelligence (đánh số mục khác hẳn), đối "
        "chiếu với 10 bài SCSO/swarm Q1 gần đây; phân loại điểm yếu theo 3 "
        "mức TỬ HUYỆT / LỚN / PHỤ.",
        bold_lead="Diem_yeu_RG-SCSO.md",
    )
    para(
        doc,
        " — cả hai tài liệu đều chỉ áp dụng cho bản Scientific Reports "
        "(build_paper_scirep.py / .docx); bản Applied Intelligence "
        "(build_paper_tex.py) được giữ nguyên, không đụng vào trong suốt "
        "toàn bộ quá trình.",
        bold_lead="Quyết định phạm vi: ",
    )

    # ------------------------------------------------------------ Giai đoạn 1
    h1(doc, "Giai đoạn 1 — Phân loại điểm yếu và sửa nhanh (“Loại A”)")
    para(
        doc,
        "Trước khi sửa bất cứ gì, mọi điểm yếu trong RG-SCSO_Q1_Review_Final.md "
        "được chia làm hai loại:",
    )
    bullet(doc, " — sửa được ngay bằng văn bản/cách trình bày, không cần số liệu mới.",
           bold_lead="Loại A: ")
    bullet(doc, " — cần chạy thí nghiệm thật mới trả lời được; nguyên tắc "
           "tuyệt đối: không được bịa số liệu, chỉ báo cáo khi có kết quả "
           "thật, kể cả khi kết quả đó bất lợi cho bài.",
           bold_lead="Loại B: ")
    para(doc, "Các sửa Loại A đã áp dụng ngay ở bước này:")
    bullet(doc, "W5/W10 — đổi khung câu chuyện từ “thắng về accuracy” "
                "sang “parsimony ở accuracy ngang bằng”; đổi cách trình bày "
                "“wins 107/108” thành “paired per-dataset comparisons”.")
    bullet(doc, "W2 — thêm đoạn “Scope of this result” ngay sau Lemma 1, "
                "nói rõ đây là sensitivity bound cục bộ, không phải bằng "
                "chứng cho toàn bộ hiện tượng washout.")
    bullet(doc, "W3 — thêm cảnh báo ngưỡng ρ=0.5 không có cơ sở lý thuyết "
                "phổ quát, phụ thuộc số lớp/entropy/cỡ mẫu/estimator MI.")
    bullet(doc, "W4 — hạ mức độ nổi bật của UMR trong phần đóng góp, nêu rõ "
                "RMS mới là centerpiece có ablation xác nhận.")
    bullet(doc, "W12 — đổi tên gọi “deployment speedup” thành “projected, "
                "reconstructed inference-cost analysis” cho đúng bản chất.")

    # ------------------------------------------------------------ Giai đoạn 2
    h1(doc, "Giai đoạn 2 — 5 thí nghiệm mới (Loại B, đợt 1)")
    para(
        doc,
        "Khảo sát hạ tầng thí nghiệm sẵn có trước, xác nhận chưa có sẵn "
        "harness cho các so sánh Priority 1–3 và 6, rồi thiết kế + viết mới "
        "5 pipeline thí nghiệm, mỗi pipeline đều có bước --smoke kiểm tra "
        "wiring trước khi chạy full-scale nền (background, nhiều giờ):",
    )
    add_table(
        doc,
        ["Priority", "Nội dung", "Quy mô", "Phát hiện chính"],
        [
            ("1", "So với filter/embedded kinh điển (MI-threshold, mRMR, "
                   "ReliefF, LASSO, SFS)", "750 task",
             "RG-SCSO thắng rõ trên 3/5 dataset; LASSO thắng RG-SCSO trên "
             "CẢ HAI dataset gene-expression (Leukemia, ColonCancer)"),
            ("2", "Ablation cô lập vị trí tiêm tín hiệu MI (init/objective/"
                   "transfer)", "750 task",
             "Binarization-interface thắng có ý nghĩa trên 2/5 dataset, hòa "
             "trên 3/5; ngoại lệ Leukemia — MI-init đạt accuracy ngang mà "
             "dùng ít hơn ~4 lần số feature"),
            ("3", "Shuffle-MI — can thiệp nhân quả thật (permute/invert "
                   "relevance field)", "450 task",
             "MI thật vs MI xáo trộn KHÔNG khác biệt có ý nghĩa trên 3/5 "
             "dataset — bằng chứng nhân quả yếu hơn nhiều so với câu chữ "
             "“causally” ban đầu của bài"),
            ("6a", "Biến thiên của chính relevance field qua bootstrap "
                    "resampling", "30 resample × 5 dataset",
             "2 dataset gene-expression có độ ổn định thấp hơn hẳn 3 "
             "dataset còn lại — xác nhận đúng lo ngại overfitting của "
             "reviewer"),
            ("6b", "Nested cross-validation pilot (thay vì chỉ 1 lần chia "
                    "80/20)", "3 dataset × 3 thuật toán × 5 run",
             "Nhất quán hướng với kết quả held-out chính, nhưng công suất "
             "thống kê thấp (n=5), công khai nêu rõ là pilot"),
        ],
    )
    para(
        doc,
        " — hai phát hiện ở Priority 1 và Priority 2 đều là kết quả bất "
        "lợi cho câu chuyện ban đầu của bài. Nguyên tắc xuyên suốt: báo cáo "
        "trung thực, không giấu, không tô hồng — đây chính là điều được "
        "viết thẳng vào Discussion của bài, không né tránh.",
        bold_lead="Lưu ý quan trọng: ",
    )

    # ------------------------------------------------------------ Giai đoạn 3
    h1(doc, "Giai đoạn 3 — Tích hợp trung thực + nâng cấp lý thuyết")
    para(
        doc,
        "Toàn bộ số liệu thật ở Giai đoạn 2 được viết vào bài: thêm bảng "
        "so sánh classic-baseline vào main text, mở rộng bảng ablation với "
        "2 hàng mới (MI-init, MI-objective), viết lại đoạn Discussion về "
        "shuffle-MI theo đúng mức độ bằng chứng thật có (không quá lời).",
    )
    bullet(
        doc,
        " — thu hẹp tiêu đề từ “...for Parsimonious High-Dimensional Feature "
        "Selection” thành “...on High-Dimensional Benchmarks”, khớp với "
        "trần 3571 feature thực tế đã kiểm chứng (không phải 10⁴–10⁵ như "
        "tiêu đề cũ ngụ ý).",
        bold_lead="Sửa tiêu đề (Priority 6d): ",
    )
    bullet(
        doc,
        " — nâng Lemma 1 thành Proposition 2 (“Transition probability and "
        "cumulative leverage”), suy ra từ đúng cơ chế flip-style V-shaped "
        "binarization thật của thuật toán (xác nhận trong pseudocode, "
        "không suy đoán): xác suất lật bit CHÍNH LÀ xác suất chuyển trạng "
        "thái, rồi mở rộng thành cận tích lũy qua N lần lấy mẫu lặp lại "
        "bằng linearity of expectation — trung thực nêu rõ đây KHÔNG phải "
        "chứng minh washout luôn xảy ra, chỉ là một liên kết toán học thật "
        "ở mức kỳ vọng.",
        bold_lead="Nâng cấp lý thuyết (Priority 4): ",
    )

    # ------------------------------------------------------------ Giai đoạn 4
    h1(doc, "Giai đoạn 4 — Xử lý tài liệu phản biện thứ hai (Diem_yeu_RG-SCSO.md)")
    para(
        doc,
        "Trước khi làm bất cứ gì mới, rà lại xem những gì Diem_yeu đòi hỏi "
        "đã có sẵn trong code chưa — tránh làm lại việc đã làm. Kết quả rà "
        "soát: 3/3 mục mức TỬ HUYỆT đầu tiên hoá ra ĐÃ có sẵn hạ tầng giải "
        "quyết từ trước khi đợt sửa này bắt đầu:",
    )
    bullet(doc, "§1.1/§1.2 (đa dạng quần thể kế thừa từ SCSO, rủi ro đóng "
                "băng bit của RMS) — đã có sẵn diversity_analysis() + dữ "
                "liệu diversity_history.csv, đã tích hợp vào Discussion "
                "từ trước.")
    bullet(doc, "§2.1 (trình bày in-sample lấn át leak-free) — kiến trúc "
                "SciRep vốn đã dẫn dắt bằng số liệu held-out (leak-free) "
                "làm bằng chứng chính ngay từ đầu, không cần sửa.")
    para(
        doc,
        "Hai mục còn lại thực sự cần thí nghiệm mới:",
    )
    bullet(
        doc,
        " — thí nghiệm mới, dùng công thức Nogueira (2018, JMLR) — bản tổng "
        "quát hoá chuẩn của Kuncheva consistency index cho trường hợp kích "
        "thước tập con thay đổi qua từng lần chạy. Kết quả trung thực: "
        "RG-SCSO ổn định hơn SCSO trên cả 5 dataset, nhưng trên 2 dataset "
        "gene-expression thì độ ổn định của CHÍNH RG-SCSO cũng gần mức "
        "ngẫu nhiên — không phải một chiến thắng sạch.",
        bold_lead="Stability index (§2.2): ",
    )
    bullet(
        doc,
        " — phát hiện harness run_fs_robustness.py đã hỗ trợ sẵn wrapper "
        "“RF” nhưng chưa từng được chạy thật. Khi chạy thì phát hiện và "
        "sửa MỘT BUG THẬT: ProcessPoolExecutor + RandomForestClassifier bị "
        "deadlock hoàn toàn (0% CPU) trên máy macOS này — một dạng lỗi "
        "kinh điển giữa fork() và threaded BLAS; sửa bằng cách set "
        "OMP_NUM_THREADS=1 (và các biến tương tự) trước khi import numpy. "
        "Do RandomForest tốn tài nguyên hơn KNN rất nhiều, giảm số run từ "
        "30 xuống 10/dataset — công khai ghi rõ đây là pilot rút gọn.",
        bold_lead="RF robustness (§2.5): ",
    )
    para(
        doc,
        "Ngoài ra còn phát hiện và sửa một lỗi có sẵn từ trước, không liên "
        "quan đến 2 tài liệu phản biện: Supplementary tex thiếu hẳn lệnh "
        "\\bibliographystyle/\\bibliography, khiến 3 trích dẫn (kể cả trích "
        "dẫn Nogueira mới thêm) hiển thị lỗi [?] — đã sửa cho cả 3.",
    )

    # ------------------------------------------------------------ Giai đoạn 5
    h1(doc, "Giai đoạn 5 — Chuẩn hoá văn phong theo đúng chuẩn Q1 thật")
    para(
        doc,
        "Thay vì chỉ tự đánh giá, chủ động đọc 4–5 bài Q1/gần-Q1 thật cùng "
        "dòng SCSO (qua bản mirror mở PMC/arXiv, vì link Springer/Nature/"
        "MDPI trực tiếp đều bị chặn) để đối chiếu thật, không đoán.",
    )
    para(
        doc,
        "Bài mình vượt hẳn mặt bằng chung ở nhiều điểm: 0/4 bài mẫu có mục "
        "tự phê bình giới hạn thực chất, 0/4 có bất kỳ lemma/proposition "
        "nào, 0/4 so với filter/embedded baseline kinh điển. Điểm lệch "
        "chuẩn duy nhất tìm được: câu kết Abstract đang kết ở một nhượng "
        "bộ (LASSO thắng), trong khi cả 4 bài mẫu đều kết ở khẳng định "
        "hoặc hướng tương lai.",
    )
    bullet(doc, "Sửa câu kết Abstract — giữ nguyên nội dung LASSO, chỉ đổi "
                "vị trí để kết ở tổng hợp accuracy-preservation/parsimony.")
    bullet(doc, "Tách Discussion từ 1 đoạn ~700 từ/13 câu thành 6 đoạn theo "
                "chủ đề — không đổi số liệu, chỉ phá vỡ cấu trúc đọc như "
                "liệt kê checklist.")

    # ------------------------------------------------------------ Giai đoạn 6
    h1(doc, "Giai đoạn 6 — Tái cấu trúc bảng/hình, giảm số mục")
    para(
        doc,
        "Sau khi người dùng trực tiếp mở PDF xem, phản hồi: quá nhiều mục, "
        "đọc giống AI viết vì lạm dụng chú thích kiểu “(Methods)”/"
        "“(Results)”, và thiếu hình/bảng so sánh mạnh. Ràng buộc cứng "
        "phải tôn trọng: Scientific Reports giới hạn tối đa 8 hình+bảng "
        "gộp lại trong bài chính (đã tra thật từ nature.com, không đoán) — "
        "đã dùng hết 8/8 từ trước. Người dùng chọn phương án: gộp thành "
        "hình/bảng nhiều-panel thay vì vượt giới hạn.",
    )
    bullet(doc, "Gộp Table 1 (accuracy) + Table 2 (số feature) thành 1 "
                "bảng dạng “accuracy (features)” — giải phóng 1 slot.")
    bullet(doc, "Dùng slot đó thêm Table 5 mới: Robustness across "
                "classifiers (KNN/SVM/RF gộp lại) — đúng phần so sánh "
                "mạnh còn thiếu.")
    bullet(doc, "Fig. 3 (mechanism) thành 2-panel: panel (a) enrichment "
                "cũ, panel (b) mới — biểu đồ stability index.")
    bullet(doc, "Methods giảm từ 10 mục xuống 7 mục (gộp RMS+online-"
                "extension, gộp algorithm-budget+computational-complexity, "
                "gộp datasets+baselines+protocol).")
    bullet(doc, "Giảm/viết lại các chú thích kiểu “(Methods)”/“(Results)” "
                "lặp lại nhiều lần — 10 chỗ ở tex, 14 chỗ ở docx, giữ lại "
                "đúng 1 chỗ thật cần thiết mỗi bản.")
    para(
        doc,
        "Trong quá trình này bắt và sửa 2 lỗi thật: bảng Table 1 gộp bị "
        "tràn lề ~198pt ở cỡ chữ \\scriptsize (sửa bằng cách cắt bớt ±std "
        "trong ô, ghi rõ trong caption, không âm thầm bỏ thông tin); và "
        "file overleaf.zip bị lỗi thời (vẫn còn cache hình mechanism.pdf "
        "phiên bản 1-panel cũ) — phát hiện qua sai lệch kích thước file, "
        "rebuild lại.",
    )

    # ------------------------------------------------------------ Giai đoạn 7
    h1(doc, "Giai đoạn 7 — Nén độ dài")
    para(
        doc,
        "Người dùng tiếp tục phản hồi: bài vẫn dài, mục vẫn chưa thật "
        "chuẩn quốc tế. Đọc thêm 1 bài Q1 thật nữa (PMC11591711) xác nhận: "
        "bài SCSO Q1 thật thường 18–22 trang, 0 định lý/chứng minh nào, "
        "không có mục Discussion/Limitations riêng — càng xác nhận phần "
        "lý thuyết + tự phê bình của mình là điểm mạnh, không phải điều "
        "cần cắt.",
    )
    bullet(doc, "Chuyển toàn bộ phần chứng minh (proof) của Lemma 1 và "
                "Proposition 2 sang Supplementary — main text chỉ giữ "
                "phát biểu định lý + ý nghĩa, kèm chú thích trỏ sang "
                "Supplementary.")
    bullet(doc, "Gộp tiếp RMS+UMR thành 1 mục “The RG-SCSO mechanism” — "
                "Methods còn 6 mục.")
    bullet(doc, "Cắt đoạn lặp lại thông tin dataset đã có sẵn ở Table 1/"
                "Results.")
    para(
        doc,
        " 17 trang → 15 trang (main); Supplementary 8→9 trang; "
        "docx-PDF 12→11 trang. KHÔNG đạt đúng mục tiêu ~11 trang mà chính "
        "Scientific Reports khuyến nghị — lý do trung thực: 2 đợt sửa lớn "
        "trong phiên này đã thêm 7 thí nghiệm mới + 1 định lý mới + 2 mục "
        "hình/bảng mới, đều là nội dung thật, không phải nội dung thừa. "
        "Muốn cắt tiếp sẽ phải cắt nội dung thật.",
        bold_lead="Kết quả: ",
    )
    para(
        doc,
        " dừng ở 15 trang, không cắt nội dung thật thêm nữa. Đây là điểm "
        "dừng có chủ đích, không phải bỏ cuộc giữa chừng.",
        bold_lead="Quyết định cuối cùng của người dùng: ",
    )

    # ------------------------------------------------------------ Giai đoạn 8
    h1(doc, "Giai đoạn 8 — Rà và sửa tiêu đề mục cuối cùng")
    para(
        doc,
        "Người dùng hỏi trực tiếp: tiêu đề các mục đã chuẩn chưa, từ ngữ "
        "đã đủ sâu về chuyên ngành chưa. Rà toàn bộ tiêu đề main text + "
        "Supplementary (tex và docx). Kết luận: từ ngữ chuyên ngành đã đủ "
        "chính xác/sâu (đối chiếu trực tiếp với đoạn Problem formulation: "
        "“stratified 5-fold cross-validation”, “cardinality weight”, "
        "“sensitivity range”, “Lipschitz”... đều đúng chuẩn). Chỉ có "
        "đúng 1 tiêu đề lệch chuẩn:",
    )
    bullet(doc, "“Why continuous enhancements wash out” — giọng văn hỏi-"
                "đáp, thông tục, không đúng chuẩn tiêu đề học thuật.")
    para(
        doc,
        " → ",
        bold_lead="Đã sửa thành: ",
    )
    doc.paragraphs[-1].add_run("“Theoretical motivation: transfer-function "
                                "washout”").italic = True
    para(doc, " — sửa ở cả tex và docx, biên dịch lại xác nhận vẫn 15 "
              "trang, không có mục nào khác cần sửa thêm.")

    # ------------------------------------------------------------ Kết quả
    h1(doc, "Kết quả cuối cùng")
    add_table(
        doc,
        ["File", "Nội dung", "Trạng thái"],
        [
            ("RG-SCSO_SciRep.pdf", "Bài chính (tiếng Anh)", "15 trang, đã verify đầy đủ"),
            ("RG-SCSO_SciRep_Supplementary.pdf", "Phụ lục", "9 trang"),
            ("RG-SCSO_SciRep.docx / _word.pdf", "Bản Word", "11 trang (docx-PDF)"),
            ("RG-SCSO_SciRep_overleaf.zip", "Gói nộp Overleaf", "md5-verified khớp nguồn"),
            ("RG-SCSO_SciRep_Revision_Summary.md", "Bảng đối chiếu từng điểm phản biện", "Đầy đủ W1-W13 + Priority 1-6 + Diem_yeu"),
            ("RG-SCSO_SciRep_VI.tex/.pdf", "Bản dịch tiếng Việt (đang thực hiện song song)", "Xem báo cáo riêng khi hoàn tất"),
        ],
    )
    para(
        doc,
        "Mọi lần biên dịch trong toàn bộ quy trình đều dùng tectonic với "
        "sn-jnl.cls/sn-nature.bst tải mới từ nguồn Springer, kiểm tra 0 "
        "lỗi biên dịch, đúng 8/8 mục hình+bảng, Abstract đúng 200/200 từ, "
        "tiêu đề ≤20 từ, 0 dấu gạch ngang dài (em-dash), pytest 10/10 — "
        "lặp lại quy trình này sau MỖI lần sửa, không chỉ ở bước cuối.",
    )
    para(
        doc,
        "Nguyên tắc xuyên suốt toàn bộ quy trình, không có ngoại lệ: "
        "không bịa số liệu, không giấu kết quả bất lợi, mọi thay đổi đều "
        "được biên dịch lại và kiểm tra thật trước khi báo là xong.",
        bold_lead="Nguyên tắc cốt lõi: ",
    )

    doc.save(OUT_DOCX)
    print(f"Đã ghi {OUT_DOCX}")


if __name__ == "__main__":
    build()
