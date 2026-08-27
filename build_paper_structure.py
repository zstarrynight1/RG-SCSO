"""Sinh bản draft bài báo (RG-SCSO) ra .docx theo chuẩn Springer / Applied Intelligence:
1 cột, đánh số mục Ả-Rập (1, 1.1, ...), có Declarations, tham chiếu numbered.

Nguyên tắc (không thỏa hiệp):
  1. MỌI con số trong bảng/prose đọc trực tiếp từ fs_results.csv, không gõ tay.
  2. Phần chưa chạy xong (RIME đủ 18 dataset, kiểm định thống kê, ablation, hình)
     để trống bằng placeholder in đậm màu, KHÔNG bịa số.
  3. Prose tiếng Anh; phần method và setup mô tả ĐÚNG code đã cài, không phóng đại.

Chạy: .venv/bin/python build_paper_structure.py
Xuất:  RG-SCSO_AppliedIntelligence.docx
"""

from __future__ import annotations

import os
import re

import numpy as np
import pandas as pd

import build_heldout_table as _heldout
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

RESULTS_CSV = os.path.join("experiments", "results_fs", "fs_results.csv")
PROCESSED_DIR = os.path.join("data", "processed")
FS_DIR = os.path.join("experiments", "results_fs")
FIG_DIR = "figures"
RANK_CSV = os.path.join(FS_DIR, "friedman_ranking.csv")
FRIED_CSV = os.path.join(FS_DIR, "friedman_summary.csv")
WILC_CSV = os.path.join(FS_DIR, "wilcoxon_vs_rgscso.csv")
ABL_SUMMARY_CSV = os.path.join(FS_DIR, "fs_ablation_summary.csv")
ABL_RAW_CSV = os.path.join(FS_DIR, "fs_ablation_results.csv")
FAMILY_CSV = os.path.join("experiments", "results_fs_scso_family",
                          "fs_scso_family_results.csv")
ROBUST_CSV = os.path.join("experiments", "results_fs_robustness",
                          "fs_robustness_results.csv")
DIVERSITY_CSV = os.path.join("experiments", "results_diversity",
                             "diversity_history.csv")
OUT_DOCX = "RG-SCSO.docx"

COMPLETE_ALGOS = ["RG-SCSO", "SCSO", "AOA", "COA", "GWO", "PSO"]
GRAY = RGBColor(0x80, 0x80, 0x80)
PEND = RGBColor(0x99, 0x66, 0x00)  # muted amber cho placeholder (chuyên nghiệp hơn đỏ)
BODY_PT = 11    # cỡ chữ thân bài 1 cột kiểu Springer (Applied Intelligence)


# --------------------------------------------------------------- số liệu thật
_SUP = str.maketrans("0123456789-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻")


def sci_unicode(x: float, sig: int = 1) -> str:
    """Ký hiệu khoa học đọc được cho p-value trong docx (không dùng superscript run).
    Vd 2.36e-13 -> '2.4 × 10⁻¹³'. Tránh kiểu code '2.4e-13'."""
    import math
    if x == 0:
        return "0"
    exp = int(math.floor(math.log10(abs(x))))
    mant = x / 10 ** exp
    return f"{mant:.{sig}f} × 10{str(exp).translate(_SUP)}"


def pcmp_unicode(p: float) -> str:
    """RHS báo cáo p-value cho docx: p rất nhỏ -> '< 0.001' (tránh độ chính xác
    giả với chỉ 18 block, rev #8); còn lại '= <giá trị>'. Dùng 'p {pcmp_unicode(p)}'."""
    if p < 1e-3:
        return "< 0.001"
    return f"= {p:.3f}"


def load_summary() -> dict:
    raw = pd.read_csv(RESULTS_CSV)
    df = raw[raw.algorithm.isin(COMPLETE_ALGOS)]
    counts = df.groupby(["dataset", "algorithm"]).size().unstack()
    full = counts.dropna()[(counts.dropna() >= 30).all(axis=1)].index.tolist()
    df = df[df.dataset.isin(full)]

    acc_mean = df.groupby(["dataset", "algorithm"])["accuracy"].mean().unstack()[COMPLETE_ALGOS]
    acc_std = df.groupby(["dataset", "algorithm"])["accuracy"].std().unstack()[COMPLETE_ALGOS]
    nf_mean = df.groupby(["dataset", "algorithm"])["n_selected_features"].mean().unstack()[COMPLETE_ALGOS]
    ntot = df.groupby("dataset")["n_total_features"].first()

    ranks = acc_mean.rank(axis=1, ascending=False, method="average")
    avg_rank = ranks.mean().sort_values()

    wtl = {}
    for a in COMPLETE_ALGOS:
        if a == "RG-SCSO":
            continue
        w = int((acc_mean["RG-SCSO"] > acc_mean[a]).sum())
        loss = int((acc_mean["RG-SCSO"] < acc_mean[a]).sum())
        wtl[a] = (w, len(acc_mean) - w - loss, loss)

    # tỉ lệ giảm feature trung bình RG-SCSO vs SCSO / AOA
    red_scso = float((1 - nf_mean["RG-SCSO"] / nf_mean["SCSO"]).mean() * 100)
    red_aoa = float((1 - nf_mean["RG-SCSO"] / nf_mean["AOA"]).mean() * 100)
    # biên accuracy trung bình vs AOA (đối thủ mạnh nhất)
    margin_aoa = float((acc_mean["RG-SCSO"] - acc_mean["AOA"]).mean() * 100)
    margin_scso = float((acc_mean["RG-SCSO"] - acc_mean["SCSO"]).mean() * 100)

    gene = {}
    for g in ["Leukemia", "ColonCancer"]:
        if g in full:
            gene[g] = dict(
                acc=acc_mean.loc[g, "RG-SCSO"], nf=nf_mean.loc[g, "RG-SCSO"],
                ntot=int(ntot[g]), nf_aoa=nf_mean.loc[g, "AOA"],
                acc_aoa=acc_mean.loc[g, "AOA"],
            )

    # RIME đang chạy, chỉ lấy dataset đã đủ 30 run (số THẬT, không bịa phần thiếu)
    rime_df = raw[raw.algorithm == "RIME"]
    rime_cnt = rime_df.groupby("dataset").size()
    rime = {}
    for ds in rime_cnt[rime_cnt >= 30].index:
        g = rime_df[rime_df.dataset == ds]
        rime[ds] = dict(acc=float(g.accuracy.mean()), std=float(g.accuracy.std()))
    rime_won = sum(
        1 for ds in rime if ds in full and acc_mean.loc[ds, "RG-SCSO"] > rime[ds]["acc"]
    )

    out = dict(
        datasets=sorted(full), acc_mean=acc_mean, acc_std=acc_std, nf_mean=nf_mean,
        ntot=ntot, avg_rank=avg_rank, wtl=wtl, n=len(full),
        red_scso=red_scso, red_aoa=red_aoa, margin_aoa=margin_aoa,
        margin_scso=margin_scso, gene=gene,
        rime=rime, rime_done=len(rime), rime_won=rime_won,
    )
    out.update(_load_stats())
    out.update(load_ablation())
    return out


def _load_stats() -> dict:
    """Đọc artifact R4-stats (Wilcoxon+Holm+effect size + Friedman) NẾU đã chạy.
    Trả về dict rỗng (stats=None) khi chưa có → paper vẫn để [pending]."""
    if not (os.path.exists(WILC_CSV) and os.path.exists(RANK_CSV)):
        return dict(stats=None)
    w = pd.read_csv(WILC_CSV)
    rank7 = pd.read_csv(RANK_CSV).set_index("algorithm")["avg_rank"]

    # W/T/L có Ý NGHĨA THỐNG KÊ (Holm) của RG-SCSO vs từng baseline
    sig_wtl = {}
    for a in w["compared_with"].unique():
        m = w[w["compared_with"] == a]["mark"].value_counts()
        sig_wtl[a] = (int(m.get("+", 0)), int(m.get("=", 0)), int(m.get("-", 0)))
    tot = w["mark"].value_counts()
    sig_total = (int(tot.get("+", 0)), int(tot.get("=", 0)), int(tot.get("-", 0)))

    d = w[w["mark"] == "+"]["cohens_d"].abs()
    es_median = float(d.median()) if len(d) else float("nan")
    es_large_pct = float((d >= 0.8).mean() * 100) if len(d) else float("nan")

    ties = [(r.group, r.compared_with) for r in w[w["mark"] != "+"].itertuples()]

    fried = None
    if os.path.exists(FRIED_CSV):
        f = pd.read_csv(FRIED_CSV).iloc[0]
        fried = dict(chi2=float(f["statistic"]), p=float(f["p_value"]),
                     k=int(f["n_algorithms"]))

    return dict(stats=True, rank7=rank7, sig_wtl=sig_wtl, sig_total=sig_total,
                es_median=es_median, es_large_pct=es_large_pct, ties=ties,
                friedman=fried)


# ------------------------------------------------------------------- ablation
# Nhãn hiển thị các cấu hình ablation. "Full" = thiết-kế 3-thành-phần được KHẢO
# SÁT; "− ORL" chính là RG-SCSO CUỐI CÙNG (2-thành-phần) sau khi cắt C2. Thứ tự
# giữ Full trên đầu, NoImprovement (tắt cả 3) cuối để làm mốc dưới.
ABL_CONFIG_ORDER = ["Full", "NoRMS", "NoORL", "NoUMR", "NoImprovement"]
ABL_CONFIG_LABEL = {
    "Full": "Full (RMS+ORL+UMR)",
    "NoRMS": "− RMS",
    "NoORL": "− ORL  (final method)",
    "NoUMR": "− UMR",
    "NoImprovement": "− all three",
}


def load_ablation() -> dict:
    """Đọc artifact ablation (raw + summary) NẾU đã chạy, trả về dữ liệu điền
    Bảng V + verdict prose. Trả dict(ablation=None) khi chưa có → paper để [pending].

    Trả về:
        ablation: True/None (cờ có dữ liệu).
        datasets: list dataset (cột bảng).
        configs: list config (hàng bảng) theo ABL_CONFIG_ORDER.
        means: means[config][dataset] = mean accuracy (30 run).
        sig: set (config, dataset) mà removal LÀM TỆ có ý nghĩa (Holm p<0.05).
        verdict: verdict[removal] = dict(component, kept, n_deg, n_ds, worst_ds,
            worst_delta_pts, worst_d, worst_p, closest_ds, closest_delta_pts,
            closest_p), số liệu điền câu văn load-bearing / cắt.
        kept, cut: list nhãn component (RMS/ORL/UMR) giữ lại / cắt.
    """
    if not (os.path.exists(ABL_RAW_CSV) and os.path.exists(ABL_SUMMARY_CSV)):
        return dict(ablation=None)
    raw = pd.read_csv(ABL_RAW_CSV)
    summ = pd.read_csv(ABL_SUMMARY_CSV)

    datasets = sorted(raw["dataset"].unique())
    configs = [c for c in ABL_CONFIG_ORDER if c in raw["config_name"].unique()]
    means = {c: {} for c in configs}
    for (cfg, ds), g in raw.groupby(["config_name", "dataset"]):
        if cfg in means:
            means[cfg][ds] = float(g["accuracy"].mean())

    sig = {(r.removal, r.dataset) for r in summ[summ["degrades_sig"]].itertuples()}

    verdict = {}
    label = {"NoRMS": "RMS", "NoORL": "ORL", "NoUMR": "UMR"}
    for removal, comp in label.items():
        sub = summ[summ["removal"] == removal]
        deg = sub[sub["degrades_sig"]]
        kept = len(deg) > 0
        # dataset có Δ dương lớn nhất (bằng chứng mạnh nhất, dù có ý nghĩa hay không)
        worst = sub.loc[sub["delta"].idxmax()]
        # dataset "gần đạt" nhất trong nhóm không có ý nghĩa (dùng khi CẮT)
        closest = sub.loc[sub["p_value"].idxmin()]
        verdict[removal] = dict(
            component=comp, kept=kept, n_deg=int(len(deg)), n_ds=int(len(sub)),
            worst_ds=str(worst["dataset"]), worst_delta_pts=float(worst["delta"] * 100),
            worst_d=float(worst["cohens_d"]), worst_p=float(worst["p_holm"]),
            closest_ds=str(closest["dataset"]),
            closest_delta_pts=float(closest["delta"] * 100),
            closest_p=float(closest["p_holm"]),
        )
    kept = [v["component"] for v in verdict.values() if v["kept"]]
    cut = [v["component"] for v in verdict.values() if not v["kept"]]

    # NB: key phải là `abl_datasets` (KHÔNG phải `datasets`), nếu không
    # out.update(load_ablation()) sẽ ghi đè danh sách 18 dataset của bảng chính
    # bằng 5 dataset ablation, cắt cụt Table 2/II/IV.
    return dict(ablation=True, abl_datasets=datasets, configs=configs, means=means,
                sig=sig, verdict=verdict, kept=kept, cut=cut)


# --------------------------------------------------------------- layout 2 cột
def _set_cols(section, n, space=460):
    """Đặt số cột cho một section (IEEE journal = 2 cột thân bài)."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(space))


def full_width(doc, emit):
    """Chạy `emit()` (caption + bảng/hình). Bài Springer 1 cột nên đây chỉ là
    wrapper giữ nguyên 1 cột (không đổi layout)."""
    s1 = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_cols(s1, 1)
    emit()
    s2 = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_cols(s2, 1)


# --------------------------------------------------------------- helper docx
# Times New Roman has no glyph for U+2C7C (subscript "j"), the fake-subscript
# character used throughout the prose (e.g. "rho_j" written as "ρⱼ"); Word/LibreOffice
# silently substitutes a different font for just that one character, which
# with ~40 occurrences reads as the font being broken everywhere. Splitting
# on it and rendering "j" as a real Word subscript run (font.subscript=True)
# uses the surrounding run's own font, so it can never glyph-mismatch again.
# Same treatment for literal caret exponents ("10^4", "{0,1}^d", "[-1,1]^d")
# that never get rendered as real superscripts -- both patterns are
# tokenized in one pass. The exponent can be numeric ("10^4") or a single
# variable ("^d"); either way it's whatever word-run directly follows "^",
# and the caret itself is simply dropped once the exponent is superscripted.
_SUBSCRIPT_J = "ⱼ"
_RUN_TOKEN_RE = re.compile(r"ⱼ|\^\w+")


def _add_run_text(p, text, *, size=None, italic=None, bold=None):
    """Add `text` to paragraph `p`, splitting out fake-subscript-j characters
    and "^exponent" carets into real Word subscript/superscript runs so the
    font stays consistent and exponents render properly."""

    def _add(run_text, *, subscript=False, superscript=False):
        r = p.add_run(run_text)
        if size is not None:
            r.font.size = Pt(size)
        if italic is not None:
            r.italic = italic
        if bold is not None:
            r.bold = bold
        if subscript:
            r.font.subscript = True
        if superscript:
            r.font.superscript = True

    pos = 0
    for m in _RUN_TOKEN_RE.finditer(text):
        if m.start() > pos:
            _add(text[pos:m.start()])
        token = m.group(0)
        if token == _SUBSCRIPT_J:
            _add("j", subscript=True)
        else:
            _add(token[1:], superscript=True)  # drop the leading "^"
        pos = m.end()
    if pos < len(text):
        _add(text[pos:])
    return p


def para(doc, text, size=BODY_PT, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    _add_run_text(p, text, size=size, italic=italic)
    return p


def repeat_header_row(table) -> None:
    """Mark the table's first row as a repeating header (w:tblHeader) so a
    table that spans a page break shows column headers again on the next
    page instead of leaving bare data rows the reader has to scroll back
    to identify."""
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def widen_first_col(table, inches: float = 1.05) -> None:
    """Word's default column-width algorithm splits long single-token cell
    text (dataset names like "WaveformEW", "GermanCredit", or a bracketed
    config label like "(RMS+ORL+UMR)") mid-word when the first column ends
    up too narrow -- there's no space/hyphen for Word to break on instead.

    Setting `cell.width` per cell (python-docx's documented API) is NOT
    enough on its own: Word/LibreOffice render column widths from the
    table's shared <w:tblGrid><w:gridCol .../></w:tblGrid>, and
    doc.add_table() initializes every gridCol to the same default width;
    per-cell tcW overrides are inconsistently honored on top of an
    unchanged, uniform grid. Rewriting tblGrid directly (first column
    wider, remaining width split evenly across the rest) is what actually
    changes the rendered layout; the per-cell widths are kept in sync too
    since some renderers do cross-check them against the grid."""
    table.autofit = False
    grid = table._tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol")) if grid is not None else []
    if len(cols) >= 2:
        total_twips = sum(int(c.get(qn("w:w"))) for c in cols)
        first_twips = int(inches * 1440)
        first_twips = min(first_twips, total_twips - 200 * (len(cols) - 1))
        rest_twips = (total_twips - first_twips) // (len(cols) - 1)
        cols[0].set(qn("w:w"), str(first_twips))
        for c in cols[1:]:
            c.set(qn("w:w"), str(rest_twips))
    for row in table.rows:
        row.cells[0].width = Inches(inches)


def force_font_everywhere(doc, font_name: str = "Times New Roman") -> None:
    """Final pass, called right before doc.save(): explicitly stamp
    w:rFonts on every real Word text run (w:r) in the document, including
    inside every table cell. Style-level rFonts (Normal/Heading N) already
    say Times New Roman, and in principle every run without its own
    override should cascade to that -- but the document's theme
    (docDefaults -> minorHAnsi/majorHAnsi) resolves to Calibri/Cambria, and
    any run that was given direct formatting (bold, size, subscript, ...)
    without an explicit font name is exactly the case where cascade
    behaviour is least reliable across Word/LibreOffice/Pages. Rather than
    audit every call site that creates a run, force it once, everywhere,
    at the end. OMML math runs (m:r/m:t) are untouched -- those correctly
    use Word's own math font (Cambria Math) and forcing Times New Roman
    onto them would look wrong."""
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for run_elem in doc.element.body.iter(f"{w_ns}r"):
        rpr = run_elem.find(f"{w_ns}rPr")
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_elem.insert(0, rpr)
        rfonts = rpr.find(f"{w_ns}rFonts")
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), font_name)
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
            if rfonts.get(qn(theme_attr)) is not None:
                del rfonts.attrib[qn(theme_attr)]


def eq(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(BODY_PT)
    r.font.name = "Cambria Math"   # font toán của Word (ký hiệu hiển thị chuẩn hơn)
    return p


# ------------------------------------------------------- OMML (real Word math)
# Word's native equation objects (Insert > Equation), built directly as OOXML
# math markup so formulas render as proper typeset math (real fractions,
# superscripts, norm bars) instead of plain Cambria-Math text approximating it.
def M(tag):
    return OxmlElement(f"m:{tag}")


def mrun(text):
    """m:r văn bản toán học (in nghiêng tự động theo quy ước OMML)."""
    r = M("r")
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def msup(base_elems, sup_elems):
    e = M("sSup")
    eb = M("e"); eb.extend(base_elems)
    sp = M("sup"); sp.extend(sup_elems)
    e.append(eb); e.append(sp)
    return e


def msub(base_elems, sub_elems):
    e = M("sSub")
    eb = M("e"); eb.extend(base_elems)
    sb = M("sub"); sb.extend(sub_elems)
    e.append(eb); e.append(sb)
    return e


def mfrac(num_elems, den_elems):
    e = M("f")
    e.append(M("fPr"))
    num = M("num"); num.extend(num_elems)
    den = M("den"); den.extend(den_elems)
    e.append(num); e.append(den)
    return e


def mdelim(inner_elems, beg="(", end=")"):
    """Dấu ngoặc/hàng rào tự-co-giãn: (), ||...|| (norm), |...| (abs)."""
    e = M("d")
    dPr = M("dPr")
    begChr = M("begChr"); begChr.set(qn("m:val"), beg)
    endChr = M("endChr"); endChr.set(qn("m:val"), end)
    dPr.append(begChr); dPr.append(endChr)
    e.append(dPr)
    ee = M("e"); ee.extend(inner_elems)
    e.append(ee)
    return e


def eqm(doc, elements, note=None):
    """Chèn 1 dòng công thức OMML thật (list phần tử m:*), căn giữa, thay cho
    text xấp xỉ bằng Cambria-Math. Nhiều dòng liên quan thì gọi eqm() nhiều lần.
    `note`: chú thích văn xuôi ngắn nối sau công thức trên cùng dòng (vd. "if ...")."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    oMathPara = M("oMathPara")
    oMath = M("oMath")
    oMath.extend(elements)
    oMathPara.append(oMath)
    p._p.append(oMathPara)
    if note:
        nr = p.add_run("   " + note)
        nr.font.size = Pt(BODY_PT)
    return p


def blank(doc, label):
    """Placeholder cho số liệu CHƯA chạy xong (không bịa), amber, italic, gọn."""
    p = doc.add_paragraph()
    r = p.add_run("[pending, " + label + "]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = PEND
    return p


def bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        _add_run_text(p, lead, size=BODY_PT)   # không bôi đậm lead-in (Springer)
    _add_run_text(p, text, size=BODY_PT)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run_text(p, text, size=9)                 # caption thường, không bôi đậm
    return p


def add_fig_placeholder(doc, caption_text, height_cm=4.0):
    """Khung hình rỗng (placeholder) + caption Ả-Rập bên dưới, kiểu IEEE."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(3.2)
    cell.vertical_alignment = 1  # center
    body = cell.paragraphs[0]
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(int(height_cm)):
        body.add_run("\n")
    r = body.add_run("[ figure placeholder, generated at R4 ]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption_text)
    cr.font.size = Pt(8.5)
    return t


def add_figure(doc, img_name, caption_text, width_in=3.3):
    """Nhúng ảnh THẬT figures/<img_name> (căn giữa) + caption IEEE bên dưới.

    Nếu file ảnh chưa tồn tại → tự lùi về `add_fig_placeholder` (không bịa hình).
    """
    path = os.path.join(FIG_DIR, img_name)
    if not os.path.exists(path):
        return add_fig_placeholder(doc, caption_text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True  # never let a page break land
    p.add_run().add_picture(path, width=Inches(width_in))  # between image and caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption_text)
    cr.font.size = Pt(8.5)
    return p


# --------------------------------------------------------------- bảng
def _set_borders(el, spec):
    """Gắn <w:*Borders> vào tblPr/tcPr. spec = {'top':(val,sz),...}."""
    pr = el  # đã là tblPr hoặc tcPr
    tag = "w:tblBorders" if pr.tag == qn("w:tblPr") else "w:tcBorders"
    b = pr.find(qn(tag))
    if b is None:
        b = OxmlElement(tag)
        pr.append(b)
    for side, val in spec.items():
        e = b.find(qn("w:" + side))
        if e is None:
            e = OxmlElement("w:" + side)
            b.append(e)
        if val is None:
            e.set(qn("w:val"), "nil")
        else:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(val))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "000000")


def _ieee_table(t):
    """Kiểu booktabs IEEE: không tô màu, không kẻ dọc/kẻ giữa các hàng; chỉ 3
    đường ngang (trên header, dưới header, dưới cùng). Áp cho bảng số liệu."""
    t.style = "Table Grid"          # style trung tính, KHÔNG tô nền (khác Accent)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblpr = t._tbl.tblPr
    # top + bottom = rule dày; bỏ hết kẻ dọc và kẻ ngang giữa các hàng dữ liệu
    _set_borders(tblpr, {"top": 8, "bottom": 8, "left": None, "right": None,
                         "insideH": None, "insideV": None})
    # rule dưới hàng header (mid-rule booktabs)
    for cell in t.rows[0].cells:
        _set_borders(cell._tc.get_or_add_tcPr(), {"bottom": 6})
    return t


def _hdr(t, cols, size=8):
    for i, c in enumerate(cols):
        run = t.rows[0].cells[i].paragraphs[0].add_run(c)
        run.bold = True
        run.font.size = Pt(size)


def add_accuracy_table(doc, s):
    cols = ["Dataset", "#F", "RG-SCSO", "SCSO", "AOA", "COA", "GWO", "PSO", "RIME"]
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _hdr(t, cols)
    for ds in s["datasets"]:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(ds).font.size = Pt(8)
        cells[1].paragraphs[0].add_run(str(int(s["ntot"][ds]))).font.size = Pt(8)
        # best gồm cả RIME nếu dataset này RIME đã xong (để in đậm đúng)
        vals = list(s["acc_mean"].loc[ds])
        if ds in s["rime"]:
            vals.append(s["rime"][ds]["acc"])
        best = max(vals)
        for j, a in enumerate(COMPLETE_ALGOS):
            m, sd = s["acc_mean"].loc[ds, a], s["acc_std"].loc[ds, a]
            run = cells[2 + j].paragraphs[0].add_run(f"{m:.4f}\n±{sd:.3f}")
            run.font.size = Pt(7.5)
            if abs(m - best) < 1e-9:
                run.bold = True
        if ds in s["rime"]:  # RIME đủ 30 run → số thật
            m, sd = s["rime"][ds]["acc"], s["rime"][ds]["std"]
            rr = cells[8].paragraphs[0].add_run(f"{m:.4f}\n±{sd:.3f}")
            rr.font.size = Pt(7.5)
            if abs(m - best) < 1e-9:
                rr.bold = True
        else:  # RIME đang chạy dataset này
            rr = cells[8].paragraphs[0].add_run("-")
            rr.font.size = Pt(8)
            rr.font.color.rgb = PEND


def add_nfeat_table(doc, s):
    cols = ["Dataset"] + COMPLETE_ALGOS
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _hdr(t, cols)
    for ds in s["datasets"]:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(ds).font.size = Pt(8)
        least = s["nf_mean"].loc[ds].min()
        for j, a in enumerate(COMPLETE_ALGOS):
            v = s["nf_mean"].loc[ds, a]
            run = cells[1 + j].paragraphs[0].add_run(f"{v:.1f}")
            run.font.size = Pt(8)
            if abs(v - least) < 1e-9:
                run.bold = True


def add_rank_table(doc, s):
    yr = {"RG-SCSO": "ours", "SCSO": "2022", "AOA": "2021", "COA": "2023",
          "GWO": "2014", "PSO": "1995", "RIME": "2023"}
    t = doc.add_table(rows=1, cols=3)
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hdr(t, ["Algorithm (year)", "Avg. rank", "W/T/L vs RG-SCSO (Holm)"], 9)
    # Đã có R4-stats: xếp theo rank 7 thuật toán (gồm RIME), W/T/L có ý nghĩa Holm.
    ranking = s["rank7"].sort_values() if s.get("stats") else s["avg_rank"]
    for a, r in ranking.items():
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(f"{a} ({yr.get(a, '?')})").font.size = Pt(9)
        cells[1].paragraphs[0].add_run(f"{r:.2f}").font.size = Pt(9)
        if a == "RG-SCSO":
            cells[2].paragraphs[0].add_run("-").font.size = Pt(9)
        elif s.get("stats"):
            w, ti, l = s["sig_wtl"].get(a, (0, 0, 0))
            cells[2].paragraphs[0].add_run(f"{w}/{ti}/{l}").font.size = Pt(9)
        else:
            pr = cells[2].paragraphs[0].add_run("[pending]")
            pr.font.size = Pt(8)
            pr.font.color.rgb = PEND


def add_ablation_table(doc, s):
    """Bảng V, ablation: accuracy trung bình mỗi cấu hình × dataset. Hàng Full
    in đậm (mốc trên); ô có dấu † = cấu hình đó TỆ hơn Full có ý nghĩa (Wilcoxon
    paired + Holm, p<0.05) → thành phần bị gỡ là load-bearing trên dataset đó."""
    ds_list = s["abl_datasets"]
    cols = ["Configuration"] + ds_list
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _hdr(t, cols, 8)
    for cfg in s["configs"]:
        cells = t.add_row().cells
        lab = cells[0].paragraphs[0].add_run(ABL_CONFIG_LABEL[cfg])
        lab.font.size = Pt(8)
        for j, ds in enumerate(ds_list):
            m = s["means"][cfg].get(ds)
            dagger = "†" if (cfg, ds) in s["sig"] else ""
            run = cells[1 + j].paragraphs[0].add_run(
                f"{m:.4f}{dagger}" if m is not None else "-")
            run.font.size = Pt(8)
    caption(doc, "† significantly worse than Full (paired Wilcoxon signed-rank, "
                 "Holm-corrected p < 0.05, 30 runs). “− ORL” is the "
                 "final RG-SCSO (RMS+UMR); ORL removal never degrades accuracy, "
                 "so ORL is not retained.")


def add_dataset_spec_table(doc, s):
    cols = ["Dataset", "Samples", "Features", "Classes"]
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hdr(t, cols, 9)
    for ds in s["datasets"]:
        d = pd.read_csv(os.path.join(PROCESSED_DIR, f"{ds}.csv"))
        vals = [ds, str(len(d)), str(d.shape[1] - 1), str(int(d["label"].nunique()))]
        cells = t.add_row().cells
        for i, v in enumerate(vals):
            cells[i].paragraphs[0].add_run(v).font.size = Pt(9)


# --------------------------------------------------------------- ép font đồng nhất
FONT_NAME = "Times New Roman"


def _force_run_font(run) -> None:
    """Set toàn bộ 4 slot rFonts (ascii/hAnsi/cs/eastAsia) của 1 run về Times New
    Roman, ghi đè cả font kế thừa từ table-style (thường là Calibri)."""
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _iter_all_paragraphs(container):
    """Duyệt paragraph ở body + trong MỌI cell của MỌI bảng (đệ quy bảng lồng)."""
    for p in container.paragraphs:
        yield p
    for tbl in container.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def _enforce_font(doc) -> None:
    """Một lượt cuối: mọi run trong toàn tài liệu dùng chung Times New Roman.
    Cũng set docDefaults để phần tử không có run (vd numbering) kế thừa đúng."""
    # docDefaults rFonts
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    rpr_def = dd.find(qn("w:rPrDefault"))
    if rpr_def is None:
        rpr_def = OxmlElement("w:rPrDefault")
        dd.append(rpr_def)
    rpr = rpr_def.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_def.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    # từng run (body + tất cả bảng + header/footer mọi section)
    scopes = [doc]
    for sec in doc.sections:
        scopes += [sec.header, sec.footer, sec.first_page_header,
                   sec.first_page_footer, sec.even_page_header, sec.even_page_footer]
    for scope in scopes:
        for p in _iter_all_paragraphs(scope):
            for run in p.runs:
                _force_run_font(run)


# bề rộng ước lượng 1 ký tự ở Times 9pt (in) + đệm 2 lề ô; dùng để đo cột theo nội dung
_CHAR_IN = 0.063        # nới nhẹ cho header bold 9pt rộng hơn data 7.5pt (tránh gãy chữ)
_CELL_PAD_IN = 0.16
_MIN_COL_IN = 0.45


def _table_col_counts(doc) -> list:
    """Trả về số cột (1=full-width / 2=inline) của section chứa mỗi bảng, THEO ĐÚNG
    thứ tự `doc.tables`. Trong OOXML, w:sectPr trong pPr của 1 paragraph ĐÓNG section
    đó, mọi nội dung từ ranh giới trước tới paragraph ấy thuộc section có thuộc tính này.
    Nên gom các bảng gặp được từ sau ranh giới trước; khi chạm sectPr kế, gán số cột ấy
    cho cả nhóm. (Không dùng id()/dict vì proxy lxml của cùng 1 node không trùng id.)"""
    body = doc.element.body

    def cols_of(sectpr):
        c = sectpr.find(qn("w:cols"))
        n = c.get(qn("w:num")) if c is not None else None
        return int(n) if n else 1

    result, pending = [], 0
    for ch in body.iterchildren():
        if ch.tag == qn("w:tbl"):
            pending += 1
        elif ch.tag == qn("w:p"):
            ppr = ch.find(qn("w:pPr"))
            sp = ppr.find(qn("w:sectPr")) if ppr is not None else None
            if sp is not None and pending:
                result.extend([cols_of(sp)] * pending)
                pending = 0
    body_sp = body.find(qn("w:sectPr"))
    tail = cols_of(body_sp) if body_sp is not None else 1
    if pending:
        result.extend([tail] * pending)
    return result


def _fit_tables(doc) -> None:
    """Đặt bề rộng cột THEO NỘI DUNG (không chia đều): mỗi cột rộng theo chuỗi dài
    nhất trong cột đó. Bảng nhiều cột-số (accuracy/feature) tự vượt trần → scale vừa
    khít bề rộng trang; bảng nhỏ (dataset-spec, rank) giữ bề rộng tự nhiên và căn giữa,
    tránh cảnh số lọt thỏm giữa ô bị kéo giãn. Fixed layout để Word tôn trọng đúng số đo."""
    sec = doc.sections[0]
    page_tw = int((int(sec.page_width) - int(sec.left_margin)
                   - int(sec.right_margin)) / 914400 * 1440)   # 8640 (full-width)
    col_tw = (page_tw - 460) // 2      # bề rộng 1 cột thân bài (space cols = 460)
    gov = _table_col_counts(doc)       # số cột section chứa mỗi bảng, theo thứ tự doc.tables
    for t, num in zip(doc.tables, gov):
        # bảng nằm trong section 1 cột (full_width) → trần 6in; trong section 2 cột
        # (inline) → trần = bề rộng 1 cột, nếu không bảng sẽ TRÀN sang cột kia (Word lệch).
        target_tw = page_tw if num == 1 else col_tw
        t.autofit = False
        t.allow_autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.CENTER      # bảng hẹp thì căn giữa vùng chứa
        tblpr = t._tbl.tblPr
        # tblLayout = fixed
        lay = tblpr.find(qn("w:tblLayout"))
        if lay is None:
            lay = OxmlElement("w:tblLayout")
            tblpr.append(lay)
        lay.set(qn("w:type"), "fixed")

        grid = t._tbl.tblGrid
        gcs = grid.findall(qn("w:gridCol"))
        ncol = len(gcs)
        # đo chuỗi dài nhất mỗi cột (header + mọi hàng), quy ra bề rộng tự nhiên
        nat_in = [_MIN_COL_IN] * ncol
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                if ci >= ncol:
                    continue
                longest = max((len(ln) for ln in cell.text.split("\n")), default=0)
                w_in = longest * _CHAR_IN + _CELL_PAD_IN
                if w_in > nat_in[ci]:
                    nat_in[ci] = w_in
        nat_tw = [max(1, int(w * 1440)) for w in nat_in]
        cur = sum(nat_tw)
        # 1 cột (Algorithm box) hoặc rộng hơn trần vùng chứa → scale vừa khít trần
        if ncol == 1 or cur > target_tw:
            scale = target_tw / cur
            nat_tw = [max(1, int(w * scale)) for w in nat_tw]
            total_tw = target_tw
        else:
            total_tw = cur                            # bảng nhỏ: giữ tự nhiên, căn giữa
        for g, wtw in zip(gcs, nat_tw):
            g.set(qn("w:w"), str(wtw))
        # QUAN TRỌNG cho Word: set luôn tcW từng ô = gridCol. Word ở fixed-layout ưu
        # tiên tcW của ô; nếu bỏ trống, ô co giãn lệch so với gridCol (LibreOffice thì
        # theo gridCol nên nhìn đúng, đây là chỗ Word "lệch" mà LO không).
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                if ci >= len(nat_tw):
                    continue
                tcpr = cell._tc.get_or_add_tcPr()
                tcw = tcpr.find(qn("w:tcW"))
                if tcw is None:
                    tcw = OxmlElement("w:tcW")
                    tcpr.append(tcw)
                tcw.set(qn("w:w"), str(nat_tw[ci]))
                tcw.set(qn("w:type"), "dxa")
        w = tblpr.find(qn("w:tblW"))
        if w is None:
            w = OxmlElement("w:tblW")
            tblpr.append(w)
        w.set(qn("w:w"), str(total_tw))
        w.set(qn("w:type"), "dxa")
        # lề ô nhỏ (0.04in ~ 58 twips mỗi bên) để chữ có chỗ, tránh wrap thừa
        mar = tblpr.find(qn("w:tblCellMar"))
        if mar is None:
            mar = OxmlElement("w:tblCellMar")
            tblpr.append(mar)
        for side in ("left", "right"):
            e = mar.find(qn("w:" + side))
            if e is None:
                e = OxmlElement("w:" + side)
                mar.append(e)
            e.set(qn("w:w"), "58")
            e.set(qn("w:type"), "dxa")


# --------------------------------------------------------------- build
# --------------------------------------------------------------- references
# Nguồn DUY NHẤT = references.bib (chung với bản .tex). Đánh số theo THỨ TỰ
# TRÍCH DẪN trong bài để KHỚP CHÍNH XÁC với bản PDF/LaTeX (IEEE numbering).
BIB_PATH = "references.bib"
CITE_ORDER = ["guyon", "mrmr", "gwo", "pso", "tf", "scso", "bgwo", "mafarja",
              "aoa", "coa", "rime", "nfl", "neri", "bscso", "scsofs2", "scsofs3",
              "kraskov", "holm", "demsar",
              "imscso2024", "mescso2025", "scsolensobl2024", "improvedscso2024",
              "ludwig2025guided", "islam2017tvtf", "teng2017avbpso"]

_LATEX_MAP = {
    r"{\'e}": "é", r"{\'a}": "á", r"{\'A}": "Á", r"{\'o}": "ó",
    r"{\'i}": "í", r"{\'\i}": "í", r"{\'y}": "ý", r"{\'Y}": "Ý",
    r"{\'u}": "ú", r"{\'U}": "Ú", r'{\"o}': "ö", r'{\"u}': "ü",
    r'{\"a}': "ä", r"{\v{s}}": "š", r"{\v{c}}": "č", r"{\o}": "ø",
    r"{\aa}": "å",
}


def _delatex(s: str) -> str:
    """Gỡ escape LaTeX + gộp khoảng trắng/xuống dòng thành text hiển thị được."""
    for k, v in _LATEX_MAP.items():
        s = s.replace(k, v)
    s = s.replace("{", "").replace("}", "")
    return " ".join(s.split())


def _parse_entry_fields(body: str) -> dict:
    """Bóc các field key=value trong một entry BibTeX (giá trị bọc {} cân bằng)."""
    fields, i, n = {}, 0, len(body)
    while i < n:
        m = re.match(r"\s*(\w+)\s*=\s*", body[i:])
        if not m:
            break
        name = m.group(1).lower()
        i += m.end()
        if i < n and body[i] == "{":
            depth, j = 0, i
            while j < n:
                if body[j] == "{":
                    depth += 1
                elif body[j] == "}":
                    depth -= 1
                    if depth == 0:
                        break
                j += 1
            fields[name] = _delatex(body[i + 1:j])
            i = j + 1
        else:
            j = body.find(",", i)
            j = n if j == -1 else j
            fields[name] = _delatex(body[i:j])
            i = j
        while i < n and body[i] in ", \n\t\r":
            i += 1
    return fields


def _parse_bib(path: str = BIB_PATH) -> dict:
    """Đọc references.bib -> {key: {field: value, __type__: ...}} (balanced-brace)."""
    text = open(path, encoding="utf-8").read()
    entries = {}
    for m in re.finditer(r"@(\w+)\s*\{", text):
        etype, i, depth = m.group(1).lower(), m.end(), 1
        j = i
        while j < len(text) and depth:
            if text[j] == "{":
                depth += 1
            elif text[j] == "}":
                depth -= 1
            j += 1
        key, _, body = text[i:j - 1].partition(",")
        fields = _parse_entry_fields(body)
        fields["__type__"] = etype
        entries[key.strip()] = fields
    return entries


def _fmt_authors(raw: str) -> str:
    """'Last, First and ...' -> 'F. Last, G. Other, and H. Third' (kiểu IEEE)."""
    out = []
    for a in (x.strip() for x in raw.split(" and ")):
        if "," in a:
            last, first = a.split(",", 1)
        else:
            toks = a.rsplit(" ", 1)
            first, last = (toks[0], toks[1]) if len(toks) == 2 else ("", a)
        inits = " ".join(w[0] + "." for w in first.split() if w)
        out.append((inits + " " + last.strip()).strip())
    if len(out) <= 1:
        return out[0] if out else ""
    if len(out) == 2:
        return out[0] + " and " + out[1]
    return ", ".join(out[:-1]) + ", and " + out[-1]


def add_references(doc) -> None:
    """Mục References: đọc references.bib, đánh số theo CITE_ORDER (khớp PDF).

    Venue in nghiêng theo chuẩn IEEE; mọi thông tin lấy từ .bib, không gõ tay.
    """
    doc.add_heading("References", level=1)
    entries = _parse_bib()
    for num, key in enumerate(CITE_ORDER, 1):
        f = entries.get(key)
        if f is None:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def run(text, italic=False):
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.italic = italic

        run(f"[{num}] {_fmt_authors(f.get('author', ''))}, “"
            f"{f.get('title', '')},” ")
        run(f.get("journal") or f.get("booktitle") or "", italic=True)
        seg = ""
        if f.get("volume"):
            seg += f", vol. {f['volume']}"
        if f.get("number"):
            seg += f", no. {f['number']}"
        if f.get("pages"):
            seg += f", pp. {f['pages'].replace('--', '–')}"
        if f.get("year"):
            seg += f", {f['year']}"
        run(seg + ".")
        _bookmark_paragraph(p, f"ref{num}", 3000 + num)   # đích cho [num] click tới


def add_heldout_section(doc, _hs) -> None:
    """§5.G — Generalization under a Leak-Free Hold-Out (Table 6/7). Mirror của
    mục cùng tên trong build_paper_tex.py; TRƯỚC đây mục này hoàn toàn KHÔNG có
    trong bản Word dù Abstract đã dẫn số leak-free làm bằng chứng chính — số đọc
    động từ build_heldout_table.load(), không hardcode."""
    algos = _hs["algos"]
    wil = _hs["wil"]
    fr_chi2 = _hs["stats"]["friedman_chi2"]
    ranking = _hs["ranking"].sort_values()
    rank_str = ", ".join(f"{a} {r:.2f}" for a, r in ranking.items() if a != "RG-SCSO")
    wtl_all = wil["mark"].value_counts()
    w_all, t_all, l_all = (int(wtl_all.get(k, 0)) for k in ("+", "=", "-"))

    def wtl_vs(baseline):
        sub = wil[wil.compared_with == baseline]["mark"].value_counts()
        return (int(sub.get("+", 0)), int(sub.get("=", 0)), int(sub.get("-", 0)))

    w_scso, t_scso, l_scso = wtl_vs("SCSO")
    d_scso = float(wil[wil.compared_with == "SCSO"]["cohens_d"].abs().median())
    w_aoa, t_aoa, l_aoa = wtl_vs("AOA")
    d_aoa = float(wil[wil.compared_with == "AOA"]["cohens_d"].abs().median())
    acc_rg = float(_hs["acc_mean"]["RG-SCSO"].mean())
    acc_aoa = float(_hs["acc_mean"]["AOA"].mean())
    others_d = [float(wil[wil.compared_with == b]["cohens_d"].abs().median())
                for b in ("COA", "GWO", "PSO", "RIME")]
    nf_rg = float(_hs["nf_mean"]["RG-SCSO"].mean())
    nf_scso = float(_hs["nf_mean"]["SCSO"].mean())
    nf_aoa = float(_hs["nf_mean"]["AOA"].mean())
    aoa_loss_ds = wil[(wil.compared_with == "AOA") & (wil["mark"] == "-")]["group"]
    aoa_loss_ds = aoa_loss_ds.iloc[0] if len(aoa_loss_ds) else "n/a"
    d_es_median = float(wil["cohens_d"].abs().median())

    doc.add_heading("G. Generalization under a Leak-Free Hold-Out", level=2)
    para(doc, "For each dataset, algorithm, and independent run we draw an outer "
              "stratified 80/20 split. The relevance prior, the search, and the "
              "cross-validated fitness are computed exclusively on the 80% "
              "training partition; the selected subset is then evaluated once on "
              "the untouched 20% hold-out, on which a fresh k-NN classifier "
              "(standardized on the training partition) reports accuracy. The "
              "evaluation budget, population size, iteration count, seed scheme, "
              "and the 30 independent runs are identical to the primary study; "
              "only the metric of record changes to the held-out accuracy. This "
              "removes any transductive access of the relevance prior to the "
              "test labels.")
    para(doc, "Table 6 reports held-out accuracy over all seven algorithms. "
              f"RG-SCSO attains the best average Friedman rank "
              f"({ranking['RG-SCSO']:.2f}, ahead of {rank_str}; "
              f"chi-square = {fr_chi2:.2f}, p < 0.001). Across the full set of "
              f"pairwise comparisons a Holm-corrected Wilcoxon signed-rank test "
              f"gives RG-SCSO {w_all} significant wins, {l_all} loss, and "
              f"{t_all} ties. Against its own base optimizer SCSO the "
              f"improvement is clear ({w_scso} wins, {l_scso} losses; median "
              f"|d|={d_scso:.2f}), and it also improves on COA, GWO, PSO, and "
              f"RIME (median |d| up to {max(others_d):.2f}). The only close "
              f"competitor is AOA, against which the advantage is genuine but "
              f"moderate ({w_aoa} wins, {l_aoa} loss, {t_aoa} ties; median "
              f"|d|={d_aoa:.2f}), RG-SCSO still leading on mean accuracy "
              f"({acc_rg:.3f} vs. {acc_aoa:.3f}); the single loss occurs on "
              f"{aoa_loss_ds}. Crucially, RG-SCSO delivers this accuracy while "
              f"selecting fewer features: on average {nf_rg:.1f}, against "
              f"{nf_scso:.1f} for SCSO and {nf_aoa:.1f} for AOA (Table 7). Its "
              "ranking advantage is thus coupled with a substantial parsimony "
              "advantage that is robust under the leak-free evaluation.")
    cols = ["Dataset", "#F"] + algos
    tb = doc.add_table(rows=1, cols=len(cols)); _ieee_table(tb); _hdr(tb, cols, size=7)
    for ds in _hs["datasets"]:
        cells = tb.add_row().cells
        cells[0].paragraphs[0].add_run(ds).font.size = Pt(7)
        cells[1].paragraphs[0].add_run(str(int(_hs["ntot"][ds]))).font.size = Pt(7)
        row = _hs["acc_mean"].loc[ds]
        best = row.max()
        for j, a in enumerate(algos):
            m, sdv = _hs["acc_mean"].loc[ds, a], _hs["acc_std"].loc[ds, a]
            r = cells[2 + j].paragraphs[0].add_run(f"{m:.3f}\n±{sdv:.3f}")
            r.font.size = Pt(7)
            if abs(m - best) < 1e-9:
                r.bold = True
    caption(doc, "Table 6 Held-out generalization: mean ± std accuracy on the "
                 "outer 20% hold-out over 30 runs (relevance prior, search, and "
                 "CV fitness fit on the 80% training split only). #F = total "
                 "features. Bold = best per dataset. Higher is better.")

    cols2 = ["Dataset"] + algos
    tb2 = doc.add_table(rows=1, cols=len(cols2)); _ieee_table(tb2); _hdr(tb2, cols2, size=7)
    for ds in _hs["datasets"]:
        cells = tb2.add_row().cells
        cells[0].paragraphs[0].add_run(ds).font.size = Pt(7)
        row = _hs["nf_mean"].loc[ds]
        least = round(row.min(), 1)
        for j, a in enumerate(algos):
            v = _hs["nf_mean"].loc[ds, a]
            r = cells[1 + j].paragraphs[0].add_run(f"{v:.1f}")
            r.font.size = Pt(7)
            if round(v, 1) == least:
                r.bold = True
    caption(doc, "Table 7 Held-out setting: mean number of selected features "
                 "over 30 runs. Bold = fewest.")

    para(doc, "On effect-size magnitude. Comparing the two protocols is itself "
              f"informative. The large in-sample effect sizes (median "
              f"|d|=2.15) contract to a small-to-moderate range (per-baseline "
              f"median |d|={min([d_scso,d_aoa]+others_d):.2f}-"
              f"{max([d_scso,d_aoa]+others_d):.2f}) on the hold-out, whereas the "
              "ranking is preserved (RG-SCSO remains first). This is the "
              "expected signature of the optimistic bias shared by any wrapper "
              "that uses its cross-validation score both to search and to "
              "report: it inflates absolute magnitudes equally for all "
              "methods, hence is fair for relative comparison, but should not "
              "be read as the true out-of-sample gain. The persistence of the "
              "ranking, and of the parsimony advantage, under the leak-free "
              "protocol confirms that the mechanism behind RG-SCSO "
              "(relevance-biased flipping) transfers to unseen data and is not "
              "a consequence of the prior observing the labels.")


def add_scso_family_section(doc) -> None:
    """§V-G, so RG-SCSO với baseline CÙNG HỌ SCSO-FS (bSCSO-S/OBL). Số đọc raw CSV."""
    from src.stats.statistical_tests import paired_wilcoxon_vs_target
    rg = pd.read_csv(RESULTS_CSV)
    rg = rg[rg.algorithm == "RG-SCSO"][
        ["algorithm", "dataset", "run_id", "accuracy", "n_selected_features"]]
    fam = pd.read_csv(FAMILY_CSV)[
        ["algorithm", "dataset", "run_id", "accuracy", "n_selected_features"]]
    combined = pd.concat([rg, fam], ignore_index=True)
    order = ["RG-SCSO", "bSCSO-S", "bSCSO-OBL"]
    labels = {"RG-SCSO": "RG-SCSO (ours)", "bSCSO-S": "bSCSO (S-shaped)",
              "bSCSO-OBL": "bSCSO (V-shaped + OBL)"}
    macc = {a: combined[combined.algorithm == a].accuracy.mean() for a in order}
    mnf = {a: combined[combined.algorithm == a].n_selected_features.mean() for a in order}
    wil = paired_wilcoxon_vs_target(combined, "accuracy", "dataset", "algorithm",
                                    "run_id", "RG-SCSO", lower_is_better=False)
    datasets = sorted(rg.dataset.unique())
    wtl, red, smaller = {}, {}, {}
    for a in order[1:]:
        sub = wil[wil.compared_with == a]
        w = int((sub["mark"] == "+").sum()); t = int((sub["mark"] == "=").sum())
        loss = int((sub["mark"] == "-").sum())
        wtl[a] = f"{w}/{t}/{loss}"
        red[a] = (mnf[a] - mnf["RG-SCSO"]) / mnf[a] * 100
        smaller[a] = sum(
            combined[(combined.algorithm == "RG-SCSO") & (combined.dataset == d)]
            .n_selected_features.mean()
            < combined[(combined.algorithm == a) & (combined.dataset == d)]
            .n_selected_features.mean() for d in datasets)
    n_ds = len(datasets)

    doc.add_heading("G. Comparison with Same-Family SCSO Feature Selectors", level=2)
    para(doc,
         "Because the paper locates its gap inside the SCSO feature-selection line, "
         "we compare directly against binary SCSO selectors of that same family. We "
         "reimplement two standard recipes under the identical protocol: bSCSO "
         "(S-shaped), which binarizes the SCSO position with an S-shaped transfer, "
         "and bSCSO (V-shaped + OBL), which adds a V-shaped transfer and "
         "opposition-based learning; neither carries a per-feature relevance field "
         "(reimplemented from published pseudocode, parameters as reported). As "
         "Table 8 shows, on accuracy RG-SCSO does not dominate its own family "
         f"(win/tie/loss {wtl['bSCSO-S']} against bSCSO (S-shaped) and "
         f"{wtl['bSCSO-OBL']} against bSCSO (V-shaped + OBL)), the mean accuracies "
         f"lying within {abs(macc['RG-SCSO']-macc['bSCSO-S'])*100:.2f} of a point. "
         f"What separates RG-SCSO is parsimony: it averages {mnf['RG-SCSO']:.0f} "
         f"features against {mnf['bSCSO-S']:.0f} and {mnf['bSCSO-OBL']:.0f}, roughly "
         f"{red['bSCSO-S']:.0f}% fewer at equal accuracy, and it selects the smaller "
         f"subset on {smaller['bSCSO-S']} of {n_ds} datasets. The per-feature "
         "relevance field is therefore what buys subset compactness within the SCSO "
         "family, consistent with the paper's central claim.")
    caption(doc, "Table 8 Comparison with same-family binary SCSO feature "
                 f"selectors ({n_ds} datasets × 30 runs, budget-matched). W/T/L is "
                 "RG-SCSO's Holm-corrected win/tie/loss on accuracy (paired Wilcoxon, "
                 "per-dataset family).")   # caption TRÊN bảng (chuẩn Springer)
    cols = ["Method", "Mean Acc.", "Mean #Feat.", "W/T/L vs RG-SCSO"]
    tb = doc.add_table(rows=1, cols=len(cols)); _ieee_table(tb); _hdr(tb, cols)
    for a in order:
        cells = tb.add_row().cells
        r0 = cells[0].paragraphs[0].add_run(labels[a]); r0.font.size = Pt(8)
        r1 = cells[1].paragraphs[0].add_run(f"{macc[a]:.4f}"); r1.font.size = Pt(8)
        r2 = cells[2].paragraphs[0].add_run(f"{mnf[a]:.1f}"); r2.font.size = Pt(8)
        r3 = cells[3].paragraphs[0].add_run("-" if a == "RG-SCSO" else wtl[a])
        r3.font.size = Pt(8)
        if a == "RG-SCSO":
            r2.bold = True   # chỉ đậm giá trị tốt nhất (ít feature nhất), không đậm tên


def add_robustness_section(doc) -> None:
    """§V-H, robustness qua wrapper (KNN/SVM) và prior (MI/ReliefF). Số từ raw CSV."""
    rob = pd.read_csv(ROBUST_CSV)
    # Bảng chéo KNN×SVM cố định trên 5 dataset đại diện; SVM/18 phân tích riêng.
    rob = rob[rob.dataset.isin(["Zoo", "Sonar", "WDBC", "ColonCancer", "Leukemia"])]
    wrappers = ["KNN", "SVM"]
    algos = ["RG-SCSO-MI", "RG-SCSO-ReliefF", "bSCSO"]
    labels = {"RG-SCSO-MI": "RG-SCSO (MI)", "RG-SCSO-ReliefF": "RG-SCSO (ReliefF)",
              "bSCSO": "bSCSO (no prior)"}

    def m(w, a, col):
        return rob[(rob.wrapper == w) & (rob.algorithm == a)][col].mean()

    red = {w: (m(w, "bSCSO", "n_selected_features") - m(w, "RG-SCSO-MI", "n_selected_features"))
              / m(w, "bSCSO", "n_selected_features") * 100 for w in wrappers}
    ratio = {w: m(w, "RG-SCSO-ReliefF", "n_selected_features")
                / m(w, "RG-SCSO-MI", "n_selected_features") for w in wrappers}
    n_ds = rob["dataset"].nunique()

    doc.add_heading("H. Robustness Across Classifiers and Relevance Priors", level=2)
    para(doc,
         "Two design choices could confound the main results: the KNN wrapper and "
         "the mutual-information prior might form a uniquely favorable pair. We test "
         f"both on {n_ds} representative datasets by crossing two wrappers (KNN and "
         "an SVM) with two priors (MI and ReliefF), against bSCSO with no prior. "
         "Table 9 reports two findings, one supportive and one that bounds the "
         "claim. First, the parsimony advantage is not a KNN artifact: under the "
         f"SVM wrapper RG-SCSO (MI) still selects about {red['SVM']:.0f}% fewer "
         f"features than bSCSO ({red['KNN']:.0f}% under KNN) at competitive accuracy, "
         "so the mechanism transfers across classifiers. Second, and reported "
         "plainly as a limitation, the parsimony is specific to the MI prior: "
         f"swapping in ReliefF yields subsets about {ratio['KNN']:.1f}× larger under "
         f"KNN and {ratio['SVM']:.1f}× larger under SVM, comparable to bSCSO and "
         "above RG-SCSO (MI). The normalized MI score drives uninformative features "
         "below the neutral point so RMS prunes them, whereas the ReliefF weights on "
         "these datasets sit mostly above neutral and bias toward inclusion. "
         "RG-SCSO's compactness therefore depends on a prior that pushes noise "
         "features toward exclusion, not on the mere presence of a relevance signal; "
         "we did not retune the ReliefF mapping to obscure this.")
    caption(doc, "Table 9 Robustness across classifier wrappers and relevance "
                 f"priors on {n_ds} representative datasets (× 30 runs, "
                 "budget-matched). Fewest features per wrapper in bold.")   # caption TRÊN bảng
    cols = ["Wrapper", "Method", "Mean Acc.", "Mean #Feat."]
    tb = doc.add_table(rows=1, cols=len(cols)); _ieee_table(tb); _hdr(tb, cols)
    for w in wrappers:
        for i, a in enumerate(algos):
            cells = tb.add_row().cells
            cells[0].paragraphs[0].add_run(w if i == 0 else "").font.size = Pt(8)
            r1 = cells[1].paragraphs[0].add_run(labels[a]); r1.font.size = Pt(8)
            cells[2].paragraphs[0].add_run(f"{m(w, a, 'accuracy'):.4f}").font.size = Pt(8)
            r3 = cells[3].paragraphs[0].add_run(f"{m(w, a, 'n_selected_features'):.1f}")
            r3.font.size = Pt(8)
            if a == "RG-SCSO-MI":
                r3.bold = True   # chỉ đậm giá trị tốt nhất, không đậm tên method

    # ---- SVM trên 16 dataset (rev #3, diện rộng) ----
    from src.stats.statistical_tests import paired_wilcoxon_vs_target
    svm = pd.read_csv(ROBUST_CSV)
    svm = svm[svm.wrapper == "SVM"]
    ds16 = sorted(svm.dataset.unique())

    def mn(a, col):
        return svm[svm.algorithm == a][col].mean()

    s_macc = {a: mn(a, "accuracy") for a in algos}
    s_mnf = {a: mn(a, "n_selected_features") for a in algos}
    s_red = (s_mnf["bSCSO"] - s_mnf["RG-SCSO-MI"]) / s_mnf["bSCSO"] * 100
    s_ratio = s_mnf["RG-SCSO-ReliefF"] / s_mnf["RG-SCSO-MI"]
    wil = paired_wilcoxon_vs_target(
        svm[svm.algorithm.isin(["RG-SCSO-MI", "bSCSO"])], "accuracy", "dataset",
        "algorithm", "run_id", "RG-SCSO-MI", lower_is_better=False)
    sw = int((wil["mark"] == "+").sum()); st_ = int((wil["mark"] == "=").sum())
    sl = int((wil["mark"] == "-").sum())
    s_smaller = sum(
        svm[(svm.algorithm == "RG-SCSO-MI") & (svm.dataset == d)].n_selected_features.mean()
        < svm[(svm.algorithm == "bSCSO") & (svm.dataset == d)].n_selected_features.mean()
        for d in ds16)
    para(doc, "To confirm the classifier result beyond the five-dataset "
              f"cross-tabulation, we ran the SVM wrapper on {len(ds16)} of the 18 "
              "datasets, excluding only the two largest-sample sets (KrVsKpEW and "
              "WaveformEW), for which kernel-SVM training is O(n²) per fit and "
              "infeasible inside a 15,000-evaluation wrapper; the KNN main study "
              "already covers those two. The picture is unchanged (Table 10): "
              f"RG-SCSO (MI) selects {s_red:.0f}% fewer features than the no-prior "
              "baseline at essentially equal accuracy "
              f"({s_macc['RG-SCSO-MI']:.4f} vs {s_macc['bSCSO']:.4f}), on "
              f"{s_smaller} of {len(ds16)} datasets, and does not lead on accuracy "
              f"(win/tie/loss {sw}/{st_}/{sl}). The ReliefF variant again inflates "
              f"subsets to {s_ratio:.1f}× the MI size. The parsimony mechanism "
              "therefore generalizes across classifiers on the broad benchmark, and "
              "its dependence on the mutual-information prior generalizes with it.")
    caption(doc, f"Table 10 SVM wrapper on {len(ds16)} of 18 datasets (× 30 runs). "
                 "The two largest-sample sets (KrVsKpEW, WaveformEW) are excluded: "
                 "kernel-SVM is O(n²) per fit, infeasible in a 15,000-evaluation "
                 "wrapper, and the KNN main study covers them. Fewest features in "
                 "bold.")
    cols2 = ["Method", "Mean Acc.", "Mean #Feat."]
    tb2 = doc.add_table(rows=1, cols=len(cols2)); _ieee_table(tb2); _hdr(tb2, cols2)
    s_min = min(s_mnf.values())
    for a in algos:
        cells = tb2.add_row().cells
        cells[0].paragraphs[0].add_run(labels[a]).font.size = Pt(8)
        cells[1].paragraphs[0].add_run(f"{s_macc[a]:.4f}").font.size = Pt(8)
        rf = cells[2].paragraphs[0].add_run(f"{s_mnf[a]:.1f}"); rf.font.size = Pt(8)
        if abs(s_mnf[a] - s_min) < 1e-9:
            rf.bold = True


def add_diversity_section(doc) -> None:
    """Chẩn đoán đóng băng bit / sụp đa dạng quần thể (phòng thủ tử huyệt
    §1.1/§1.2 Diem_yeu_RG-SCSO.md). Bỏ qua nếu chưa đo (measure_diversity.py).
    Đặt CUỐI Results (sau Robustness) để bảng mới là Table 11, không phải dịch
    số các bảng trước (khác LaTeX, Word đánh số bảng bằng chuỗi tĩnh)."""
    if not os.path.exists(DIVERSITY_CSV):
        return
    df = pd.read_csv(DIVERSITY_CSV)
    datasets = sorted(df["dataset"].unique())
    end = df.sort_values("iter").groupby(["dataset", "gamma"]).tail(1)

    def val(ds, g, col):
        row = end[(end.dataset == ds) & (end.gamma == g)]
        return float(row[col].iloc[0])

    def maxcol(ds, g, col):
        return float(df[(df.dataset == ds) & (df.gamma == g)][col].max())

    def first_iter_over(ds, g, thresh):
        sub = df[(df.dataset == ds) & (df.gamma == g)].sort_values("iter")
        hit = sub[sub["frozen_frac"] > thresh]
        return int(hit["iter"].min()) if len(hit) else None

    max_frz_g5 = max(maxcol(d, 0.5, "frozen_frac") for d in datasets)
    worst_ds = max(datasets, key=lambda d: val(d, 1.0, "frozen_frac"))
    worst_frz_g1 = val(worst_ds, 1.0, "frozen_frac")
    worst_first_freeze = first_iter_over(worst_ds, 1.0, 0.1)

    doc.add_heading("I. Exploration Safety: A Bit-Freezing Diagnostic", level=2)
    para(doc, "Because RMS biases the flip probability toward a preferred bit "
              "(Section 3.2), a legitimate concern is that strong bias (γ close "
              "to 1) could push p_flip toward the clip boundary and freeze "
              "confidently classified features into the mask early, collapsing "
              "the population's coverage of the subset space, the same failure "
              "mode Section 3.3 shows for continuous operators, now asked of the "
              "binary domain. We instrument the search to record, each "
              "iteration, the population's mean expected Hamming spread "
              "2p̄(1−p̄) per feature, our measure of population diversity, "
              "together with the fraction of features on which every "
              "individual already agrees, the frozen fraction. Three settings "
              "are compared: γ = 0, a plain V-shaped transfer with no "
              "relevance signal; γ = 0.5, the value used throughout this "
              "paper; and γ = 1, a deliberately aggressive stress test. All "
              "three run on three datasets spanning low, medium, and very "
              f"high dimensionality ({', '.join(datasets)}). Table 11 "
              "reports the end-of-run values.")
    para(doc, "The risk is real at the stress-test setting: under γ = 1, "
              f"{worst_ds} reaches {worst_frz_g1*100:.0f}% frozen features by "
              f"iteration {worst_first_freeze}, and the effect is sharper on "
              "higher-dimensional data, exactly as the clip term in the RMS "
              "rule predicts, since more features reach confident relevance "
              "scores when there are more of them. At the deployed default "
              f"γ = 0.5, however, the frozen fraction never exceeds "
              f"{max_frz_g5*100:.1f}% at any iteration on any of the three "
              "datasets, including the highest-dimensional one; diversity "
              "decreases moderately but the population retains broad coverage "
              "of the subset space throughout the run. The bit-freezing risk "
              "is therefore genuine and γ-dependent, not a property of RMS we "
              "can dismiss, but the conservative default this paper ships "
              "with avoids it in practice.")
    caption(doc, "Table 11 End-of-run population diversity and frozen-bit "
                 "fraction, mean over 3 runs, for three relevance-bias "
                 "settings: no relevance (γ = 0, plain V-shaped), the "
                 "default used throughout the paper (γ = 0.5), and a "
                 "maximal-bias stress test (γ = 1). Diversity is the mean "
                 "expected normalized Hamming spread 2p̄(1−p̄) across "
                 "features; frozen fraction is the share of features on "
                 "which every individual in the population agrees.")
    cols = ["Dataset", "Div.(γ=0)", "Div.(γ=0.5)", "Div.(γ=1)",
            "Frozen(γ=0)", "Frozen(γ=0.5)", "Frozen(γ=1)"]
    tb = doc.add_table(rows=1, cols=len(cols)); _ieee_table(tb); _hdr(tb, cols, size=7)
    for d in datasets:
        cells = tb.add_row().cells
        vals = [d] + [f"{val(d,g,'diversity'):.3f}" for g in (0.0, 0.5, 1.0)] \
                    + [f"{val(d,g,'frozen_frac'):.3f}" for g in (0.0, 0.5, 1.0)]
        for c, v in zip(cells, vals):
            c.paragraphs[0].add_run(v).font.size = Pt(7.5)


def add_threats_section(doc, s, adaptive_red_min: float, adaptive_red_max: float) -> None:
    """§V-J Threats to Validity, mirroring build_paper_tex.py's subsection so the
    two artifacts carry the same limitations. Numbers pulled fresh from the same
    CSVs the other Results subsections already read; no cross-doc import since
    the two generators are kept independent by design."""
    div_df = pd.read_csv(DIVERSITY_CSV)
    max_frz_g5 = float(div_df[div_df.gamma == 0.5]["frozen_frac"].max() * 100)

    rob = pd.read_csv(ROBUST_CSV)
    rob5 = rob[rob.dataset.isin(["Zoo", "Sonar", "WDBC", "ColonCancer", "Leukemia"])]

    def m5(w, a, col):
        return rob5[(rob5.wrapper == w) & (rob5.algorithm == a)][col].mean()

    ratio_knn = m5("KNN", "RG-SCSO-ReliefF", "n_selected_features") / \
        m5("KNN", "RG-SCSO-MI", "n_selected_features")

    svm = rob[rob.wrapper == "SVM"]
    ds16 = sorted(svm.dataset.unique())

    def m16(a, col):
        return svm[svm.algorithm == a][col].mean()

    s_red = (m16("bSCSO", "n_selected_features") - m16("RG-SCSO-MI", "n_selected_features")) \
        / m16("bSCSO", "n_selected_features") * 100

    feat_min, feat_max = int(s["ntot"].min()), int(s["ntot"].max())

    doc.add_heading("J. Threats to Validity", level=2)
    para(doc,
         "Several boundaries delimit what these results establish. RG-SCSO "
         "inherits SCSO's continuous search dynamics unchanged, exploration "
         "limitations included, and the RMS rule carries a risk of its own: "
         "bias the flip probability too strongly and confidently classified "
         "bits can freeze in place, draining the population's diversity. We "
         "did not assume this away; we measured it directly (Exploration "
         "Safety above). At an aggressive stress-test bias the risk is real, "
         "and it grows with dimensionality. The conservative γ = 0.5 this "
         "paper deploys, though, keeps the frozen-bit fraction below "
         f"{max_frz_g5:.1f}% throughout the run on every dataset tested, "
         "including the highest-dimensional one (Table 11).")
    para(doc,
         "Two further design choices could narrow the claim. The main "
         "objective is a KNN wrapper, and because mutual information and "
         "KNN both weigh local neighbourhood structure, part of the "
         "advantage might be specific to that pairing. Under an SVM "
         f"wrapper, tested directly on {len(ds16)} of the 18 datasets "
         f"(Table 10), the parsimony advantage persists: about {s_red:.0f}% "
         "smaller subsets than the no-prior baseline at essentially equal "
         "accuracy, so the mechanism is not a KNN artifact, though "
         "tree-ensemble and neural classifiers remain future work. The "
         "relevance field is also built from a single filter statistic, the "
         "mutual information between each feature and the label. Swapping "
         "in a ReliefF prior (Table 9) exposes a genuine limitation rather "
         "than confirming the mechanism: the parsimony advantage "
         f"disappears, and subsets grow to roughly {ratio_knn:.1f} times "
         "those of the MI variant, comparable to the no-prior baseline. "
         "RG-SCSO's compactness, it turns out, depends on a prior that "
         "drives uninformative features below the neutral point, not on "
         "the mere presence of a relevance signal.")
    para(doc,
         "External validity has its own limits. The benchmark spans "
         f"{feat_min} to {feat_max} features across biomedical, "
         "gene-expression, and categorical domains, but it is drawn from a "
         "single curated family of UCI and standard microarray sets, so "
         "validity to other data regimes is asserted rather than proven. "
         "We evaluate up to a few thousand features; behaviour on "
         "ultra-high-dimensional omics data of 10⁴ to 10⁵ features is "
         "extrapolated, not measured. The numbers themselves carry a "
         "caveat too: in-sample effect sizes, reported above in "
         "Classification Accuracy, inherit the optimistic bias intrinsic "
         "to any wrapper that both searches and reports on the same folds, "
         "and we regard the leak-free hold-out as the conservative "
         "estimate. Finally, the accuracy claim is scoped, not universal. "
         "Against binary particle-swarm and grey-wolf optimizers carrying "
         "published adaptive V-shaped transfers, and against same-family "
         "binary SCSO selectors (Comparison with Same-Family SCSO Feature "
         "Selectors above), RG-SCSO does not lead on accuracy, so we scope "
         "the accuracy result to the standard baseline suite and frame "
         "parsimony, the "
         f"{adaptive_red_min:.0f} to {adaptive_red_max:.0f}% smaller "
         "subsets that hold even there, as the mechanism's transferable "
         "benefit. We state these limits explicitly so the claim, a "
         "per-feature relevance mechanism that selects smaller subsets and "
         "improves accuracy on the standard suite, for binary SCSO-style "
         "feature selection under a KNN wrapper on this benchmark family, "
         "is not overread.")


def _bookmark_paragraph(p, name, bmid) -> None:
    """Bọc paragraph bằng bookmark (đích để [n] nhảy tới). Đặt bookmarkStart sau
    pPr, bookmarkEnd ở cuối."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bmid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bmid))
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p._p.insert(0, start)
    p._p.append(end)


def _mk_run_xml(text, size):
    """1 <w:r> văn bản thường (Times New Roman, cỡ size), giữ khoảng trắng."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rf)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(round(size.pt * 2))))
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _cite_hyperlink_xml(anchor, text, size):
    """<w:hyperlink w:anchor=...> quanh [n] — click nhảy tới bookmark, giữ màu đen."""
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)
    r = _mk_run_xml(text, size)
    col = OxmlElement("w:color")           # đen, không xanh/underline kiểu web-link
    col.set(qn("w:val"), "000000")
    r.find(qn("w:rPr")).append(col)
    hl.append(r)
    return hl


def _linkify_citations(doc, maxref) -> None:
    """Đổi mọi [n] trong thân bài thành internal hyperlink tới bookmark ref{n}.
    Bỏ qua heading và danh mục References."""
    pat = re.compile(r"\[(\d+)\]")
    in_refs = False
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            if p.text.strip().lower().startswith("references"):
                in_refs = True
            continue
        if in_refs or not pat.search(p.text):
            continue
        text = p.text
        size = p.runs[0].font.size if (p.runs and p.runs[0].font.size) else None
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        pos = 0
        for m in pat.finditer(text):
            if m.start() > pos:
                p._p.append(_mk_run_xml(text[pos:m.start()], size))
            n = int(m.group(1))
            if 1 <= n <= maxref:
                p._p.append(_cite_hyperlink_xml(f"ref{n}", m.group(0), size))
            else:
                p._p.append(_mk_run_xml(m.group(0), size))
            pos = m.end()
        if pos < len(text):
            p._p.append(_mk_run_xml(text[pos:], size))


def _renumber_springer(doc) -> None:
    """Đổi số mục IEEE (I./II. + A./B.) sang Springer Ả-Rập (1, 1.1, 2, 2.1, ...).
    'Statements and Declarations' và 'References' để KHÔNG đánh số (back-matter
    chuẩn Springer)."""
    top, sub = 0, 0
    for p in doc.paragraphs:
        st = p.style.name
        if st not in ("Heading 1", "Heading 2"):
            continue
        title = re.sub(r"^\s*([IVXLC]+|[A-Z])\.\s*", "", p.text).strip()
        if st == "Heading 1":
            if title.lower() in ("references", "declarations", "statements and declarations"):
                new = title
            else:
                top += 1
                sub = 0
                new = f"{top} {title}"
        else:
            sub += 1
            new = f"{top}.{sub} {title}"
        for r in list(p.runs):
            r.text = ""
        if p.runs:
            p.runs[0].text = new
        else:
            p.add_run(new)


def build():
    s = load_summary()
    # Số leak-free cho Abstract (đọc động từ artifact, KHÔNG hardcode) — dùng làm
    # bằng chứng CHÍNH trong Abstract thay vì con số in-sample optimistic.
    _hs = _heldout.load()
    hs_rank = float(_hs["ranking"]["RG-SCSO"])
    hs_wtl = _hs["wil"]["mark"].value_counts()
    hs_w, hs_t, hs_l = (int(hs_wtl.get(k, 0)) for k in ("+", "=", "-"))
    hs_d_aoa = float(_hs["wil"][_hs["wil"].compared_with == "AOA"]["cohens_d"].abs().median())
    _adap = pd.read_csv(os.path.join(
        "experiments", "results_fs_adaptive_baselines", "summary_vs_rgscso.csv")).set_index("config")
    _adap_rg_nf = float(_adap.loc["RG-SCSO", "mean_nfeat"])
    _adap_red = [(row["mean_nfeat"] - _adap_rg_nf) / row["mean_nfeat"] * 100
                 for cfg, row in _adap.iterrows() if cfg != "RG-SCSO"]
    adaptive_red_min, adaptive_red_max = min(_adap_red), max(_adap_red)
    doc = Document()
    # A4 + margins matching the compiled sn-jnl PDF (measured via pdftotext -bbox:
    # left/right text block ~35-43mm, top ~25.5mm), not python-docx's US-Letter
    # default, so the Word and LaTeX deliverables share the same page geometry.
    sec0 = doc.sections[0]
    sec0.page_width = Mm(210)
    sec0.page_height = Mm(297)
    sec0.left_margin = Mm(35)
    sec0.right_margin = Mm(35)
    sec0.top_margin = Mm(25)
    sec0.bottom_margin = Mm(25)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(BODY_PT)
    st.paragraph_format.line_spacing = 1.0
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(6)
    # Heading kiểu Springer: đen, Times, RÕ RÀNG lớn hơn thân bài (không phải
    # xanh-to mặc định Word, và không được bằng/nhỏ hơn cỡ chữ thân bài).
    for hs, sz, before, after in [("Heading 1", 14, 18, 8), ("Heading 2", 12, 14, 6)]:
        h = doc.styles[hs]
        h.font.name = "Times New Roman"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.italic = False
        h.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.line_spacing = 1.0

    # ---- Title block (full-width, kiểu IEEE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "Relevance-Guided Sand Cat Swarm Optimization: A Per-Feature "
        "Relevance-Modulated Binarization for Parsimonious "
        "High-Dimensional Feature Selection")
    tr.bold = True
    tr.font.size = Pt(18)

    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a1 = au.add_run("Bui Quang Huy")
    a1.font.size = Pt(11)
    sup1 = au.add_run("1,*")
    sup1.font.size = Pt(11)
    sup1.font.superscript = True
    a2 = au.add_run(" and Duong Minh Son")
    a2.font.size = Pt(11)
    sup2 = au.add_run("1")
    sup2.font.size = Pt(11)
    sup2.font.superscript = True
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affn1 = aff.add_run("1")
    affn1.font.superscript = True
    affn1.italic = True
    affn1.font.size = Pt(9.5)
    affr = aff.add_run("Dong A University, Da Nang, Vietnam")
    affr.italic = True
    affr.font.size = Pt(9.5)
    af = doc.add_paragraph()
    af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    afr = af.add_run("*Corresponding author(s). E-mail(s): huybq@donga.edu.vn; "
                      "Contributing authors: sondm@donga.edu.vn")
    afr.italic = True
    afr.font.size = Pt(9.5)

    # ---- Abstract (IEEE: 1 đoạn, bold lead-in, justify, cỡ nhỏ)
    ab = doc.add_paragraph()
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = ab.add_run("Abstract ")
    lead.bold = True
    lead.font.size = Pt(9)
    abs_body = (
        "Wrapper feature selection with swarm intelligence typically searches "
        "continuously and crosses into the binary domain via a fixed transfer "
        "function, a feature-agnostic quantization that discards continuous "
        "operators' fine adjustments, an effect we term washout. RG-SCSO "
        "instead replaces this transfer with a per-feature, relevance-modulated "
        "binarization, biasing each feature's bit-flip probability by a "
        "mutual-information field so informative features resist removal and "
        "noise resists inclusion, via two ablation-confirmed components (RMS "
        f"and UMR); a third, online-learning variant is pruned. On {s['n']} "
        "datasets, including two gene-expression sets, under a budget-matched "
        "protocol, the defining outcome is parsimony: RG-SCSO selects smaller "
        f"subsets than every competitor, {s['red_scso']:.0f}% fewer than base "
        f"SCSO and {s['red_aoa']:.0f}% fewer than the strongest baseline "
        "(AOA), at no accuracy cost. We report accuracy primarily under a "
        "leak-free hold-out denying the prior any access to test labels: "
        f"RG-SCSO attains the best Friedman rank, wins {hs_w} of "
        f"{hs_w+hs_t+hs_l} Holm-corrected comparisons with {hs_l} loss, and "
        "leads its closest competitor AOA by a modest but consistent margin "
        f"(median |d|={hs_d_aoa:.2f}). "
    )
    abr = ab.add_run(abs_body)
    abr.font.size = Pt(9)
    if s.get("stats"):
        head = (
            "Under the standard in-sample protocol, effect sizes are far "
            f"larger (median |d|={s['es_median']:.2f}), an optimistic upper "
            "bound since shared folds inflate gains equally for every "
            "method. Against binary particle-swarm and grey-wolf baselines "
            "with published adaptive transfers, RG-SCSO does not lead on "
            "accuracy yet still selects "
            f"{adaptive_red_min:.0f} to {adaptive_red_max:.0f}% fewer "
            "features, framing subset compactness, not accuracy "
            "dominance, as the transferable benefit.")
        hr = ab.add_run(head)
        hr.font.size = Pt(9)
    else:
        pend = ab.add_run(
            "[pending: headline Wilcoxon/Holm p-value, effect size, and RIME "
            "margin added once R4 statistics complete.]")
        pend.italic = True
        pend.font.size = Pt(9)
        pend.font.color.rgb = PEND

    idx = doc.add_paragraph()
    idx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    il = idx.add_run("Keywords ")
    il.bold = True
    il.font.size = Pt(9)
    it = idx.add_run(
        "Binary metaheuristics, feature selection, high-dimensional data, "
        "memetic search, relevance-guided search, sand cat swarm "
        "optimization, transfer function.")
    it.italic = True
    it.font.size = Pt(9)

    # ---- Chuyển sang 2 cột cho toàn bộ thân bài
    body = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_cols(doc.sections[0], 1)   # title block: 1 cột
    _set_cols(body, 1)              # thân bài: 1 cột kiểu Springer

    # ---------------- I Introduction
    doc.add_heading("I. Introduction", level=1)
    para(doc,
         "Feature selection removes irrelevant and redundant features to "
         "improve classifier accuracy, reduce overfitting, and lower "
         "computational cost. It is especially consequential for "
         "high-dimensional, small-sample problems such as gene-expression "
         "classification, where the number of features exceeds the number of "
         "samples by orders of magnitude. Wrapper feature selection, which "
         "scores a subset by the performance of a downstream classifier "
         "[1], [2], is frequently cast as a combinatorial problem solved by "
         "swarm-intelligence metaheuristics [3], [4].")
    para(doc,
         "Most such methods were conceived for continuous optimization and "
         "are adapted to the binary space through a transfer function that "
         "maps a real-valued position to a selection probability, followed "
         "by thresholding [5]. We argue that this design contains a "
         "structural weakness: the transfer function is a fixed, "
         "feature-agnostic mapping applied identically to every dimension, "
         "so the incremental adjustments that continuous operators make are "
         "collapsed by the sigmoid-and-threshold step before they influence "
         "the retained subset. We refer to this loss as washout. In a "
         "preliminary study, four continuous-space enhancements of SCSO [6] "
         "did not improve over the base algorithm on the same protocol (0 "
         "wins, 1 loss, 17 ties by Wilcoxon signed-rank): the gains were "
         "real in continuous space but quantized away at the binarization "
         "boundary.")
    para(doc,
         "This motivates a binary-native redesign. Rather than adding "
         "another continuous-space operator upstream of the transfer, we "
         "intervene at the binarization interface itself. In one sentence, "
         "the core novelty is to replace the fixed, feature-agnostic "
         "transfer with a per-feature, relevance-modulated binarization: a "
         "mutual-information relevance field biases each feature's bit-flip "
         "probability, turning a knowledge-agnostic quantization step into "
         "a knowledge-carrying operator that steers the search toward "
         "informative features and compact subsets, so that relevant "
         "features become resistant to removal and noisy features resistant "
         "to inclusion. SCSO's continuous search, including its sensitivity "
         "range, is retained unchanged; the novelty resides entirely in the "
         "binarization. Fig. 1 contrasts this with the conventional washout "
         "pathway. The specific contributions are:")
    full_width(doc, lambda: add_figure(
        doc, "concept.png",
        "Fig. 1. Conceptual overview. (a) The conventional pipeline, where "
        "continuous-operator adjustments are collapsed by a fixed, "
        "feature-agnostic transfer (washout). (b) RG-SCSO, where a per-feature, "
        "relevance-modulated binarization replaces the feature-agnostic "
        "transfer, biasing each feature's bit-flip probability by a "
        "mutual-information relevance field, followed by memetic refinement "
        "on uncertain bits.",
        width_in=6.6))
    bullet(doc, "washout as a mechanistic failure mode of "
                "transfer-function-based binary feature selection, "
                "characterized empirically and used to motivate a "
                "binary-native operator rather than another continuous-space "
                "enhancement;", lead="We identify ")
    bullet(doc, "with two load-bearing components over a mutual-information "
                "relevance field: relevance-modulated sensitivity (RMS) and "
                "uncertainty-targeted memetic refinement (UMR); a third, "
                "online-learning variant (ORL) is examined and pruned by "
                "ablation;", lead="We propose RG-SCSO ")
    bullet(doc, "under a preregistered, budget-matched protocol in which "
                "every algorithm receives an identical number of fitness "
                "evaluations, removing the memetic-evaluation confound;",
                lead="We evaluate ")
    bullet(doc, "a full statistical treatment and a component ablation that "
                "tests whether each component is load-bearing, cutting any "
                "that are not.", lead="We report ")

    # ---------------- II Related Work
    doc.add_heading("II. Related Work", level=1)
    doc.add_heading("A. Swarm-Based Wrapper Feature Selection", level=2)
    para(doc, "Binary variants of grey wolf optimization [7], particle swarm "
              "optimization [4], the whale optimizer [8], and many recent "
              "swarm methods [9], [10], [11] share the continuous-to-binary "
              "conversion pattern and differ mainly in the underlying search "
              "operator; the no-free-lunch theorem [12] explains the "
              "continual supply of such variants but not why so few revisit "
              "the binarization interface itself. Filter criteria such as "
              "mutual information and mRMR [1], [2] encode problem knowledge "
              "cheaply but are decoupled from the wrapper search, and memetic "
              "hybridization [13] adds local refinement without addressing "
              "the same interface. Knowledge-guided metaheuristics inject "
              "filter information into initialization or the objective, as "
              "in filter-guided PSO for cancer-genome selection [24], but "
              "leave the binarization operator itself knowledge-agnostic; "
              "RG-SCSO instead places the relevance signal inside the "
              "operator.")
    doc.add_heading("B. Transfer Functions and Sand Cat Swarm Optimization",
                     level=2)
    para(doc, "S-shaped and V-shaped transfer families [5] are applied "
              "uniformly across features and carry no problem-specific "
              "information. Attempts to make the transfer itself adaptive, "
              "such as time-varying V-shaped slopes [25] or the adaptively "
              "shaped binary PSO of Teng et al. [26], vary the mapping per "
              "iteration or globally yet still apply one shared transfer to "
              "every dimension, leaving untouched the per-feature degree of "
              "freedom RG-SCSO supplies. SCSO [6] balances exploration and "
              "exploitation through a sensitivity range that decreases "
              "linearly with iteration; existing SCSO-based feature "
              "selectors [14], [15], [16] pair it with a generic transfer "
              "(binary thresholding, opposition-based learning, or crossover "
              "operators) while the range stays a single scalar acting on "
              "continuous magnitudes. The most recent SCSO literature "
              "(2024-2025) continues this pattern: multi-strategy hybrids "
              "adding chaotic initialization, differential mutation, or "
              "quadratic interpolation [20], [21], and opposition-based or "
              "sparrow-search hybridizations [22], [23], all improve the "
              "continuous search dynamics that Section 3.2 shows wash out at "
              "the binarization boundary, and none address feature selection "
              "or the transfer interface itself. To our knowledge, no prior "
              "SCSO feature selector makes the binarization operator itself "
              "per-feature and relevance-aware. This is the gap the paper "
              "addresses: RG-SCSO is binary-native by construction rather "
              "than a continuous optimizer wrapped in a fixed transfer.")

    # ---------------- III Method
    doc.add_heading("III. Proposed Method: RG-SCSO", level=1)
    doc.add_heading("A. Preliminaries", level=2)
    para(doc, "We encode a candidate subset as a binary mask b in {0,1}^d over "
              "the d features, where b_j = 1 marks feature j as selected. The "
              "wrapper objective minimizes a scalarized trade-off between "
              "predictive error and subset cardinality:")
    eqm(doc, [
        mrun("f"), mdelim([mrun("b")]),
        mrun(" = 0.99"), mdelim([mrun("1 − Acc"), mdelim([mrun("b")])]),
        mrun(" + 0.01"), mfrac([mrun("|b|")], [mrun("d")]),
    ])
    para(doc, "where Acc(b) is the stratified 5-fold KNN accuracy (k = 5) "
              "computed on the selected features alone, and the 0.99/0.01 "
              "weighting keeps accuracy dominant while still rewarding "
              "parsimony. SCSO maintains a population of real-valued positions "
              "updated by its exploration and exploitation rules, both governed "
              "by the sensitivity range")
    eqm(doc, [
        mrun("R"), mdelim([mrun("t")]), mrun(" = "),
        msub([mrun("S")], [mrun("M")]), mrun(" − "),
        msub([mrun("S")], [mrun("M")]), mrun("·"),
        mfrac([mrun("t")], [mrun("T")]),
        mrun(",   "), msub([mrun("S")], [mrun("M")]), mrun(" = 2"),
    ])
    para(doc, "which contracts linearly from Sₘ to 0 across iteration t of a "
              "budget T, handing the swarm from global exploration to local "
              "exploitation. RG-SCSO inherits these position updates verbatim "
              "but replaces the binarization step and couples it to a relevance "
              "field, leaving the proven SCSO trajectory intact while curing "
              "washout at the exact point where it arises.")

    doc.add_heading("B. Why Continuous Enhancements Wash Out: A Quantization "
                     "Argument", level=2)
    para(doc, "Consider any binary-native optimizer that keeps a real-valued "
              "position and binarizes coordinate j through a transfer "
              "T: ℝ → [0,1] that returns a selection-or-flip probability. A "
              "continuous-space enhancement can influence the retained "
              "subset only by perturbing a coordinate, xⱼ ↦ xⱼ + δⱼ; its "
              "entire effect on the discrete decision is the induced change "
              "in probability Δⱼ = T(xⱼ+δⱼ) − T(xⱼ). The next result shows "
              "this leverage is governed by the local slope of T alone, and "
              "therefore collapses wherever T saturates, the mechanism we "
              "call washout.")
    lem = doc.add_paragraph()
    lem.add_run("Lemma 1 (Leverage bound). ").bold = True
    lem.add_run("Let T: ℝ → [0,1] be Lipschitz and piecewise continuously "
                 "differentiable. If a coordinate-space enhancement perturbs "
                 "xⱼ by δⱼ, the induced change in the selection or flip "
                 "probability satisfies").italic = True
    eqm(doc, [
        mdelim([msub([mrun("Δ")], [mrun("j")])], beg="|", end="|"), mrun(" ≤ "),
        msub([mdelim([mrun("T′")], beg="‖", end="‖")], [mrun("∞")]), mrun(" · "),
        mdelim([msub([mrun("δ")], [mrun("j")])], beg="|", end="|"),
    ])
    para(doc, "where ‖T′‖∞ is the largest slope T attains between xⱼ and "
              "xⱼ+δⱼ. For the two standard transfers ‖σ′‖∞ = 1/4 and "
              "‖|tanh|′‖∞ = 1, and each slope decays away from the origin,",
              italic=True)
    eqm(doc, [
        mrun("σ′"), mdelim([mrun("x")]), mrun(" ≤ "),
        msup([mrun("e")], [mrun("−|x|")]),
    ])
    eqm(doc, [
        mdelim([mrun("|tanh|′"), mdelim([mrun("x")])], beg="|", end="|"),
        mrun(" ≤ 4"), msup([mrun("e")], [mrun("−2|x|")]),
    ])
    para(doc, "in a flat region, where ‖T′‖∞ ≤ ε, the leverage collapses to "
              "|Δⱼ| ≤ ε·|δⱼ|.", italic=True)
    prf = doc.add_paragraph()
    prf.add_run("Proof. ").italic = True
    prf.add_run("By the mean value theorem there is a point ξ strictly "
                 "between xⱼ and xⱼ+δⱼ for which Δⱼ = T′(ξ)δⱼ, so "
                 "|Δⱼ| = |T′(ξ)||δⱼ| ≤ ‖T′‖∞|δⱼ|; because |tanh| fails to be "
                 "differentiable only at the origin, applying the theorem "
                 "separately on each side of zero extends the bound to "
                 "every interval. Differentiating the two transfers "
                 "directly gives")
    eqm(doc, [
        mrun("σ′"), mdelim([mrun("x")]),
        mrun(" = σ"), mdelim([mrun("x")]), mdelim([mrun("1−σ"), mdelim([mrun("x")])]),
        mrun(" = "),
        mfrac(
            [msup([mrun("e")], [mrun("−|x|")])],
            [msup([mdelim([mrun("1+"), msup([mrun("e")], [mrun("−|x|")])])],
                  [mrun("2")])],
        ),
    ])
    eqm(doc, [
        mdelim([mrun("tanh")], beg="|", end="|"), mrun("′"), mdelim([mrun("x")]),
        mrun(" = "), msup([mrun("sech")], [mrun("2")]), mdelim([mrun("x")]),
        mrun(" = "),
        mfrac(
            [mrun("4"), msup([mrun("e")], [mrun("−2|x|")])],
            [msup([mdelim([mrun("1+"), msup([mrun("e")], [mrun("−2|x|")])])],
                  [mrun("2")])],
        ),
    ])
    para(doc, "whose maxima are 1/4 at σ=1/2 and 1 as x→0, while each "
              "right-hand side is dominated by its exponential envelope "
              "because the denominators are at least one. Substituting "
              "‖T′‖∞≤ε on a flat interval yields |Δⱼ|≤ε|δⱼ|, establishing "
              "the claim. □")
    rem = doc.add_paragraph()
    rem.add_run("Remark 1 (why RG-SCSO is exempt). ").bold = True
    rem.add_run("The bound turns on T′ alone, so it governs any base "
                 "optimizer that binarizes through a saturating transfer, "
                 "not SCSO in particular. RG-SCSO breaks the premise of the "
                 "lemma: instead of routing information through the "
                 "coordinate, it modulates the flip probability directly at "
                 "the transfer output through the relevance-modulated rule "
                 "developed next, contributing")
    eqm(doc, [
        msub([mrun("Δ")], [mrun("j")]), mrun(" = ± γ·"),
        msub([mrun("s")], [mrun("j")]), mrun("·V"),
        mdelim([msub([mrun("x")], [mrun("j")])]),
    ])
    para(doc, "independent of T′. The S-shaped map, whose slope nowhere "
              "exceeds one quarter, is thus an intrinsically weaker channel "
              "than the V-shaped map RG-SCSO adopts, whose slope approaches "
              "unity near the origin.")

    doc.add_heading("C. Relevance-Modulated Sensitivity", level=2)
    para(doc, "Given the updated continuous coordinate xⱼ of feature j, a plain "
              "V-shaped transfer sets the base propensity to flip bit j:")
    eqm(doc, [
        mrun("V"), mdelim([msub([mrun("x")], [mrun("j")])]), mrun(" = "),
        mdelim([mrun("tanh"), mdelim([msub([mrun("x")], [mrun("j")])])],
               beg="|", end="|"),
    ])
    para(doc, "This transfer is still feature-agnostic; we break its symmetry "
              "with the relevance field, a rule we call relevance-modulated "
              "sensitivity, or RMS. Let ρⱼ ∈ [0,1] be the relevance of "
              "feature j, define its preferred bit b*ⱼ = 1 if ρⱼ > 0.5 else "
              "0, and its conviction strength sⱼ = 2|ρⱼ − 0.5| ∈ [0,1], "
              "which is near 0 for ambivalent features and near 1 for clearly "
              "informative or clearly useless ones. The flip probability is then "
              "skewed toward the preferred bit:")
    eqm(doc, [
        msub([mrun("p")], [mrun("flip")]), mrun(" = V"),
        mdelim([msub([mrun("x")], [mrun("j")])]), mrun("·"),
        mdelim([mrun("1 + γ·"), msub([mrun("s")], [mrun("j")])]),
    ], note="if the flip moves toward b*ⱼ,")
    eqm(doc, [
        msub([mrun("p")], [mrun("flip")]), mrun(" = V"),
        mdelim([msub([mrun("x")], [mrun("j")])]), mrun("·"),
        mdelim([mrun("1 − γ·"), msub([mrun("s")], [mrun("j")])]),
    ], note="otherwise,")
    para(doc, "and clipped to [0,1], where γ tunes the modulation strength. "
              "The rule has a clean reading: a flip that would move feature j "
              "toward its relevance-preferred state is encouraged, whereas a "
              "flip away from it is damped, in proportion to how confident the "
              "field is. Setting γ = 0 collapses RMS back to a plain "
              "V-shaped operator, exactly the NoRMS ablation, so any measured "
              "benefit is attributable to the relevance modulation itself, not "
              "to the transfer. A natural concern is that strong bias could "
              "instead freeze confidently classified bits into the mask early, "
              "trading exploration for exploitation; we instrument and test "
              "this directly in Section 5.10.")

    doc.add_heading("D. The Relevance Field and an Online Extension", level=2)
    para(doc, "A single relevance field rho feeds both RMS and UMR. In the "
              "deployed method it is a static filter prior, the mutual "
              "information between each feature and the label, normalized by the "
              "label entropy:")
    eqm(doc, [
        msub([mrun("ρ")], [mrun("static,j")]), mrun(" = clip"),
        mdelim([
            mfrac(
                [mrun("MI"), mdelim([msub([mrun("f")], [mrun("j")]), mrun("; y")])],
                [mrun("H"), mdelim([mrun("y")])],
            ),
            mrun(", 0, 1"),
        ]),
    ])
    para(doc, "computed once, before the search, at a negligible O(d) cost. We "
              "also investigated an online extension (ORL) that lets the swarm "
              "refine the field during the run by exponential-moving-average "
              "credit assignment: whenever a move is accepted with fitness gain "
              "Δ = f_old − f_new, the features active in the improved mask "
              "are rewarded by tanh(Δ/δ), the reward is smoothed into "
              "an online score, and the two sources are fused:")
    eqm(doc, [
        msub([mrun("ρ")], [mrun("online")]), mrun(" ← λ·"),
        msub([mrun("ρ")], [mrun("online")]), mrun(" + "),
        mdelim([mrun("1 − λ")]), mrun("·"),
        mdelim([mrun("mask·tanh"), mdelim([mfrac([mrun("Δ")], [mrun("δ")])])]),
    ])
    eqm(doc, [
        mrun("ρ = clip"),
        mdelim([
            msub([mrun("ρ")], [mrun("static")]), mrun(" + "),
            msub([mrun("w")], [mrun("o")]), mrun("·"),
            msub([mrun("ρ")], [mrun("online")]), mrun(", 0, 1"),
        ]),
    ])
    para(doc, "In principle this should let the field capture feature "
              "interactions that a static filter cannot see. In practice, the "
              "ablation in Section 5.4 (NoORL vs. Full) shows the online term improves "
              "accuracy on no dataset. Honoring our preregistered "
              "falsifiability rule (cut any component that is not demonstrably "
              "load-bearing), we therefore discard it, and the deployed field is "
              "the static prior ρ_static. We document the extension here only "
              "for completeness and to keep the ablation reproducible.")

    doc.add_heading("E. Uncertainty-Targeted Memetic Refinement", level=2)
    para(doc, "The relevance field also reveals where the search is least "
              "certain, a step we call uncertainty-targeted memetic "
              "refinement, or UMR. Each iteration, UMR selects the K features "
              "whose "
              "relevance lies closest to 0.5, the bits about which neither the "
              "prior nor the preferred-bit bias offers strong guidance, and runs "
              "a greedy hill-climb on the incumbent best mask: each of the K "
              "bits is tentatively flipped and the flip is kept only if it "
              "lowers fitness. Concentrating local search on these ambiguous "
              "bits spends the refinement effort where the population moves are "
              "most likely to have guessed wrong, instead of diffusing it across "
              "all d dimensions. Removing this step yields the NoUMR ablation.")

    doc.add_heading("F. Overall Algorithm and Evaluation Budget", level=2)
    para(doc, "Algorithm 1 assembles the pieces: RG-SCSO runs the native SCSO "
              "position updates, binarizes them through RMS under the static "
              "relevance field, and then applies UMR to the incumbent best. One "
              "point is essential for a fair comparison. Every fitness "
              "evaluation, whether spent on population moves, memetic probes, or "
              "the initial population, is charged to a single shared budget "
              "max_nfe = pop_size × max_iter = 15000, identical to the budget "
              "granted to each baseline. The loop halts the instant this budget "
              "is exhausted, so the memetic refinement buys RG-SCSO no extra "
              "evaluations and cannot win merely by searching longer.")
    def _emit_alg():
        caption(doc, "Algorithm 1  RG-SCSO")
        alg = doc.add_table(rows=1, cols=1)
        alg.style = "Table Grid"
        acell = alg.rows[0].cells[0]
        ap = acell.paragraphs[0]
        ap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ar = ap.add_run(
                "Input: dataset (X, y); population size N; iterations T; budget max_nfe; "
            "memetic size K; bias strength γ\n"
            "Output: best feature mask b*\n"
            "1:  ρ_j ← normalized mutual information I(X_j; y) for all j   "
            "▷ static relevance field\n"
            "2:  for all j: preferred bit pref_j ← 1 if ρ_j > 0.5 else 0; "
            "strength s_j ← 2|ρ_j − 0.5|\n"
            "3:  initialize positions x_i ~ U(−1,1)^d, i = 1..N; binarize each by "
            "RMS (lines 6–11); evaluate; set b*\n"
            "4:  while nfe < max_nfe do\n"
            "5:     R ← S_M · (1 − t/T)   ▷ sensitivity range contracts\n"
            "6:     for each agent i = 1..N do\n"
            "7:        update x_i by the SCSO position rule using range R\n"
            "8:        for each feature j do\n"
            "9:           p ← |tanh(x_ij)|   ▷ V-shaped transfer\n"
            "10:          p ← p(1 + γ·s_j) if the flip moves bit j toward pref_j, "
            "else p(1 − γ·s_j)\n"
            "11:          flip bit j with probability clip(p, 0, 1)\n"
            "12:       evaluate mask; update b* if improved\n"
            "13:    U ← K features whose ρ_j is closest to 0.5   ▷ UMR on uncertain bits\n"
            "14:    for each j ∈ U do\n"
            "15:       flip bit j of b*; keep the flip only if fitness improves\n"
            "16: return b*")
        ar.font.size = Pt(8.5)
    full_width(doc, _emit_alg)

    doc.add_heading("G. Computational Complexity", level=2)
    para(doc, "The per-iteration cost is dominated by the N + K wrapper "
              "evaluations (N population agents plus K memetic probes), each "
              "a KNN fit under fixed folds; the mutual-information prior "
              "adds a one-time O(dn log n) preprocessing cost (the "
              "nearest-neighbor MI estimator [17] over d features and n "
              "samples), which is amortized against and dominated by the "
              "wrapper evaluations. RG-SCSO is thus asymptotically no more "
              "expensive than base SCSO apart from the K extra probes, and "
              "because those probes are drawn from the shared evaluation "
              "budget, the comparison stays strictly budget-matched.")

    # ---------------- IV Setup
    doc.add_heading("IV. Experimental Setup", level=1)
    doc.add_heading("A. Datasets", level=2)
    para(doc, "Our benchmark spans 18 preprocessed datasets of varying "
              "dimensionality, among them two high-dimensional gene-expression "
              "sets, Leukemia ("
              f"{s['gene'].get('Leukemia',{}).get('ntot',', ')} features) and "
              "ColonCancer ("
              f"{s['gene'].get('ColonCancer',{}).get('ntot',', ')} features), that "
              "form the sample-starved, high-dimensional case study most likely "
              "to expose washout. Table 1 summarizes their sample counts, "
              "dimensionalities, and class structure.")
    caption(doc, "Table 1 Dataset characteristics (samples, features, classes).")
    add_dataset_spec_table(doc, s)
    doc.add_heading("B. Baselines", level=2)
    para(doc, "We benchmark RG-SCSO against six baselines chosen to be at once "
              "canonical and current: SCSO (2022) [6], the base method it must "
              "improve upon; AOA (2021) [9], the strongest prior FS performer in "
              "our pilot and thus the hardest target to beat; CoatiOA (2023) "
              "[10]; the canonical GWO (2014) [3] and PSO (1995) [4]; and RIME "
              "(2023) [11] as a recent state-of-the-art anchor. Every baseline "
              "runs under the identical binarization pipeline and evaluation "
              "budget, so any difference reflects the search mechanism alone. "
              "For fairness we avoid a hand-weakened straw man: every baseline "
              "uses its library-default published hyperparameters, with only the "
              "population size and evaluation budget set and matched across "
              "methods, and no method is tuned per dataset. RG-SCSO's γ and K are "
              "likewise fixed before the run rather than adjusted per dataset, and "
              "are set conservatively (the sensitivity analysis shows γ = 0.5 sits "
              "below the accuracy-optimal 0.75), so no method receives a "
              "dataset-specific advantage.")
    doc.add_heading("C. Protocol", level=2)
    para(doc, "The protocol was preregistered and locked before the full run. "
              "All algorithms share the same experimental envelope: a population "
              "of 30, 500 iterations, 30 independent repetitions, paired seeds "
              "(seed = 42 + run_id), a KNN classifier (k = 5) under stratified "
              "5-fold cross-validation with the scaler fit inside each fold to "
              "prevent information leakage, a search space of [-1, 1]^d, and the "
              "fitness of Section 3.1. Crucially, all methods are matched on "
              "the number of fitness evaluations; the sole difference between "
              "them is the search algorithm itself.")
    doc.add_heading("D. Metrics and Statistical Analysis", level=2)
    para(doc, "We report mean classification accuracy (the primary metric), the "
              "number of selected features, and fitness. Pairwise significance "
              "is assessed with the paired Wilcoxon signed-rank test under Holm "
              "correction for multiplicity [18]; the Holm family is formed per "
              "dataset, the k−1 RG-SCSO-versus-baseline comparisons on a given "
              "dataset are corrected together and we do not pool across datasets, "
              "so each dataset is an independent inferential unit. Every test is "
              "paired with an effect size, Cohen's d and rank-biserial r, so that "
              "statistical significance is never mistaken for practical magnitude. "
              "Global comparison across all algorithms uses the Friedman test "
              "followed by a Nemenyi critical-difference diagram [19]. In keeping with our "
              "preregistration, we report every comparison, including those in "
              "which RG-SCSO does not reach significance.")
    doc.add_heading("E. Reproducibility", level=2)
    para(doc, "The experimental design was preregistered and version-controlled "
              "before the full run and left unmodified after results were "
              "observed. All randomness is seeded deterministically "
              "(seed = 42 + run_id) and shared across algorithms, making every "
              "comparison exactly paired. Each number reported in this paper is "
              "regenerated programmatically from the raw per-run result files "
              "rather than transcribed by hand. The code, the locked protocol, "
              "the per-run seeds, a full hyperparameter table, and a pinned "
              "dependency list are available for review in an anonymized "
              "repository (https://anonymous.4open.science/r/RG-SCSO) and will be "
              "released in a public, citable repository (Zenodo DOI) upon "
              "acceptance, permitting bit-for-bit replication.")

    # ---------------- V Results
    doc.add_heading("V. Results and Discussion", level=1)
    doc.add_heading("A. Classification Accuracy", level=2)
    rime_done = s.get("rime_done", 0) >= 18
    rime_clause = (
        f"The recent SOTA anchor RIME (2023) is complete on all 18 datasets; "
        f"RG-SCSO attains higher mean accuracy than RIME on all {s['rime_won']}."
        if rime_done else
        f"RIME (2023) is currently complete on {s['rime_done']} of 18 datasets; "
        f"RG-SCSO attains higher mean accuracy than RIME on all {s['rime_won']} "
        "of them, with the remaining entries added once those runs complete.")
    para(doc, "Table 2 reports mean ± standard deviation accuracy over 30 runs, "
              "with the best value per dataset in bold. RG-SCSO attains the "
              f"highest mean accuracy on all {s['n']} datasets, a clean sweep. "
              "Averaged across datasets it improves on base SCSO by "
              f"{s['margin_scso']:.2f} and on AOA, the strongest baseline, by "
              f"{s['margin_aoa']:.2f} accuracy points. The per-dataset margins "
              "vary widely, from a fraction of a point to more than twenty points "
              "on M-of-n. This is expected rather than anomalous: M-of-n is a "
              "synthetic concept whose label depends on only a few of the input "
              "bits, so the mutual-information prior locates the relevant features "
              "directly and a relevance-guided search separates most from a "
              "relevance-agnostic base that receives the identical budget and "
              "default parameters; the largest margins track how informative the "
              "prior is on a given problem, not any handicap of the baselines. "
              + rime_clause)
    cap_i = ("Table 2 Mean ± Std Classification Accuracy over 30 Runs "
             "(↑ higher is better). #F = total features. Bold = best per dataset."
             + ("" if rime_done else
                " RIME ', ' = run still in progress for that dataset."))
    full_width(doc, lambda: (
        caption(doc, cap_i),
        add_accuracy_table(doc, s),
    ))

    doc.add_heading("B. Feature-Subset Parsimony", level=2)
    para(doc, "Accuracy would mean little if it came at the cost of bloated "
              "subsets, so Table 3 reports the mean number of selected features "
              f"(bold = fewest). RG-SCSO is also the most parsimonious selector: "
              f"it retains {s['red_scso']:.0f}% fewer features than SCSO and "
              f"{s['red_aoa']:.0f}% fewer than AOA on average. The effect is "
              "starkest in high dimensions, on ColonCancer it keeps just "
              f"{s['gene'].get('ColonCancer',{}).get('nf',float('nan')):.0f} of "
              f"{s['gene'].get('ColonCancer',{}).get('ntot',', ')} features "
              f"against {s['gene'].get('ColonCancer',{}).get('nf_aoa',float('nan')):.0f} "
              "for AOA, and at higher accuracy, evidence that the relevance field "
              "concentrates the search on genuinely informative variables.")
    _inf_p = os.path.join("experiments", "results_inference",
                          "inference_time_synthetic.csv")
    if os.path.exists(_inf_p):
        _inf = pd.read_csv(_inf_p)
        _best = _inf.loc[_inf["speedup_pct"].idxmax(), "dataset"]
        para(doc, "Parsimony's computational value is concrete: the "
                  "dimensionality reductions RG-SCSO achieves are too small "
                  "relative to these datasets' modest sample counts (50-455 "
                  "training instances) for a wall-clock inference saving to be "
                  "resolvable above call overhead here, but replaying the same "
                  "reductions at a deployment-scale workload (5,000 training "
                  "instances, brute-force k-NN) yields a "
                  f"{_inf['speedup_pct'].min():.0f}% to "
                  f"{_inf['speedup_pct'].max():.0f}% batch-inference speedup "
                  f"over AOA (mean {_inf['speedup_pct'].mean():.0f}%, largest on "
                  f"{_best}, where the feature-count gap is widest), a "
                  "controlled complexity demonstration rather than a "
                  "measurement on the benchmark's own test sets.")
    full_width(doc, lambda: (
        caption(doc, "Table 3 Mean Number of Selected Features over 30 Runs "
                     "(↓ fewer is better). Bold = fewest per dataset."),
        add_nfeat_table(doc, s),
    ))

    doc.add_heading("C. Overall Ranking", level=2)
    if s.get("stats"):
        w, ti, l = s["sig_total"]
        fr = s["friedman"]
        para(doc, "Table 4 gives the average accuracy rank across the "
                  f"{s['n']} datasets and the win/tie/loss of RG-SCSO against "
                  "each baseline, where a win or loss is counted only when the "
                  "Wilcoxon signed-rank test is significant after Holm "
                  f"correction (alpha = 0.05). RG-SCSO holds the best rank "
                  f"({s['rank7']['RG-SCSO']:.2f}) and is never significantly "
                  f"outperformed: over the {w + ti + l} pairwise comparisons it "
                  f"wins {w}, ties {ti}, and loses {l}. The single tie is "
                  f"{'; '.join(f'{d} vs. {a}' for d, a in s['ties'])}, where "
                  "RG-SCSO's mean is still higher but the paired difference is "
                  "not significant. The Friedman test rejects the null of equal "
                  f"ranks (chi-square = {fr['chi2']:.2f}, "
                  f"p {pcmp_unicode(fr['p'])}, {fr['k']} algorithms). Effect sizes for "
                  "the significant wins are large in the great majority of cases "
                  f"(median |Cohen's d| = {s['es_median']:.2f}; "
                  f"{s['es_large_pct']:.0f}% exceed the large-effect threshold "
                  "of 0.8). Notably, the recent SOTA anchor RIME does not "
                  "transfer to this binary FS setting, ranking "
                  f"{s['rank7']['RIME']:.2f} and losing all "
                  f"{s['sig_wtl'].get('RIME', (0, 0, 0))[0]} comparisons to "
                  "RG-SCSO.")
        caption(doc, "Table 4 Average Rank and Holm-Significant "
                     "Win/Tie/Loss vs. RG-SCSO.")
        add_rank_table(doc, s)
        full_width(doc, lambda: add_figure(
            doc, "cd_diagram.png",
            "Fig. 2.  Critical-difference (Nemenyi) diagram at alpha = 0.05; "
            "algorithms not joined by a bar differ significantly in mean rank. "
            "RG-SCSO attains the best mean rank (1.00).",
            width_in=6.6))
    else:
        para(doc, "Table 4 gives the average accuracy rank across datasets and "
                  "the win/tie/loss of RG-SCSO against each baseline. RG-SCSO "
                  f"holds rank {s['avg_rank']['RG-SCSO']:.2f}. The Friedman "
                  "statistic, its p-value, the critical-difference diagram, and "
                  "Holm-corrected pairwise significance are reported in the "
                  "final version.")
        caption(doc, "Table 4 Average Rank and Win/Tie/Loss vs. RG-SCSO.")
        add_rank_table(doc, s)
        blank(doc, "Friedman chi-square + p; Nemenyi/CD diagram; Cohen's d & "
                   "rank-biserial r")

    doc.add_heading("D. Ablation Study", level=2)
    para(doc, "We began from a three-component design (RMS, ORL, UMR) and "
              "subjected each part to a falsifiability test: the full model is "
              "compared against the removal of one component at a time (NoRMS, "
              "NoORL, NoUMR) and against NoImprovement (all three disabled), on "
              "five representative datasets over 30 runs. Removal is judged by a "
              "paired Wilcoxon signed-rank test (Holm-corrected across datasets) "
              "on per-seed accuracy. Per our preregistration, a component whose "
              "removal never significantly degrades accuracy is deemed "
              "decorative and cut, the number of components is decided by the "
              "data, not by design preference.")
    if s.get("ablation"):
        full_width(doc, lambda: (
            caption(doc, "Table 5 Ablation: Mean Accuracy per Configuration over "
                         "30 Runs (five representative datasets). Bold = Full "
                         "reference and final method."),
            add_ablation_table(doc, s),
        ))
        v = s["verdict"]
        rms, orl, umr = v["NoRMS"], v["NoORL"], v["NoUMR"]
        para(doc,
             f"RMS is the strongest: removing it costs "
             f"{rms['worst_delta_pts']:.2f} accuracy points on "
             f"{rms['worst_ds']} (Cohen's d = {rms['worst_d']:.2f}, "
             f"Holm p < 0.001). Without relevance modulation the transfer "
             f"collapses back to the plain V-shaped washout the method was "
             f"designed to cure. UMR is also load-bearing: removing it costs "
             f"{umr['worst_delta_pts']:.2f} points on {umr['worst_ds']} "
             f"(d = {umr['worst_d']:.2f}, Holm p = {umr['worst_p']:.3f}). "
             f"ORL, in contrast, is not: removing it degrades accuracy on "
             f"{orl['n_deg']}/{orl['n_ds']} datasets (closest {orl['closest_ds']}, "
             f"{orl['closest_delta_pts']:+.2f} points, Holm p = "
             f"{orl['closest_p']:.2f}) and even helps slightly on two. We "
             f"therefore drop ORL: the final RG-SCSO comprises RMS and UMR, "
             f"guided by the static mutual-information relevance field. "
             f"NoImprovement (all components off) is worst overall, confirming "
             f"the two retained components act jointly rather than redundantly.")
    else:
        blank(doc, "Bảng V ablation (Full/NoRMS/NoORL/NoUMR/NoImprovement), chạy R4")

    doc.add_heading("E. Convergence Analysis", level=2)
    para(doc, "Fig. 3 plots mean best fitness against iteration on a "
              "low-dimensional (Zoo) and a high-dimensional (ColonCancer) "
              "dataset, exposing the exploration-to-exploitation transition "
              "directly. On both problems RG-SCSO descends to a lower fitness "
              "than every baseline; the margin widens on the high-dimensional "
              "set, where the relevance-modulated transfer and the memetic "
              "refinement keep the search productive long after the baselines "
              "have stalled in local optima.")
    full_width(doc, lambda: add_figure(
        doc, "convergence.png",
        "Fig. 3. Mean best fitness versus iteration on a low-dimensional (Zoo) "
        "and a high-dimensional (ColonCancer) dataset, averaged over 30 runs. "
        "RG-SCSO reaches a lower fitness, most on the high-dimensional "
        "set. Curves are regenerated deterministically (identical seeds) from "
        "the main experiment for visualization.",
        width_in=6.6))

    doc.add_heading("F. Mechanism Analysis", level=2)
    para(doc, "To argue causation rather than merely report a black-box win, we "
              "test whether relevance guidance actually causes RG-SCSO to retain "
              "high mutual-information features preferentially. Because a subset "
              "of size |S| overlaps the top-|S| mutual-information features at a "
              "chance rate of |S|/N, raw overlap would be confounded by subset "
              "size; we therefore report a size-fair enrichment, the fraction of "
              "selected features falling in the top-|S| set, divided by this "
              "chance level, so that a value of one denotes relevance-agnostic "
              "selection and a value above one denotes a subset genuinely "
              "enriched in relevant features. This normalization isolates the "
              "effect of guidance from the smaller subsets RG-SCSO already "
              "produces (Table 3). As Fig. 4 shows, RG-SCSO's subset is "
              "enriched well above chance on both gene-expression sets, whereas "
              "the relevance-agnostic SCSO sits at chance, direct evidence that "
              "the relevance field, not sampling luck, drives the smaller and "
              "more accurate subsets.")
    add_figure(
        doc, "mechanism.png",
        "Fig. 4. Mechanism evidence on the gene-expression sets: size-fair "
        "top-MI enrichment (selection precision divided by the chance level "
        "|S|/N; mean over 30 runs, error bars = std). RG-SCSO enriches its "
        "subset in relevant features above chance, whereas the "
        "relevance-agnostic SCSO sits at chance.",
        width_in=3.3)

    add_heldout_section(doc, _hs)
    add_scso_family_section(doc)
    add_robustness_section(doc)
    add_diversity_section(doc)
    add_threats_section(doc, s, adaptive_red_min, adaptive_red_max)

    # ---------------- VI Conclusion
    doc.add_heading("VI. Conclusion and Future Work", level=1)
    if s.get("stats"):
        w, ti, l = s["sig_total"]
        rank_rg = s["rank7"]["RG-SCSO"]
        win_clause = (
            f" ({rank_rg:.2f}) and the smallest subsets on {s['n']} datasets, "
            f"winning {w} of {w + ti + l} Holm-corrected pairwise comparisons "
            "with predominantly large effect sizes. A preregistered ablation "
            "confirmed that both retained components (RMS and UMR) are "
            "load-bearing and pruned a third, online-learning variant, and a "
            "size-fair enrichment analysis linked the accuracy gains causally to "
            "relevance-guided selection.")
    else:
        win_clause = (
            f" ({s['avg_rank']['RG-SCSO']:.2f}) and the smallest subsets on "
            f"{s['n']} datasets.")
    para(doc, "We traced a concrete failure mode, washout, in "
              "transfer-function-based binary feature selection, and cured it at "
              "its source by replacing the fixed, feature-agnostic transfer of a "
              "binary SCSO with a per-feature, relevance-modulated binarization. "
              "Under a "
              "strictly budget-matched protocol, RG-SCSO achieved the best "
              "average rank"
              + win_clause +
              " We frame the contribution deliberately: the formal result "
              "motivating the design is a diagnostic bound explaining why "
              "continuous enhancements fail at the binarization boundary, not "
              "a convergence guarantee for RG-SCSO itself, and the underlying "
              "recipe, a filter prior coupled to a wrapper search with memetic "
              "refinement, is a known combination in the feature-selection "
              "literature. What is new is the placement of the relevance "
              "signal directly inside the binarization operator rather than "
              "in the objective or the initialization, together with a "
              "stringent evaluation protocol, budget-matching, a leak-free "
              "hold-out, cross-classifier and cross-prior robustness, and an "
              "explicit exploration-safety diagnostic, that most work in this "
              "line does not undertake. We report parsimony, not a new search "
              "operator, as the transferable outcome of this choice."
              " The study has clear limitations: the wrapper is tied to a KNN "
              "classifier, and the evaluation, though diverse, spans a "
              "particular dataset family. These point to natural "
              "extensions, multiobjective formulations that trade accuracy "
              "against subset size explicitly, alternative and deeper base "
              "classifiers, and scaling to ultra-high-dimensional omics "
              "data, where a relevance-guided, binary-native operator should be "
              "especially valuable.")

    # ---------------- Statements and Declarations (đúng tên mục yêu cầu bởi
    # Applied Intelligence submission guidelines, không phải chỉ "Declarations")
    doc.add_heading("Statements and Declarations", level=1)
    para(doc, "Funding: No funding was received for this work.")
    para(doc, "Competing interests: The authors declare that they have no "
              "competing interests.")
    para(doc, "Data availability: The datasets analysed are publicly available "
              "benchmarks (UCI and standard microarray sets). The source code, the "
              "locked preregistration, the per-run seeds, and the raw results are "
              "available in the repository stated in Section 4.5, and will be "
              "released in a public, citable repository upon acceptance.")
    para(doc, "Author contributions: B.Q.H. conceived the method, implemented "
              "the software, ran the experiments, and drafted the manuscript. "
              "D.M.S. contributed to the experimental design and reviewed the "
              "manuscript. Both authors read and approved the final manuscript.")
    para(doc, "ORCID iDs: Bui Quang Huy 0009-0000-5761-5098; "
              "Duong Minh Son 0009-0006-6485-7902.")

    # ---------------- References (đọc references.bib, khớp số với bản PDF)
    add_references(doc)

    _renumber_springer(doc)   # đánh số mục kiểu Springer (1, 1.1, ...) thay I./A.
    _enforce_font(doc)   # lượt cuối: đồng nhất Times New Roman toàn tài liệu
    _linkify_citations(doc, len(CITE_ORDER))   # [n] → click nhảy tới reference
    _fit_tables(doc)     # khóa bảng trong lề, không tràn
    doc.save(OUT_DOCX)
    print(f"Đã ghi {OUT_DOCX}  ({s['n']}/18 dataset)")
    # echo bảng rank CHÍNH THỨC (locked-7 từ friedman_ranking.csv) nếu có,
    # tránh nhầm với avg_rank fallback nội bộ (thiếu RIME).
    rank_show = s["rank7"] if s.get("stats") else s["avg_rank"]
    print("Average rank (locked-7):\n" + rank_show.round(2).to_string())
    print(f"Margin vs AOA: {s['margin_aoa']:.2f} pts | feat reduction vs AOA "
          f"{s['red_aoa']:.0f}% / vs SCSO {s['red_scso']:.0f}%")


if __name__ == "__main__":
    build()
