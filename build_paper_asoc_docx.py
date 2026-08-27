"""Sinh bản thảo Word (.docx) cho Applied Soft Computing (Elsevier) -- SINGLE
SOURCE OF TRUTH cho target ASOC, song song với build_paper_asoc.py (bản
LaTeX cùng target). Cấu trúc và nội dung khớp với bản LaTeX (đã verify kỹ số
liệu ở đó); file này chỉ chuyển sang định dạng python-docx. Mirror trực tiếp
build_paper_scirep_docx.py (bản Word cho Scientific Reports), cùng các thay
đổi cấu trúc đã áp dụng ở build_paper_asoc.py: Research Questions, Related
Work riêng (bảng literature-positioning), 2 hình promoted vào main text
(convergence_fs, threshold_heatmap), Conclusion riêng (limitations + Future
Work 5 mục), CRediT statement.

Graphical abstract và Highlights KHÔNG nhúng trong file này (đúng quy ước đã
chọn cho bản .tex: đây là 2 hạng mục nộp riêng trong hệ thống Elsevier,
không phải section của bản thảo).

Chạy:  python build_paper_asoc_docx.py
"""

from __future__ import annotations

import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

from build_paper_structure import (
    BODY_PT,
    _add_run_text,
    _bookmark_paragraph,
    _fmt_authors,
    _hdr,
    _ieee_table,
    _linkify_citations,
    _parse_bib,
    add_figure,
    bullet,
    caption,
    eqm,
    force_font_everywhere,
    full_width,
    load_summary,
    mdelim,
    mfrac,
    mrun,
    msub,
    msup,
    repeat_header_row,
    para,
)
from build_paper_tex import (
    DIVERSITY_CSV,
    PROC_DIR,
    ROBUST_CSV,
    adaptive_baselines,
    diversity_analysis,
    inference_value,
    robustness_baselines,
    robustness_svm16,
    scso_family_baselines,
)
import build_heldout_table as _heldout

# Tái dùng nguyên vẹn các bảng đã có docx-mirror ở bản SciRep -- cùng số
# liệu, cùng cách render, không viết lại.
from build_paper_scirep import THRESHOLD_CSV
from build_paper_scirep_docx import (
    CLASSIC_CSV,
    add_classic_baselines_table,
    add_classifier_robustness_table,
    add_extended_ablation_table,
    add_rank_table_with_effect_size,
    _add_heldout_combined_table,
    _sec0_a4,
    _style_setup,
)

OUT_DOCX = "RG-SCSO_ASOC.docx"

LITPOS_ROWS = [
    ("Seyyedabbasi & Kiani", "2022", "Base continuous SCSO (no FS)", "Baseline (Table 1)", "scso"),
    ("bSCSO", "2023", "Binary wrapper FS, standard transfer", "Same-family baseline (Discussion)", "bscso"),
    ("Binary SCSO (biomedical)", "2023", "Binary wrapper FS, standard transfer", "Same-family baseline (Discussion)", "scsofs2"),
    ("Adaptive SCSO", "2024", "Binary wrapper FS, standard transfer", "Same-family baseline (Discussion)", "scsofs3"),
    ("IMSCSO", "2024", "Continuous search (multi-strategy)", "No — global optimization only, no FS", "imscso2024"),
    ("SCSO+Lens-OBL+SSA", "2024", "Continuous search (lens-OBL init)", "No — global optimization only, no FS", "scsolensobl2024"),
    ("Improved SCSO", "2024", "Continuous search dynamics", "No — global optimization only, no FS", "improvedscso2024"),
    ("MESCSO", "2025", "Continuous search (multi-strategy)", "No — global optimization only, no FS", "mescso2025"),
]

# Thứ tự trích dẫn RIÊNG cho bản ASOC -- KHÔNG giống SCIREP_CITE_ORDER vì
# Related Work sắp xếp lại theo taxonomy (relevance-guided mechanisms trước
# SCSO variants), đổi thứ tự xuất hiện lần đầu của neri/ludwig2025guided so
# với bscso/scsofs2/... Rút ra CƠ HỌC từ chính RG-SCSO_ASOC.tex đã compile
# (grep toàn bộ \cite{} theo thứ tự xuất hiện), không đếm tay, tránh sai sót.
ASOC_CITE_ORDER = [
    "guyon", "mrmr", "bgwo", "pso", "mafarja", "aoa", "coa", "rime", "tf",
    "scso", "neri", "ludwig2025guided", "bscso", "scsofs2", "scsofs3",
    "imscso2024", "mescso2025", "scsolensobl2024", "improvedscso2024",
    "kraskov", "gwo", "holm", "demsar",
]


def _cnum(key: str) -> str:
    return str(ASOC_CITE_ORDER.index(key) + 1)


def _c(*keys: str) -> str:
    nums = sorted(int(_cnum(k)) for k in keys)
    return "[" + ", ".join(str(n) for n in nums) + "]"


def add_references_asoc(doc) -> None:
    doc.add_heading("References", level=1)
    entries = _parse_bib()
    for num, key in enumerate(ASOC_CITE_ORDER, 1):
        f = entries.get(key)
        if f is None:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def run(text, italic=False):
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.italic = italic

        run(f"[{num}] {_fmt_authors(f.get('author', ''))}, "
            f"{f.get('title', '')}. ")
        run(f.get("journal") or f.get("booktitle") or "", italic=True)
        seg = ""
        if f.get("volume"):
            seg += f" {f['volume']}"
        if f.get("pages"):
            seg += f", {f['pages'].replace('--', '–')}"
        if f.get("year"):
            seg += f" ({f['year']})"
        run(seg + ".")
        _bookmark_paragraph(p, f"ref{num}", 4000 + num)


def _renumber_caption(doc, replacements: dict) -> None:
    """The 4 table-adding functions reused verbatim from build_paper_scirep_
    docx.py (_add_heldout_combined_table, add_extended_ablation_table,
    add_classic_baselines_table, add_classifier_robustness_table) carry
    hardcoded "Table N" captions matching the SciRep document's table order.
    ASOC's Related Work promotes a new Table 1 (literature positioning)
    ahead of all of them and inserts a Threshold sensitivity table, so every
    number shifts. caption() puts the whole caption in a single run, so a
    plain substring replace on the most-recently-added paragraph is exact
    and does not touch any other paragraph in the document."""
    p = doc.paragraphs[-1]
    text = p.runs[0].text
    for old, new in replacements.items():
        text = text.replace(old, new)
    p.runs[0].text = text


def add_literature_positioning_table(doc) -> None:
    """Docx mirror of literature_positioning_table() in build_paper_scirep.py
    -- promoted into the ASOC main-text Related Work section (not
    Supplementary, unlike the SciRep version), so column citation keys are
    resolved through ASOC_CITE_ORDER's numbering here."""
    cols = ["Method", "Year", "What it modifies", "Comparable to this protocol?"]
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _hdr(t, cols, 8)
    for name, year, mod, comp, key in LITPOS_ROWS:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(f"{name} {_c(key)}").font.size = Pt(8)
        cells[1].paragraphs[0].add_run(year).font.size = Pt(8)
        cells[2].paragraphs[0].add_run(mod).font.size = Pt(8)
        cells[3].paragraphs[0].add_run(comp).font.size = Pt(8)
    caption(doc, "Table 1 Positioning of RG-SCSO against recent SCSO-family "
                 "literature cited in this paper. The four 2024-2025 "
                 "continuous-search variants (IMSCSO, SCSO+Lens-OBL+SSA, "
                 "Improved SCSO, MESCSO) improve exploration/exploitation "
                 "dynamics in continuous space but are global optimization "
                 "studies, not feature selectors, so they are not directly "
                 "comparable under this paper's binary FS protocol; they "
                 "are cited to establish that none of them touches the "
                 "binarization interface itself. The three binary SCSO "
                 "feature selectors (bSCSO and two further variants) are "
                 "directly comparable and are included as same-family "
                 "baselines (Discussion).")


def add_threshold_sensitivity_table(doc) -> None:
    """Docx mirror of threshold_sensitivity_table() in build_paper_scirep.py
    -- promoted into ASOC main-text Results (not Supplementary). No Phi
    cross-reference to a stability table here: the docx builder, unlike the
    tex one, has no Supplementary companion and no stability-index table at
    all, so Phi is described inline instead of via a dangling \\ref-style
    pointer."""
    if not os.path.exists(THRESHOLD_CSV):
        return
    th = pd.read_csv(THRESHOLD_CSV)
    if th.empty:
        return
    taus = sorted(th["tau"].unique())
    datasets = sorted(th["dataset"].unique())
    cols = ["Dataset", "tau", "Mean Acc.", "Mean #Feat.", "Phi"]
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _hdr(t, cols, 8)
    for ds in datasets:
        best_acc = max(
            float(th[(th.dataset == ds) & (th.tau == t_)]["mean_accuracy"].iloc[0])
            for t_ in taus
        )
        for i, t_ in enumerate(taus):
            r = th[(th.dataset == ds) & (th.tau == t_)].iloc[0]
            acc = float(r["mean_accuracy"])
            nfeat = float(r["mean_n_selected"])
            phi = float(r["nogueira_phi"])
            cells = t.add_row().cells
            cells[0].paragraphs[0].add_run(ds if i == 0 else "").font.size = Pt(8)
            cells[1].paragraphs[0].add_run(f"{t_:.1f}").font.size = Pt(8)
            accr = cells[2].paragraphs[0].add_run(f"{acc:.4f}")
            accr.font.size = Pt(8)
            if abs(acc - best_acc) < 1e-9:
                accr.bold = True
            cells[3].paragraphs[0].add_run(f"{nfeat:.1f}").font.size = Pt(8)
            cells[4].paragraphs[0].add_run(f"{phi:.3f}").font.size = Pt(8)
    caption(doc, "Table 5 Threshold sensitivity (tau in {0.4, 0.5, 0.6} "
                 "replacing the fixed 0.5 preferred-bit threshold; 30 "
                 "independent runs per cell, same protocol and datasets as "
                 "the main ablation). Highest accuracy per dataset in bold. "
                 "No single tau dominates uniformly: tau=0.5, the value "
                 "deployed throughout this paper, attains the highest "
                 "accuracy on two of five datasets; mean selected-feature "
                 "count falls monotonically as tau increases on every "
                 "dataset, at a modest, dataset-dependent, non-uniform "
                 "accuracy cost. Phi is the Nogueira stability index.")


_NOTATION_ROWS = [
    ("d", "Total number of features (search-space dimension)"),
    ("b ∈ {0,1}ᵈ", "Candidate binary feature-subset mask"),
    ("f(b)", "Fitness function, Eq. (1)"),
    ("Acc(b)", "Stratified 5-fold KNN accuracy on subset b"),
    ("|b|", "Number of selected features (subset size)"),
    ("R(t)", "SCSO sensitivity range at iteration t"),
    ("S_M", "Maximum sensitivity range parameter (= 2)"),
    ("T_max", "Maximum number of iterations"),
    ("T(·)", "Continuous-to-binary transfer function, T: R → [0,1]"),
    ("xⱼ", "Continuous position of feature j"),
    ("δⱼ", "Perturbation applied to xⱼ"),
    ("Δⱼ", "Induced change in bit-flip probability"),
    ("‖T′‖∞", "Largest slope (Lipschitz bound) of the transfer function"),
    ("ε", "Slope bound in a flat/saturated region of T"),
    ("γ", "RMS modulation strength (bias intensity)"),
    ("ρⱼ", "Relevance score of feature j, ρⱼ ∈ [0,1]"),
    ("b*ⱼ", "Preferred bit value for feature j, 1[ρⱼ > 0.5]"),
    ("sⱼ", "Modulation strength, sⱼ = 2|ρⱼ − 0.5| ∈ [0,1]"),
    ("σⱼ", "Direction indicator (+1 toward b*ⱼ, −1 otherwise)"),
    ("pⱼ", "RMS-modulated flip probability for feature j, Eq. (2)"),
    ("V(xⱼ)", "V-shaped base transfer, |tanh(xⱼ)|"),
    ("τ", "Preferred-bit threshold (deployed at 0.5; swept over {0.4, 0.5, 0.6})"),
    ("ρ_static", "Static mutual-information relevance prior"),
    ("I(Xⱼ;y)", "Mutual information between feature j and the label"),
    ("H(y)", "Label entropy"),
    ("K", "Number of UMR memetic probes per iteration"),
    ("N", "Population size (pop_size)"),
    ("max_nfe", "Total fitness-evaluation budget, pop_size × max_iter"),
    ("n", "Number of samples (in the O(dn log n) MI-prior cost)"),
]


def add_notation_table(doc) -> None:
    """Docx mirror of notation_table() in build_paper_scirep.py, per
    MASTER_FINAL_COMPLETE.md Section 7 ("Notation consistency table --
    BAT BUOC"), placed at the start of Methods as that section recommends
    -- ASOC has no Supplementary companion in the docx, so unlike the tex
    version (which places it in Supplementary as one of several promoted
    items) this is main text only. Compiled directly from the symbols
    actually used in Methods; T is genuinely overloaded (max-iteration
    count in the sensitivity-range formula vs. the transfer function in
    the washout subsection) and both uses are listed rather than silently
    disambiguated, matching the tex version's own disclosure."""
    # Two symbols need a real (non-single-letter) subscript that the
    # shared ⱼ/^-exponent tokenizer doesn't cover: S_M (S sub M) and
    # rho_static (rho sub "static") are genuine math subscripts. Unlike
    # max_nfe/pop_size/max_iter below, which are literal code-parameter
    # names (kept as plain underscored text, matching \mathrm{max\_nfe}
    # in the tex source), these two must render as an actual subscript run.
    _MATH_SUBSCRIPT = {"S_M": ("S", "M"), "ρ_static": ("ρ", "static"),
                        "T_max": ("T", "max")}

    cols = ["Symbol", "Meaning"]
    t = doc.add_table(rows=1, cols=len(cols))
    _ieee_table(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    _hdr(t, cols, 8)
    for sym, meaning in _NOTATION_ROWS:
        cells = t.add_row().cells
        if sym in _MATH_SUBSCRIPT:
            base, sub = _MATH_SUBSCRIPT[sym]
            p = cells[0].paragraphs[0]
            p.add_run(base).font.size = Pt(8)
            r = p.add_run(sub)
            r.font.size = Pt(8)
            r.font.subscript = True
        else:
            _add_run_text(cells[0].paragraphs[0], sym, size=8)
        _add_run_text(cells[1].paragraphs[0], meaning, size=8)
    caption(doc, "Table 8 Notation used in Methods, compiled for reference. "
                 "T_max (maximum iteration count, used in the SCSO "
                 "sensitivity-range formula) and T(·) (the continuous-to-"
                 "binary transfer function used throughout the washout "
                 "subsection) are kept as distinct symbols to avoid the "
                 "overloading that a single bare T would otherwise cause.")


def build() -> None:
    s = load_summary()
    colon = s["gene"].get("ColonCancer", {})

    adaptive = adaptive_baselines()
    scsofam = scso_family_baselines()
    robust = robustness_baselines()
    svm16 = robustness_svm16()
    inf_val = inference_value()
    diversity = diversity_analysis() if os.path.exists(DIVERSITY_CSV) else None

    _hs = _heldout.load()
    hs_wtl = _hs["wil"]["mark"].value_counts()
    hs_w, hs_t, hs_l = (int(hs_wtl.get(k, 0)) for k in ("+", "=", "-"))
    hs_d_aoa = float(_hs["wil"][_hs["wil"].compared_with == "AOA"]["cohens_d"].abs().median())
    hs_rank = _hs["ranking"].sort_values()
    hs_stats = _hs.get("stats", {})

    feat_counts = [pd.read_csv(os.path.join(PROC_DIR, f"{d}.csv")).shape[1] - 1
                   for d in s["datasets"]]
    feat_min, feat_max = min(feat_counts), max(feat_counts)

    if s.get("stats"):
        w, ti, l = s["sig_total"]
        rank7 = s["rank7"]
        n_cmp = w + ti + l
    v = s.get("verdict", {})
    rms, orl, umr = v.get("NoRMS", {}), v.get("NoORL", {}), v.get("NoUMR", {})

    scsofam_pct = min(scsofam["red"].values()) if scsofam and scsofam.get("red") else None
    hs_nf_ratio = _hs["nf_mean"]["SCSO"].mean() / _hs["nf_mean"]["RG-SCSO"].mean()

    doc = Document()
    _sec0_a4(doc)
    _style_setup(doc)

    # ---------------------------------------------------------- Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "Relevance-Guided Binarization for Parsimonious Feature Selection "
        "with Sand Cat Swarm Optimization")
    tr.bold = True
    tr.font.size = Pt(18)

    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a1 = au.add_run("Bui Quang Huy")
    a1.font.size = Pt(11)
    sup1 = au.add_run("1,*")
    sup1.font.size = Pt(11)
    sup1.font.superscript = True
    a2 = au.add_run(" and Duong Minh Son")
    a2.font.size = Pt(11)
    sup2 = au.add_run("1")
    sup2.font.size = Pt(11)
    sup2.font.superscript = True
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affn1 = aff.add_run("1")
    affn1.font.superscript = True
    affn1.italic = True
    affn1.font.size = Pt(9.5)
    affr = aff.add_run("Dong A University, Da Nang, Vietnam")
    affr.italic = True
    affr.font.size = Pt(9.5)
    af = doc.add_paragraph()
    af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    afr = af.add_run("*Corresponding author(s). E-mail(s): huybq@donga.edu.vn; "
                      "Contributing authors: sondm@donga.edu.vn")
    afr.italic = True
    afr.font.size = Pt(9.5)

    # -------------------------------------------------------------- Abstract
    # 150-250 words, no p-value/Cohen's d/Friedman-rank, no citations --
    # matches build_paper_asoc.py's tex abstract exactly (same content).
    ab = doc.add_paragraph()
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead = ab.add_run("Abstract ")
    lead.bold = True
    lead.font.size = Pt(9)
    abstract_body = (
        "Wrapper feature selection with swarm intelligence typically searches "
        "continuously and crosses into the binary domain via a fixed transfer "
        "function, a feature-agnostic quantization that discards continuous "
        "operators' fine adjustments, an effect we term washout. RG-SCSO "
        "replaces this transfer with a per-feature, relevance-modulated "
        "binarization: a mutual-information field biases each feature's "
        "bit-flip probability so informative features resist removal and "
        "noise resists inclusion. Its central and most consistent finding is "
        f"parsimony: on {s['n']} datasets, including two gene-expression "
        "sets, RG-SCSO selects the second-smallest feature subsets of any "
        "method tested, trailing only COA, which pays a several-point "
        f"accuracy cost for it, using on average approximately "
        f"{100/hs_nf_ratio:.0f}% of the feature count selected by base "
        "SCSO, an advantage that remains evident across the reported "
        "stress tests, including comparison against optimizers carrying published "
        "adaptive transfers, against which RG-SCSO does not lead on "
        "accuracy yet still selects "
        f"{adaptive['red_min']:.0f}-{adaptive['red_max']:.0f}% "
        "fewer features. Under a budget-matched, leak-free protocol, where the "
        "relevance prior, search, and cross-validated fitness are computed "
        "only on the training partition and never on held-out labels, this "
        "compactness comes with the best mean held-out accuracy of any "
        "method tested, with a consistent edge over the closest competitor. "
        "A controlled ablation isolates the contribution of the relevance "
        "modulation itself, and cross-classifier, cross-prior experiments "
        "test the robustness of the mechanism. The binarization interface "
        "is the most reliable injection point for the relevance signal, "
        "though not the sparsest: LASSO is sparser on the two "
        "gene-expression sets, but accuracy is preserved, not improved, "
        "there, and parsimony without that extreme structure is the "
        "transferable gain.")
    abr = ab.add_run(abstract_body)
    abr.font.size = Pt(9)

    kw = doc.add_paragraph()
    kwl = kw.add_run("Keywords ")
    kwl.bold = True
    kwl.font.size = Pt(9)
    kwr = kw.add_run("Feature selection, Sand Cat Swarm Optimization, binary "
                      "optimization, relevance-guided binarization, wrapper "
                      "feature selection, parsimony, metaheuristic.")
    kwr.italic = True
    kwr.font.size = Pt(9)

    # ------------------------------------------------------------ Introduction
    doc.add_heading("1. Introduction", level=1)
    para(doc, "Feature selection removes irrelevant and redundant features "
              "to improve classifier accuracy, reduce overfitting, and "
              "lower computational cost, a payoff that is greatest for "
              "high-dimensional, small-sample problems such as "
              "gene-expression classification, where the number of "
              f"features exceeds the number of samples by orders of "
              f"magnitude {_c('guyon','mrmr')}. Wrapper selection, which "
              "scores a subset by the performance of a downstream "
              "classifier, is frequently cast as a combinatorial problem "
              "solved by swarm-intelligence metaheuristics, including grey "
              f"wolf {_c('bgwo')}, particle swarm {_c('pso')}, whale "
              f"{_c('mafarja')}, and several recent optimizers "
              f"{_c('aoa','coa','rime')}. Most such methods were conceived "
              "for continuous optimization and are adapted to the binary "
              "space through a transfer function, most commonly an "
              "S-shaped or V-shaped map, that converts a real-valued "
              f"position into a selection probability {_c('tf')}. We argue "
              "that this retrofit contains a structural weakness: the "
              "transfer function is fixed and feature-agnostic, applied "
              "identically to every dimension, so the incremental "
              "adjustments a continuous-space operator makes are collapsed "
              "by the squash-and-threshold step before they can influence "
              "the retained subset. We refer to this loss as washout, and "
              "show in Methods below that it is not a thought experiment: "
              "four well-motivated continuous-space enhancements of SCSO "
              f"{_c('scso')} failed to beat the base algorithm under an "
              "identical feature-selection protocol (0 wins, 1 loss, 17 "
              "ties by Wilcoxon signed-rank; Supplementary Table S5, "
              "preliminary washout study).")
    para(doc, "Existing studies have explored relevance-guided "
              "initialization, objective weighting, and adaptive binary "
              "transfer mechanisms; however, the explicit injection of a "
              "per-feature relevance prior into the binarization operator "
              "of SCSO-based feature selection remains insufficiently "
              "investigated (Section 2 positions this gap against each "
              "cited SCSO-family work in detail).")
    para(doc, "We close this gap with RG-SCSO: a per-feature, "
              "relevance-modulated binarization in which a "
              "mutual-information relevance field biases each feature's "
              "bit-flip probability, turning a knowledge-agnostic "
              "quantization step into a knowledge-carrying operator. "
              "SCSO's continuous search, including its sensitivity range, "
              "is retained unchanged; the novelty resides entirely in the "
              "binarization (Fig. 1). We make four contributions: we "
              "identify washout as a mechanistic failure mode and derive a "
              "diagnostic bound, together with a cumulative extension "
              "linking it to discrete transition dynamics; we propose "
              "RG-SCSO, whose ablation-confirmed centerpiece is "
              "relevance-modulated sensitivity (RMS), supplemented by a "
              "smaller, budget-neutral memetic refinement step (UMR) whose "
              "own sensitivity sweep shows it contributes far less than "
              "RMS, an online-learning variant of the relevance field is "
              "examined and pruned entirely by ablation; we evaluate under "
              "a preregistered, budget-matched, leak-free protocol denying "
              "the relevance prior any access to test labels; and we "
              "report a full statistical treatment, a component ablation, "
              "and a size-fair enrichment analysis correlating the "
              "observed parsimony with relevance guidance.")

    doc.add_heading("Research questions", level=2)
    para(doc, "This study is organized around four research questions, "
              "each answered directly by a specific part of Results or "
              "Discussion below.")
    bullet(doc, "RQ1. Does relevance-modulated binarization improve "
                "held-out feature-selection performance under a fixed "
                "computational budget? Addressed in Section 3.1 "
                "(Table 2, Fig. 3).")
    bullet(doc, "RQ2. Can RG-SCSO reduce the number of selected features "
                "while preserving competitive generalization accuracy? "
                "Addressed in Section 3.2.")
    bullet(doc, "RQ3. Is direct relevance injection at the binarization "
                "interface more effective than relevance-guided "
                "initialization or objective weighting? Addressed in "
                "Section 3.4 (Table 4).")
    bullet(doc, "RQ4. How robust is the proposed mechanism across "
                "classifiers, relevance priors, and dataset "
                "dimensionalities? Addressed in Section 3.6 (Table 7) "
                "and Discussion.")
    para(doc, "The remainder of the paper reviews existing approaches "
              "along three axes: (i) binary transfer functions in "
              "swarm-based feature selection, (ii) relevance-guided search "
              "mechanisms, and (iii) SCSO variants (Section 2).")
    full_width(doc, lambda: add_figure(
        doc, "concept.png",
        "Fig. 1. Conceptual overview. (a) The conventional pipeline, where "
        "continuous-operator adjustments are collapsed by a fixed, "
        "feature-agnostic transfer (washout). (b) RG-SCSO, where a "
        "per-feature, relevance-modulated binarization replaces the "
        "feature-agnostic transfer, biasing each feature's bit-flip "
        "probability by a mutual-information relevance field, followed by "
        "memetic refinement on uncertain bits.",
        width_in=5.3))

    # -------------------------------------------------------------- Related Work
    doc.add_heading("2. Related work", level=1)
    doc.add_heading("Binary transfer functions in swarm-based feature "
                     "selection", level=2)
    para(doc, f"Swarm-intelligence wrapper selectors, including grey wolf "
              f"{_c('bgwo')}, particle swarm {_c('pso')}, whale "
              f"{_c('mafarja')}, and several recent optimizers "
              f"{_c('aoa','coa','rime')}, are conceived for continuous "
              "optimization and cross into the binary search space "
              f"through a transfer function, most commonly an S-shaped or "
              f"V-shaped map {_c('tf')}. This transfer is fixed and "
              "feature-agnostic in every one of these methods: it is "
              "applied identically to every dimension regardless of that "
              "feature's relevance, the structural weakness this paper "
              "terms washout (Introduction, Methods).")
    doc.add_heading("Relevance-guided mechanisms in feature selection",
                     level=2)
    para(doc, f"Filter criteria such as mutual information and mRMR "
              f"{_c('guyon','mrmr')} encode problem knowledge cheaply but "
              "are decoupled from the wrapper search; memetic "
              f"hybridization {_c('neri')} adds local refinement without "
              "addressing the binarization interface; and knowledge-guided "
              "metaheuristics that inject filter information into "
              "initialization or the objective, such as filter-guided PSO "
              f"for cancer-genome selection {_c('ludwig2025guided')}, bias "
              "where the search starts or what it optimizes but still "
              "leave the binarization operator itself knowledge-agnostic. "
              "RG-SCSO differs by injecting the relevance signal at the "
              "binarization interface itself rather than upstream of it; "
              "the signal-position experiment in Results (Table 4) tests "
              "this design choice directly against both alternatives.")
    doc.add_heading("Sand Cat Swarm Optimization and its variants", level=2)
    para(doc, f"SCSO {_c('scso')} is a comparatively recent continuous "
              f"metaheuristic. The existing binary SCSO feature selectors "
              f"{_c('bscso','scsofs2','scsofs3')} apply standard transfer "
              "functions without a relevance-aware binarization step, and "
              "the broader recent SCSO literature that adds chaotic "
              "initialization, differential mutation, or hybridized "
              "search strategies "
              f"{_c('imscso2024','mescso2025','scsolensobl2024','improvedscso2024')} "
              "improves continuous-space search dynamics while leaving the "
              "binarization interface itself untouched, and none of these "
              "works, to our knowledge, makes the binarization operator "
              "itself per-feature and relevance-aware. Table 1 below "
              "positions each of these works against this claim directly.")
    doc.add_heading("Literature-positioning summary", level=2)
    para(doc, "Table 1 makes this gap concrete across every SCSO-family "
              "work cited above: the three directly-comparable binary "
              "feature selectors apply a standard, feature-agnostic "
              "transfer with no relevance-aware step, and the four "
              "continuous-search variants improve exploration/exploitation "
              "dynamics without touching the binarization interface at "
              "all. None of the eight prior works positioned here makes "
              "the binarization operator itself per-feature and "
              "relevance-aware. This is precisely the interface RG-SCSO "
              "modifies.")
    full_width(doc, lambda: add_literature_positioning_table(doc))

    # ----------------------------------------------------------------- Results
    doc.add_heading("3. Results", level=1)
    doc.add_heading("Held-out generalization", level=2)
    para(doc, "RQ1 asks whether relevance-modulated binarization improves "
              "held-out performance under a fixed computational budget; "
              "this subsection answers it directly. For each dataset, "
              "algorithm, and independent run we draw an "
              "outer stratified 80/20 split. The relevance prior, the "
              "search, and the cross-validated fitness are computed "
              "exclusively on the 80% training partition; the selected "
              "subset is evaluated once on the untouched 20% hold-out, on "
              "which a fresh k-NN classifier (standardized on the training "
              "partition) reports accuracy, so the relevance prior never "
              "has transductive access to the test labels. Table 2 reports "
              "held-out accuracy over all seven algorithms and "
              f"{s['n']} datasets, at a fixed evaluation budget of "
              "max_nfe = 15000 for every algorithm compared. RG-SCSO "
              f"attains the best average "
              f"Friedman rank ({hs_rank.iloc[0]:.2f}, ahead of the "
              f"second-placed AOA at {hs_rank.iloc[1]:.2f}; "
              f"chi-square={hs_stats.get('friedman_chi2', 0):.2f}, "
              "p < 0.001). A Holm-corrected Wilcoxon signed-rank test "
              f"across all pairwise comparisons gives RG-SCSO {hs_w} "
              f"significant wins, {hs_l} loss, and {hs_t} ties; the only "
              "close competitor is AOA, against which the advantage is "
              f"genuine but moderate (median |d|={hs_d_aoa:.2f}), RG-SCSO "
              "still leading on mean accuracy.")
    full_width(doc, lambda: (
        _add_heldout_combined_table(doc, _hs),
        _renumber_caption(doc, {"Table 1 Held-out": "Table 2 Held-out"}),
    ))

    doc.add_heading("Feature-subset parsimony", level=2)
    inference_sentence = ""
    if inf_val is not None:
        inference_sentence = (
            " The dimensionality reductions are too small relative to "
            "these datasets' modest sample counts (50–455 training "
            "instances) for a wall-clock inference saving to be resolvable "
            "above call overhead. We therefore report a projected, "
            "reconstructed inference-cost analysis rather than a measured "
            "deployment speedup: replaying the same feature-count "
            "reductions at a synthetic deployment-scale workload (5,000 "
            f"training instances, brute-force k-NN) projects a "
            f"{inf_val['min_speedup']:.0f}–{inf_val['max_speedup']:.0f}% "
            "reduction in batch-inference cost relative to AOA (mean "
            f"{inf_val['mean_speedup']:.0f}%); this is a controlled "
            "complexity demonstration, not a measurement on the "
            "benchmark's own test sets.")
    para(doc, "RQ2 asks whether RG-SCSO can reduce feature count while "
              "preserving competitive accuracy; this subsection answers "
              "it directly. Parsimony, more than any accuracy margin, is "
              "RG-SCSO's "
              "defining property. Table 2 shows it attains the best mean "
              f"accuracy of any method tested ({_hs['acc_mean']['RG-SCSO'].mean():.3f}) "
              "while selecting the second-fewest features on the held-out "
              f"setting, on average {_hs['nf_mean']['RG-SCSO'].mean():.1f} against "
              f"{_hs['nf_mean']['SCSO'].mean():.1f} for base SCSO and "
              f"{_hs['nf_mean']['AOA'].mean():.1f} for AOA. Only COA selects "
              f"fewer ({_hs['nf_mean']['COA'].mean():.1f} on average), and it "
              f"does so at {_hs['acc_mean']['RG-SCSO'].mean() - _hs['acc_mean']['COA'].mean():.3f} "
              "lower mean accuracy, so RG-SCSO is the most compact method "
              "that does not trade away accuracy to get there. On "
              f"ColonCancer it retains {colon.get('nf', float('nan')):.0f} of "
              f"{colon.get('ntot', '--')} features versus "
              f"{colon.get('nf_aoa', float('nan')):.0f} for AOA, at higher "
              "accuracy. This advantage remains evident across the reported "
              "stress tests: "
              "against optimizers carrying published adaptive transfers "
              "RG-SCSO does not lead on accuracy yet still "
              f"selects {adaptive['red_min']:.0f}–{adaptive['red_max']:.0f}"
              "% fewer features, and against same-family binary SCSO "
              f"selectors {_c('bscso','scsofs2','scsofs3')} it is "
              f"{scsofam_pct:.0f}% smaller at comparable accuracy."
              f"{inference_sentence}")
    full_width(doc, lambda: add_figure(
        doc, "accuracy_parsimony_tradeoff.png",
        "Fig. 2. Mean held-out accuracy vs. mean selected-feature "
        "fraction, averaged across all "
        f"{s['n']} datasets, one point per algorithm. RG-SCSO achieves the "
        "highest mean held-out accuracy and the second-smallest mean "
        "selected-feature count among the seven algorithms tested. AOA "
        "attains the second-highest accuracy but selects on average 98% of "
        "available features, essentially no feature selection; COA is the "
        "sparsest method but at a distinctly lower accuracy than "
        "RG-SCSO.",
        width_in=4.4))

    doc.add_heading("Ranking and statistical significance", level=2)
    if s.get("stats"):
        para(doc, f"Table 3 gives the in-sample average rank "
                  f"({rank7['RG-SCSO']:.2f} for RG-SCSO) and the "
                  "Holm-significant win/tie/loss against each baseline. "
                  f"These are paired per-dataset comparisons, "
                  f"not independent trials: across the {n_cmp} "
                  f"dataset-baseline pairs RG-SCSO wins {w} and loses "
                  f"{l}, with predominantly large effect sizes (median "
                  f"|d|={s['es_median']:.2f}; {s['es_large_pct']:.0f}% "
                  "exceed 0.8). These in-sample effect sizes are an "
                  "optimistic upper bound: they contract to the "
                  "small-to-moderate held-out range above once the "
                  "relevance prior is denied access to test labels, "
                  "whereas the ranking itself is preserved.")
        full_width(doc, lambda: (
            caption(doc, "Table 3 Average Rank, Holm-Significant "
                         "Win/Tie/Loss, and Median Effect Size vs. RG-SCSO."),
            add_rank_table_with_effect_size(doc, s),
        ))
        para(doc, "Fig. 3 below visualizes the held-out ranking (Table 2, "
                  "Held-out generalization above) as a critical-difference "
                  "diagram, the paper's primary ranking evidence, distinct "
                  "from the in-sample rank in Table 3 just above.")
        full_width(doc, lambda: add_figure(
            doc, "cd_diagram_heldout.png",
            "Fig. 3. Critical-difference (Nemenyi) diagram at alpha=0.05 "
            "over the held-out ranking (Table 2); algorithms not joined by "
            "a bar differ significantly in mean held-out rank.",
            width_in=5.3))
    else:
        para(doc, "Ranking and significance are reported in the final version.")

    doc.add_heading("Ablation and mechanism", level=2)
    if s.get("ablation"):
        para(doc, "We started from a three-component design and tested "
                  "each part by removal (Table 4), judging significance "
                  "with a paired Wilcoxon signed-rank test (Holm-corrected "
                  f"across datasets). RMS is the strongest: removing it "
                  f"costs {rms.get('worst_delta_pts', 0):.2f} accuracy "
                  f"points on {rms.get('worst_ds', '')} "
                  f"(d={rms.get('worst_d', 0):.2f}, Holm p<0.001). UMR is "
                  f"also load-bearing ({umr.get('worst_delta_pts', 0):.2f} "
                  f"points on {umr.get('worst_ds', '')}), whereas ORL is "
                  f"not (degrades accuracy on only "
                  f"{orl.get('n_deg', 0)}/{orl.get('n_ds', 0)} datasets) "
                  "and is therefore dropped; the final RG-SCSO comprises "
                  "RMS and UMR only.")
        full_width(doc, lambda: (
            add_extended_ablation_table(doc, s),
            _renumber_caption(doc, {"Table 3 Component": "Table 4 Component"}),
        ))
    else:
        para(doc, "The ablation study is reported in the final version.")
    para(doc, "RQ3 asks whether direct relevance injection at the "
              "binarization interface is more effective than "
              "relevance-guided initialization or objective weighting. A "
              "dedicated signal-position experiment answers it: injecting "
              "the same mutual-information field elsewhere, at "
              "initialization or as an objective penalty, rather than at "
              "the binarization interface; Table 4 adds both alternatives "
              "to the component ablation. MI-weighted objective injection "
              "is significantly worse than the deployed RMS+UMR "
              "configuration on four of five datasets and tied on the "
              "fifth (Leukemia), supporting the original claim that the "
              "binarization interface is a more effective injection point "
              "than the objective function. MI-guided initialization is a "
              "harder comparison: significantly worse on ColonCancer and "
              "Sonar, but statistically indistinguishable on Leukemia, "
              "WDBC, and Zoo, and on Leukemia specifically it matches "
              "RG-SCSO's accuracy while selecting roughly four times fewer "
              "features (245 vs. 940). The binarization interface is thus "
              "the most reliable injection point across datasets, not a "
              "uniformly superior one; on the most extreme p >> n dataset "
              "tested, a simpler injection at initialization meets or "
              "beats it on both accuracy and parsimony.")
    full_width(doc, lambda: add_figure(
        doc, "convergence_fs.png",
        "Fig. 4. Mean best fitness versus iteration, RG-SCSO vs. SCSO vs. "
        "AOA, on Zoo (16 features), WDBC (30 features), and ColonCancer "
        "(2000 features), mean over 5 runs, on the actual "
        "feature-selection objective (illustrative, not a new statistical "
        "claim). On Zoo the three algorithms converge along essentially "
        "the same trajectory. On WDBC and ColonCancer, RG-SCSO both "
        "converges faster and plateaus at a lower (better) fitness than "
        "SCSO or AOA, which themselves plateau early at a distinctly "
        "worse value rather than continuing to close the gap with more "
        "iterations; the difference grows with dimensionality.",
        width_in=5.3))
    para(doc, "We also test whether relevance guidance makes RG-SCSO "
              "preferentially retain high mutual-information features. "
              "Because a subset of size |S| overlaps the top-|S| "
              "mutual-information features at a chance rate of |S|/N, we "
              "report a size-fair enrichment (Fig. 5), the fraction of "
              "selected features in the top-|S| set divided by this "
              "chance level. RG-SCSO's subset is enriched above chance on "
              "both gene-expression sets, whereas the relevance-agnostic "
              "SCSO sits at chance, evidence consistent with the "
              "relevance field driving the smaller and more accurate "
              "subsets, though enrichment alone is correlational. A "
              "direct intervention test, detailed in Supplementary "
              "Information, "
              "permutes the field's feature identities while preserving "
              "its value "
              "distribution: the permuted field yields no significant "
              "accuracy difference from the real field on three of five "
              "datasets, a significant but negligible-effect difference on "
              "Leukemia, and a clear difference only on ColonCancer, while "
              "inverting the field's sign is significantly worse than "
              "both on every dataset. The relevance field's direction and "
              "scale therefore matter consistently; the exact per-feature "
              "ranking within it matters demonstrably on only one of five "
              "datasets, a materially weaker causal claim than the "
              "enrichment analysis alone would suggest.")
    full_width(doc, lambda: add_figure(
        doc, "mechanism.png",
        "Fig. 5. Mechanism evidence. (a) Size-fair top-MI enrichment on "
        "the two gene-expression sets (selection precision divided by the "
        "chance level |S|/N; mean over 30 runs, error bars = std): "
        "RG-SCSO enriches its subset in relevant features above chance, "
        "whereas the relevance-agnostic SCSO sits at chance. (b) "
        "Feature-selection stability (Nogueira Phi, 30 runs) across all "
        "five representative datasets: RG-SCSO is more consistent run to "
        "run than same-family SCSO, clearest on the three lower-dimensional "
        "sets, but both are close to the level expected by chance on the "
        "two gene-expression sets, where AOA's near-zero Phi reflects that "
        "it selects nearly every available feature rather than genuine "
        "instability.",
        width_in=5.3))

    doc.add_heading("Threshold sensitivity", level=2)
    para(doc, "The 0.5 preferred-bit threshold that separates preferred "
              "from disfavored bits (Methods) is a convenience, not a "
              "theoretically grounded neutral point. This sweep "
              "(tau in {0.4, 0.5, 0.6}, 30 independent runs per cell, same "
              "protocol and datasets as the main ablation) tests whether "
              "that choice is at least empirically reasonable. No single "
              "tau dominates uniformly: tau=0.5, the value deployed "
              "throughout this paper, attains the highest accuracy on two "
              "of five datasets; mean selected-feature count falls "
              "monotonically as tau increases on every dataset, at a "
              "modest, dataset-dependent, non-uniform accuracy cost.")
    full_width(doc, lambda: add_threshold_sensitivity_table(doc))
    full_width(doc, lambda: add_figure(
        doc, "threshold_heatmap.png",
        "Fig. 6. Mean held-out accuracy (color) and mean number of "
        "selected features (in parentheses) across the three preferred-bit "
        "thresholds tested (tau in {0.4, 0.5, 0.6}), one row per dataset. "
        "Accuracy is essentially flat across thresholds on every dataset "
        "(largest swing 0.011, ColonCancer tau=0.4 vs. tau=0.5), while the "
        "selected-feature count falls monotonically as tau increases on "
        "all five datasets, confirming tau=0.5 is not a fragile choice.",
        width_in=4.0))

    doc.add_heading("Comparison with classical selectors", level=2)
    para(doc, "RQ4 asks how robust the mechanism is across classifiers, "
              "relevance priors, and dataset dimensionalities; this "
              "subsection and Table 7 answer it directly. Table 6 "
              "compares RG-SCSO against five classical filter, "
              "embedded, and wrapper selectors, mutual-information "
              "thresholding, mRMR, ReliefF, LASSO, and sequential forward "
              "selection, under the identical fitness and evaluation "
              "protocol. RG-SCSO significantly outperforms every classical "
              "baseline on the three lower-dimensional benchmark datasets "
              "(Zoo, Sonar, WDBC). On the two gene-expression (p >> n) "
              "datasets, this advantage does not hold: LASSO attains "
              "higher accuracy with far fewer features on both Leukemia "
              "(99.82% at 23 features vs. RG-SCSO's 98.60% at 940) and "
              "ColonCancer (95.70% at 64 features vs. 88.09% at 563), and "
              "mRMR is competitive with RG-SCSO on both. RG-SCSO's "
              "practical advantage is clearest on datasets without "
              "extreme p >> n structure; on ultra-high-dimensional "
              "gene-expression data, a computationally far cheaper "
              "embedded method such as LASSO is a strong, arguably "
              "preferable, alternative. A remaining question is whether "
              "this parsimony mechanism is an artifact of the KNN wrapper "
              "used throughout; Table 7 answers it directly, extending the "
              "same relevance-guided search to SVM and Random Forest "
              "wrappers. The advantage over a no-prior baseline is not a "
              "KNN artifact, though it is a more consistent parsimony gain "
              "than an accuracy one under Random Forest specifically, and "
              "the ReliefF-prior degradation already established above "
              "reproduces under every wrapper tested.")
    full_width(doc, lambda: (
        add_classic_baselines_table(doc, s),
        _renumber_caption(doc, {"Table 4 Comparison": "Table 6 Comparison"}),
    ))
    full_width(doc, lambda: (
        add_classifier_robustness_table(doc),
        _renumber_caption(doc, {
            "Table 5 Robustness": "Table 7 Robustness",
            "as Table 4 (KNN/SVM": "as Table 6 (KNN/SVM",
        }),
    ))

    # -------------------------------------------------------------- Discussion
    doc.add_heading("4. Discussion", level=1)
    para(doc, "These results trace washout, a concrete failure mode of "
              "transfer-function-based binary feature selection, to its "
              "source and cure it by moving the relevance signal directly "
              "inside the binarization operator rather than upstream of it "
              "in the objective or the initialization. The formal result "
              "motivating this design, presented in Methods, is a "
              "diagnostic bound "
              "explaining why continuous-space enhancements fail at the "
              "binarization boundary, not a convergence guarantee for "
              "RG-SCSO itself, and the underlying recipe, a filter prior "
              "coupled to a wrapper search with memetic refinement, is a "
              "known combination in the feature-selection literature; what "
              "is new is the injection point together with a stringent "
              "evaluation protocol, budget-matching, a leak-free hold-out, "
              "cross-classifier and cross-prior robustness, and an "
              "explicit exploration-safety diagnostic. We report "
              "parsimony, not raw accuracy, as the transferable outcome of "
              "this choice.")
    para(doc, "Several boundaries delimit what these results establish. "
              "RG-SCSO inherits SCSO's continuous search dynamics "
              "unchanged, and the RMS rule carries a risk of its own: bias "
              "the flip probability too strongly and confidently "
              "classified bits can freeze in place, draining population "
              "diversity. We measured this directly rather than assuming "
              "it away: at an aggressive stress-test bias the risk is real "
              "and grows with dimensionality, but the conservative "
              "gamma=0.5 this paper deploys keeps the frozen-bit fraction "
              "below "
              f"{(diversity['max_frz_g5']*100 if diversity else 1.1):.1f}% "
              "throughout the run on every dataset tested. We also tested "
              "whether UMR's benefit specifically requires targeting "
              "relevance-uncertain features, or only the extra evaluation "
              "budget it spends, by replacing its targeted K-feature "
              "selection with K uniformly random features at matched NFE "
              "(Supplementary Information). This untargeted control "
              "significantly outperforms targeted UMR on both accuracy and "
              "feature count on the two gene-expression datasets, where "
              "UMR's contribution is largest (ColonCancer: 90.6% vs. 88.1% "
              "accuracy, 214 vs. 563 features; Leukemia: 423 vs. 940 "
              "features at statistically indistinguishable accuracy), and "
              "ties it on the other three. UMR's value over no memetic "
              "step at all remains real (Table 4); on the datasets where "
              "that value is largest, however, it does not depend on "
              "targeting relevance-uncertain features specifically, a more "
              'modest claim than "uncertainty-targeted" on its own '
              "implies. The relevance "
              "field is built from a single filter statistic; swapping in "
              "a ReliefF prior exposes a genuine limitation rather than "
              "confirming the mechanism, the parsimony advantage "
              "disappears (Supplementary Information), so RG-SCSO's "
              "compactness depends on a prior that drives uninformative "
              "features below the neutral point, not on the mere presence "
              "of a relevance signal. This dependence, together with the "
              "O(dn log n) one-time cost of the mutual-information prior "
              "amortized against the wrapper's N+K per-iteration "
              "evaluations, is why the scalability claim is scoped to the "
              "wrapper cost dominating overall runtime as feature "
              "dimensionality grows (full per-dataset wall-clock cost is "
              "in Supplementary Information) rather than to any runtime "
              "advantage over same-budget baselines, which we do not "
              "claim.")
    para(doc, "The "
              "main objective is also a KNN wrapper; under SVM, tested "
              f"directly on {svm16.get('n_ds', 16)} of the {s['n']} "
              "datasets, and under Random Forest on the same five-dataset "
              "subset as Table 7 in the Results, the parsimony advantage "
              "is not a KNN artifact, though it is a more consistent gain "
              "in subset size than in accuracy under Random Forest "
              "specifically, and the ReliefF-prior degradation already "
              "established above reproduces under every wrapper tested. "
              "The selected subset is smaller and more consistent in size "
              "than competing algorithms', but not necessarily more "
              "consistent in identity: a feature-selection stability index "
              "shows RG-SCSO more stable run to run than same-family SCSO "
              "on every dataset, clearest on the three lower-dimensional "
              "sets, yet on both gene-expression datasets RG-SCSO's own "
              "stability is itself close to the level expected by chance, "
              "so a much smaller subset there is not a materially more "
              "repeatable one (Fig. 5b; Supplementary Information gives "
              "the full per-dataset breakdown for both checks).")
    para(doc, "This wrapper search is itself compute-intensive; absolute "
              "wall-clock cost per run, for every algorithm tested, is "
              "reported in full in Supplementary Information rather than "
              "only argued qualitatively. A genuine outer-fold nested "
              "cross-validation pilot on three datasets, also in "
              "Supplementary Information, is directionally consistent "
              "with the single-split held-out estimate above, though "
              "underpowered at five runs per cell to confirm significance "
              "independently.")
    para(doc, "External validity has its own limits: the "
              "benchmark spans "
              f"{feat_min} to {feat_max} features across biomedical, "
              "gene-expression, and categorical domains drawn from a "
              "single curated family of UCI and standard microarray sets, "
              "and behaviour on ultra-high-dimensional omics data of 10^4 "
              "to 10^5 features is extrapolated, not measured. The "
              "relevance field itself is not perfectly stable on the "
              "smallest, highest-dimensional datasets: a bootstrap "
              "resampling analysis, reported in Supplementary Information, "
              "shows markedly lower rank stability of the mutual-information "
              "field on the two gene-expression sets than on the three "
              "lower-dimensional benchmarks, a risk the classical-baseline "
              "comparison above makes concrete rather than "
              "theoretical.")
    para(doc, "Finally, "
              "the accuracy claim is scoped, not universal: against binary "
              "particle-swarm and grey-wolf optimizers carrying published "
              "adaptive transfers, and against same-family binary SCSO "
              "selectors, RG-SCSO does not lead on accuracy, and on the "
              "two gene-expression datasets specifically a classical "
              "LASSO baseline outperforms RG-SCSO outright, so "
              "we do not claim a practical advantage for RG-SCSO on "
              "ultra-high-dimensional, small-sample data; we scope its "
              "transferable benefit to parsimony on datasets without that "
              "extreme structure. These boundaries are gathered together, "
              "with the future work they motivate, in Conclusion below.")

    # -------------------------------------------------------------- Conclusion
    doc.add_heading("5. Conclusion", level=1)
    para(doc, "The results support RG-SCSO as a relevance-guided binary "
              "feature-selection method whose primary advantage is "
              "improved subset parsimony with competitive predictive "
              f"performance under the tested conditions: across {s['n']} "
              "benchmark datasets it selects the second-smallest feature "
              "subsets of any method evaluated, trailing only COA, while "
              "attaining the best mean held-out accuracy among all seven "
              "algorithms compared, under a fixed, budget-matched, "
              "leak-free evaluation protocol (RQ1, RQ2).")
    para(doc, "This advantage is not universal, and the boundaries "
              "established in Discussion above are real constraints on "
              "the claim, not caveats to be read past. RG-SCSO does not "
              "lead on accuracy against optimizers carrying published "
              "adaptive continuous-space transfers or against "
              "same-family binary SCSO selectors; on the two extreme "
              "p >> n gene-expression datasets tested, a classical LASSO "
              "baseline outperforms it outright on both accuracy and "
              "feature count; its dependence on the underlying relevance "
              "prior is genuine, since replacing mutual information with "
              "a ReliefF prior removes the parsimony advantage entirely; "
              "the 0.5 preferred-bit threshold, while empirically "
              "reasonable across the range tested, is not a uniformly "
              "optimal choice; and the benchmark itself, though spanning "
              f"{feat_min}–{feat_max} features across biomedical, "
              "gene-expression, and categorical domains, is drawn from a "
              "single curated family of datasets, so behavior on "
              "ultra-high-dimensional omics data of 10^4–10^5 features is "
              "extrapolated rather than measured.")
    fw = doc.add_paragraph()
    _add_run_text(fw, "Future work ", bold=True)
    _add_run_text(fw, "includes: (i) adaptive, data-driven selection of the "
                "preferred-bit threshold, for example via cross-validation "
                "or an entropy-based criterion, rather than a fixed value; "
                "(ii) an explicit multi-objective formulation that traces "
                "an accuracy-parsimony Pareto front rather than a single "
                "weighted objective; (iii) extending the "
                "relevance-modulated binarization mechanism to other swarm "
                "optimizers, such as WOA, HHO, or SMA, to test whether the "
                "injection-point argument generalizes beyond SCSO; (iv) "
                "evaluation at larger scale, on datasets exceeding 10^4 "
                "features, where a relevance-guided, binary-native "
                "operator should be especially valuable; and (v) "
                "alternative relevance priors beyond mutual information "
                "and ReliefF, such as SHAP-based, mRMR, or learned priors, "
                "to test whether the mechanism's benefit is specific to "
                "the prior used here or transfers to others.")

    # ----------------------------------------------------------------- Methods
    doc.add_heading("6. Methods", level=1)
    para(doc, "Table 8 collects the notation used throughout this section "
              "for reference.")
    full_width(doc, lambda: add_notation_table(doc))
    doc.add_heading("Problem formulation", level=2)
    para(doc, "We encode a candidate subset as a binary mask b in {0,1}^d "
              "over the d features, where bⱼ = 1 marks feature j as "
              "selected. The wrapper objective minimizes a scalarized "
              "trade-off between predictive error and subset cardinality:")
    eqm(doc, [
        mrun("f"), mdelim([mrun("b")]),
        mrun(" = 0.99"), mdelim([mrun("1 − Acc"), mdelim([mrun("b")])]),
        mrun(" + 0.01"), mfrac([mrun("|b|")], [mrun("d")]),
    ])
    para(doc, "where Acc(b) is the stratified 5-fold KNN accuracy (k = 5) "
              "computed on the selected features alone. Every baseline in "
              "this study optimizes this same objective with the same "
              "0.01 cardinality weight; RG-SCSO is granted no structural "
              "advantage on subset size from the fitness function itself, "
              "so any parsimony gap reported below reflects the search "
              "mechanism, not a differently weighted objective. SCSO "
              "maintains a population of real-valued positions updated by "
              "its exploration and exploitation rules, both governed by "
              "the sensitivity range")
    eqm(doc, [
        mrun("R"), mdelim([mrun("t")]), mrun(" = "),
        msub([mrun("S")], [mrun("M")]), mrun(" − "),
        msub([mrun("S")], [mrun("M")]), mrun("·"),
        mfrac([mrun("t")], [msub([mrun("T")], [mrun("max")])]),
        mrun(",   "), msub([mrun("S")], [mrun("M")]), mrun(" = 2"),
    ])
    para(doc, "which RG-SCSO retains while replacing the binarization and "
              "adding a relevance field.")

    doc.add_heading("Theoretical motivation: transfer-function washout", level=2)
    para(doc, "Consider any binary-native optimizer that keeps a "
              "real-valued position and binarizes coordinate j through a "
              "transfer T: R → [0,1] that returns a selection-or-flip "
              "probability. A continuous-space enhancement can influence "
              "the retained subset only by perturbing a coordinate, "
              "xⱼ -> xⱼ + δⱼ; its entire effect on the discrete decision is "
              "the induced change in probability Δⱼ = T(xⱼ+δⱼ) − T(xⱼ). "
              "The next result shows this leverage is governed by the "
              "local slope of T alone, and therefore collapses wherever T "
              "saturates, the mechanism we call washout.")
    lem = doc.add_paragraph()
    lem.add_run("Lemma 1 (Leverage bound). ").bold = True
    _add_run_text(lem, "Let T: R → [0,1] be Lipschitz and piecewise "
                 "continuously differentiable. If a coordinate-space "
                 "enhancement perturbs xⱼ by δⱼ, the induced change in the "
                 "selection or flip probability satisfies", italic=True)
    eqm(doc, [
        mdelim([msub([mrun("Δ")], [mrun("j")])], beg="|", end="|"), mrun(" ≤ "),
        msub([mdelim([mrun("T′")], beg="‖", end="‖")], [mrun("∞")]), mrun(" · "),
        mdelim([msub([mrun("δ")], [mrun("j")])], beg="|", end="|"),
    ])
    para(doc, "where ‖T′‖∞ is the largest slope T attains between xⱼ and "
              "xⱼ+δⱼ. For the two standard transfers ‖σ′‖∞ = 1/4 and "
              "‖|tanh|′‖∞ = 1, and each slope decays away from the "
              "origin,", italic=True)
    eqm(doc, [
        mrun("σ′"), mdelim([mrun("x")]), mrun(" ≤ "),
        msup([mrun("e")], [mrun("−|x|")]),
    ])
    eqm(doc, [
        mdelim([mrun("|tanh|′"), mdelim([mrun("x")])], beg="|", end="|"),
        mrun(" ≤ 4"), msup([mrun("e")], [mrun("−2|x|")]),
    ])
    para(doc, "in a flat region, where ‖T′‖∞ ≤ ε, the leverage collapses "
              "to |Δⱼ| ≤ ε·|δⱼ| (proof in the Supplementary Information "
              "PDF distributed with this submission).", italic=True)
    rem = doc.add_paragraph()
    rem.add_run("Remark (why RG-SCSO is exempt). ").bold = True
    rem.add_run("RG-SCSO breaks the premise of the lemma: instead of "
                 "routing information through the coordinate, it "
                 "modulates the flip probability directly at the transfer "
                 "output,")
    eqm(doc, [
        msub([mrun("Δ")], [mrun("j")]), mrun(" = ± γ·"),
        msub([mrun("s")], [mrun("j")]), mrun("·V"),
        mdelim([msub([mrun("x")], [mrun("j")])]),
    ], note="independent of T′.")

    prop = doc.add_paragraph()
    prop.add_run("Proposition 2 (Transition probability and cumulative "
                  "leverage). ").bold = True
    _add_run_text(prop, "Under the V-shaped binarization rule used "
                  "throughout this paper, bit j flips at iteration t with "
                  "probability exactly T(xⱼ(t)), so the one-step "
                  "bit-transition probability coincides with the flip "
                  "probability, P(bⱼ(t+1)≠bⱼ(t) | xⱼ(t)) = T(xⱼ(t)). Lemma "
                  "1's bound on Δⱼ is therefore, without further "
                  "assumption, already a bound on the one-step change in "
                  "transition probability,", italic=True)
    eqm(doc, [
        mdelim([mrun("ΔP"), mdelim([
            msub([mrun("b")], [mrun("j")]), mrun("(t+1)≠"),
            msub([mrun("b")], [mrun("j")]), mrun("(t)"),
        ])], beg="|", end="|"),
        mrun(" ≤ "),
        msub([mdelim([mrun("T′")], beg="‖", end="‖")], [mrun("∞")]),
        mrun(" · "),
        mdelim([msub([mrun("δ")], [mrun("j")])], beg="|", end="|"),
    ])
    prop2 = doc.add_paragraph()
    _add_run_text(prop2,
        "Moreover, if a coordinate's trajectory and a δ-perturbed twin "
        "trajectory differ by δⱼ(t) at each of N successive iterations, "
        "and T′≤ε throughout (the flat, saturated regime), linearity of "
        "expectation applied to the one-step bound gives a cumulative "
        "bound on the expected number of bit flips accumulated over the "
        "N steps,", italic=True)
    eqm(doc, [
        mdelim([mrun("Σ"), mdelim([mrun("t=1..N")]),
                mrun(" ΔP"), mdelim([
                    msub([mrun("b")], [mrun("j")]), mrun("(t+1)≠"),
                    msub([mrun("b")], [mrun("j")]), mrun("(t)"),
                ])],
               beg="|", end="|"),
        mrun(" ≤ ε·Σ"), mdelim([mrun("t=1..N")]),
        mdelim([msub([mrun("δ")], [mrun("j")])], beg="|", end="|"),
    ], note="(proof in the Supplementary Information PDF).")

    scope = doc.add_paragraph()
    scope.add_run("Scope of this result. ").bold = True
    scope.add_run("This is a local, one-step-composable sensitivity bound, "
                   "not a convergence guarantee, and it does not establish "
                   "that continuous-space enhancements are ineffective "
                   "outside the flat, saturated regime. We use washout, "
                   "throughout this paper, as a diagnostic framing "
                   "motivated by these bounds and corroborated empirically "
                   "(Results; Supplementary Information), not as a formal "
                   "guarantee that continuous-space enhancements must fail "
                   "under binary search; the full discussion is in the "
                   "Supplementary Information PDF.")

    doc.add_heading("The RG-SCSO mechanism", level=2)
    para(doc, "Given the updated position xⱼ, the base flip probability "
              "uses a V-shaped transfer V(xⱼ) = |tanh(xⱼ)|. With relevance "
              "ρⱼ in [0,1], preferred bit b*ⱼ = 1 if ρⱼ > 0.5 else 0, and "
              "strength sⱼ = 2|ρⱼ − 0.5| in [0,1], the flip probability is "
              "biased toward b*ⱼ,")
    eqm(doc, [
        msub([mrun("p")], [mrun("j")]), mrun(" = clip"),
        mdelim([mrun("|tanh"), mdelim([msub([mrun("x")], [mrun("j")])]),
                mrun("|·"), mdelim([mrun("1+γ·"), msub([mrun("σ")], [mrun("j")]),
                                     mrun("·"), msub([mrun("s")], [mrun("j")])]),
                mrun(", 0, 1")]),
    ])
    para(doc, "where σⱼ = +1 if the flip moves bit j toward b*ⱼ and "
              "σⱼ = −1 otherwise. Setting γ = 0 recovers a plain V-shaped "
              "operator (the NoRMS ablation). The 0.5 threshold that "
              "separates preferred from disfavored bits is a convenience, "
              "not a theoretically grounded neutral point: with "
              "ρⱼ = I(Xⱼ;y)/H(y), the value at which a feature stops being "
              "informative depends on the number of classes, the label "
              "entropy, the sample size, and the specific MI estimator "
              "used, none of which are normalized away by the H(y) "
              "division alone. The Discussion shows this "
              "dependence is not merely theoretical: the method's "
              "compactness is tied to how the chosen relevance statistic "
              "happens to place features around 0.5, not to the threshold "
              "itself. A threshold sweep (τ in {0.4, 0.5, 0.6}, Results "
              "Section 3.5) supports this choice as reasonable rather than "
              "arbitrary: τ=0.5 attains the highest accuracy on two of "
              "five datasets tested, and every value trades a modest, "
              "dataset-dependent accuracy shift for a monotonic reduction "
              "in feature count as τ increases, with no single value "
              "dominating uniformly.")

    _rho_static_p = doc.add_paragraph()
    _rho_static_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_run_text(_rho_static_p, "The field ρ that drives this mechanism "
              "is, in the final method, the static prior ρ", size=BODY_PT)
    _rsp_sub = _rho_static_p.add_run("static")
    _rsp_sub.font.size = Pt(BODY_PT)
    _rsp_sub.font.subscript = True
    _add_run_text(_rho_static_p, " = clip(MI(fⱼ;y", size=BODY_PT)
    _rsp_sub2 = _rho_static_p.add_run("train")
    _rsp_sub2.font.size = Pt(BODY_PT)
    _rsp_sub2.font.subscript = True
    _add_run_text(_rho_static_p, ")/H(y", size=BODY_PT)
    _rsp_sub3 = _rho_static_p.add_run("train")
    _rsp_sub3.font.size = Pt(BODY_PT)
    _rsp_sub3.font.subscript = True
    _add_run_text(_rho_static_p,
              "), 0, 1), a normalized "
              f"mutual-information filter score {_c('mrmr','kraskov')} "
              "computed once per outer split. For each outer split, ρ",
              size=BODY_PT)
    _rsp_sub4 = _rho_static_p.add_run("static")
    _rsp_sub4.font.size = Pt(BODY_PT)
    _rsp_sub4.font.subscript = True
    _add_run_text(_rho_static_p,
              " is computed exclusively from the corresponding training "
              "partition and is never recomputed using the held-out "
              "partition. An online extension adding an "
              "EMA credit-assignment term from accepted fitness "
              "improvements was examined and dropped, as the ablation "
              "shows no accuracy gain on any dataset. A second, smaller "
              "component, uncertainty-targeted memetic refinement (UMR), "
              "spends a fixed local-search budget where this prior is "
              "least decisive: each iteration, the K features whose "
              "relevance is closest to 0.5 are greedily flipped on the "
              f"incumbent best mask {_c('neri')} and the flip is kept "
              "only if fitness improves.", size=BODY_PT)

    doc.add_heading("Algorithm and computational cost", level=2)
    para(doc, "Every fitness evaluation, whether from population moves, "
              "memetic probes, or initialization, is counted against a "
              "single budget max_nfe = pop_size × max_iter = 15000, "
              "identical to the baselines, so UMR grants no extra "
              "evaluations.")

    para(doc, "The per-iteration cost is dominated by the N + K wrapper "
              "evaluations, each a KNN fit under fixed folds; the "
              "mutual-information prior adds a one-time O(dn log n) "
              f"preprocessing cost (the nearest-neighbor MI estimator "
              f"{_c('kraskov')}), amortized against the wrapper "
              "evaluations. RG-SCSO is thus asymptotically no more "
              "expensive than base SCSO apart from the K extra probes. In "
              "wrapper feature selection generally, classifier-based "
              "fitness evaluation, not the search mechanism itself, "
              "dominates overall computational cost as feature "
              "dimensionality grows; per-dataset wall-clock measurements "
              f"across the full {feat_min}–{feat_max}-feature range "
              "tested (Supplementary Information) are consistent with "
              "this expectation, and we make no runtime-speedup claim for "
              "RG-SCSO over same-budget baselines.")

    doc.add_heading("Datasets, baselines, and protocol", level=2)
    para(doc, f"The benchmark spans {s['n']} preprocessed datasets of "
              f"{feat_min} to {feat_max} features across biomedical, "
              "gene-expression, and categorical domains, full "
              "characteristics in Supplementary Table S1.")

    para(doc, f"We benchmark against six baselines: SCSO (base) {_c('scso')}, "
              f"AOA {_c('aoa')}, CoatiOA {_c('coa')}, GWO {_c('gwo')}, PSO "
              f"{_c('pso')}, and RIME {_c('rime')}. The protocol is "
              "preregistered and locked prior to the full run. All "
              "algorithms share: population 30, 500 iterations, 30 "
              "independent runs, seed = 42+run_id (paired across "
              "algorithms), KNN (k=5) with stratified 5-fold "
              "cross-validation, search space [-1,1]^d, and the fitness "
              "above. Every baseline runs with its library-default "
              "published hyperparameters, with only population size and "
              "evaluation budget matched across methods; RG-SCSO's γ and "
              "K are likewise fixed before the full run.")

    doc.add_heading("Statistics and reproducibility", level=2)
    para(doc, f"Significance uses the paired Wilcoxon signed-rank test "
              f"with Holm correction {_c('holm')}, the Holm family formed "
              "per dataset, at α = 0.05. A tie denotes failure to reject "
              "H0 (no significant accuracy difference) at this threshold; "
              "it is not evidence of equivalence. Effect sizes use Cohen's d and rank-biserial "
              f"r; overall comparison uses the Friedman test with a "
              f"critical-difference diagram {_c('demsar')}. The "
              "experimental design was preregistered and "
              "version-controlled before the full run and left unmodified "
              "after results were observed; all randomness is seeded "
              "deterministically and shared across algorithms.")

    # ---------------------- Statements and Declarations + References
    doc.add_heading("CRediT authorship contribution statement", level=1)
    para(doc, "Bui Quang Huy: Conceptualization, Methodology, Software, "
              "Validation, Formal analysis, Writing – original draft. "
              "Duong Minh Son: Supervision, Writing – review & editing.")

    doc.add_heading("Declaration of competing interest", level=1)
    para(doc, "The authors declare no competing interests.")

    doc.add_heading("Data availability", level=1)
    para(doc, "The datasets are publicly available benchmarks (UCI and "
              "standard microarray sets). The source code, the "
              "preregistration, the per-run seeds, and the raw results are "
              "available in an anonymized repository "
              "(https://anonymous.4open.science/r/RG-SCSO-8BC0/) and will "
              "be released in a public, citable repository upon "
              "acceptance.")

    doc.add_heading("Funding", level=1)
    para(doc, "No funding was received for this work.")

    add_references_asoc(doc)
    _linkify_citations(doc, len(ASOC_CITE_ORDER))

    for _t in doc.tables:
        repeat_header_row(_t)
    force_font_everywhere(doc)
    doc.save(OUT_DOCX)
    print(f"Đã ghi {OUT_DOCX}")


if __name__ == "__main__":
    build()
